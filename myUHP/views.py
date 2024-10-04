from django.shortcuts import render

# Create your views here.
from io import BytesIO
from math import ceil
from django.contrib.auth.decorators import login_required, permission_required, user_passes_test
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from .models import Feature, Outputworkplan, Kpi, GsmWorkplan,Toptask,Operworkplan, Country, KpiAchieve, Units,IndividualReport,SurveyProject, SurveyDataset,Statutworkplan,Subscribers, DocSave,ReportSave,TypeMeeting, MeetingProject, MeetingDiscussion,BriefingProject,BriefingBackground,RiskIdentification
from django.contrib.auth.models import User
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from uhpDataBase import settings
from django.core.mail import send_mail
from django.contrib.sites.shortcuts import get_current_site
from django.template.loader import render_to_string, get_template
from django.utils.http import urlsafe_base64_decode, urlsafe_base64_encode
from django.utils.encoding import force_bytes, force_str
try:
    from django.utils.encoding import force_text
except:
    from django.utils.encoding import force_str as force_text
    
from . tokens import generate_token
from django.core.mail import EmailMessage, send_mail
from . forms import SubscribersForm, MailMessageForm, CountryForm, StatutForm, OutputForm, KpiForm, GsmwpForm, ToptaskForm,WorkplanForm, KpiAchieveForm,OperworkplanSearchForm, IndivReportForm,SurveyProjectForm, SurveyDatasetForm, DocSaveForm,CoverPageForm, UnitsForm, DateForm, ReportSaveForm,SelectUnitForm, TypeMeetingForm, MeetingProjectForm, MeetingDiscussionForm,BriefingProjectForm,BriefingBackgroundForm,RiskIdentificationForm, SelectSurveyForm, SelectMeetingForm
from django.http import HttpResponse, JsonResponse, FileResponse
from django.views.generic import DetailView, ListView, TemplateView, View
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.template.defaultfilters import urlencode
from django.urls import reverse_lazy, reverse
from .resources import SurveyResource
from tablib import Dataset
from openpyxl import Workbook
from openpyxl.styles import Color, PatternFill, Font, Border, Alignment, Border, Side
from openpyxl.styles import colors
from openpyxl.cell import Cell
from openpyxl.formatting.rule import ColorScaleRule, CellIsRule, FormulaRule
from openpyxl.styles.differential import DifferentialStyle
from openpyxl.formatting.rule import Rule
from openpyxl.utils import FORMULAE
import plotly.express as px
import pandas as pd 
import json
from django.db.models import Count, Avg, Min, Max, Sum, CharField, Value, Q, Case, When, F, When, OuterRef, CheckConstraint,  BooleanField
from django.utils import timezone
from django.db.models.functions import Upper, Length, Concat, Lower
import itertools
from datetime import datetime, time, date, timedelta

from xhtml2pdf import pisa
import re, os
from django.views.decorators.csrf import csrf_exempt
from django.views.generic.detail import BaseDetailView
from lockdown.decorators import lockdown
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns

import numpy as np
from matplotlib import pyplot as plt
from .pdf import htmlTopdf
import seaborn as sns 
import xlwt

from .decorators import group_required


#from bokeh.plotting import figure, show
#from bokeh.models import ColumnDataSource, CDSView, GroupFilter
def is_visitors(user):
    return user.groups.filter(name='Visitors').exists()

# CREATE NEW COVER PAGE------------------------------------------

def coverPage(request):
    context={}
    return render(request, 'cover/index.html', context)

# CREATE VIEWS FOR COVER PAGE---------------------------------------

def index(request):   
    features = Feature.objects.all()
    feature1=features[0]
    feature2=features[1]
    feature3=features[2]
    feature4=features[3]
    feature5=features[4]
    feature6=features[5]
    feature7=features[6]
    feature8=features[7]
    feature9=features[8]
    feature10=features[9]
    feature11=features[10]
    feature12=features[11]
    feature13=features[12]
    feature14=features[13]
    feature15=features[14]
    feature16=features[15]
    feature17=features[16]
    feature18=features[17]
    feature19=features[18]
    feature20=features[19]
    feature21=features[20]

    data = {'feature1' : feature1, 'feature2' : feature2, 'feature3' : feature3, 'feature4' : feature4, 'feature5' : feature5, 'feature6' : feature6, 'feature7' : feature7,
            'feature8' : feature8, 'feature9' : feature9, 'feature10' : feature10, 'feature11' : feature11, 'feature12' : feature12, 'feature13' : feature13, 'feature14' : feature14,
            'feature15' : feature15, 'feature16' : feature16, 'feature17' : feature17, 'feature18' : feature18, 'feature19' : feature19, 'feature20' : feature20, 'feature21' : feature21
            } 
    
    return render(request,'index.html', data)

# UPDATE ITEMS IN COVER PAGE------------------
def is_valid_queryparam(param):
    return param!='' and param is not None
# display cover page items in Datatable
@lockdown()
def cover_list(request):
    qs = Feature.objects.order_by('name')
    
    name = request.GET.get('name')
    details = request.GET.get('details')
    
    request.session['name'] = name
    request.session['details'] = details
    
    if is_valid_queryparam(name):
        qs = qs.filter(name_icontains = name)
        
    if is_valid_queryparam(details):
        qs = qs.filter(details = details)
    
    page = request.GET.get('page', 1)
    paginator = Paginator(qs, 30)
    
    try:
        qs = paginator.page(page)
    except PageNotAnInteger:
        qs = paginator.page(1)
    except EmptyPage:
        qs = paginator.page(paginator.num_pages)
        
    context = {
        'items_list': qs,
        'name': name,
        'details': details
    }
    
    return render(request, 'pages/coverPages/cover_index.html', context)

# add new items in cover page

# Method for add new items for cover Page
@permission_required('myuhp.add_feature',raise_exception=True)
def add_items_cover_page(request):
    if request.method == "POST":
        form = CoverPageForm(request.POST)
        if form.is_valid():
            form.save(commit=True)
            messages.success(request, "Items for cover page Added Successfully")
            return redirect('cover_list')
        else:
            print("ERROR FORM INVALID")
            return JsonResponse({'msg': 'ERROR FORM INVALID'})
    else:
        context = {
            'form': form
        }
    return render(request, 'pages/coverPages/add_covertPageItems_modal.html', context)

# Method for edit items for cover Page
@permission_required('myuhp.change_feature',raise_exception=True)
def edit_items_cover_page(request, pk): 
    features = Feature.objects.get(id=pk)
    form = CoverPageForm(instance=features)
    
    if request.method == "POST":
        form = CoverPageForm(request.POST, instance=features)
        if form.is_valid():
            form.save()
            messages.success(request, "Items for covert page Edit Successfully")
            return redirect('cover_list')
    messages.warning(request, "ERROR FORM INVALID") 
    
    return render(request, 'pages/coverPages/edit_covertPageItems_modal.html', {'form': form})

# Method for delete items in the cover page
@permission_required('myuhp.delete_feature',raise_exception=True)
def delete_items_cover_page(request,pk):
    data=Feature.objects.filter(id=pk)
    data.delete()
    messages.warning(request, 'Items for covert page deleted successfully')
   
    return redirect('cover_list')

# export datatable to excel

# CREATE VIEWS FOR USER ACCOUNT---------------------------------------- 

def signup(request):
    if request.method=="POST":
        username=request.POST.get('username')
        fname=request.POST.get('fname')
        lname=request.POST.get('lname')
        email=request.POST.get('email')
       
       
      #  username=request.POST['username']
      #  fname= request.POST['fname']
      #  lname= request.POST['lname']
      #  email= request.POST['email']
        pass1= request.POST['pass1']
        pass2= request.POST['pass2']
        
        if User.objects.filter(username = username):
            messages.error(request, "Username already exist. Please try some other username")
            #return redirect('index')
           # return render(request,'signup.html')
            return render(request,'cover/signup.html')
        
        if User.objects.filter(email = email):
            messages.error(request, "Email address already registered. Please try some other email")
           # return redirect('signup')
            return render(request,'cover/signup.html')
        
        if len(username)>10:
            messages.error(request,"Username must be under 10 characters")
            return render(request,'cover/signup.html')
            
        if pass1 != pass2:
            messages.error(request, "Passwords didn't match")
            return render(request,'cover/signup.html')
        if not username.isalnum():
            messages.error(request, "Username must be Alpha-numeric")
           # return redirect('signup')
            return render(request,'cover/signup.html')
        
        myuser = User.objects.create_user(username, email, pass1)
        myuser.first_name = fname
        myuser.last_name = lname
        myuser.is_active = False
        
        myuser.save()
        messages.success(request, "Your Account has been successfully created")
        
      #  # Welcome email
      #  subject = "Welcome to the CLUSTER AFRO UHP data center"
      #  message = "Dear" + myuser.first_name + "!!\n" + "Welcome to the CLUSTER AFRO UHP data center \n Thank you for visiting our website \n We have also sent you a confirmation email, please confirm your email address in oder to activate your account. \n\n Thanking you \n Wilson Ondon"
      #  from_email=settings.EMAIL_HOST_USER
      #  to_list = [myuser.email]
      #  send_mail(subject, message, from_email, to_list, fail_silently = True)
        
      #  # Email Address Confirmation Email
      #  current_site = get_current_site(request)
      #  email_subject = "Confim your @ UHP - Website Login!!"
      #  message2 = render_to_string('email_confirmation.html',{
      #      'name':myuser.first_name,
      #      'domain':current_site.domain,
      #      'uid':urlsafe_base64_encode(force_bytes(myuser.pk)),
      #      'token': generate_token.make_token(myuser)
      #  })
      #  email = EmailMessage(
      #      email_subject,
      #      message2,
      #      settings.EMAIL_HOST_USER,
      #      [myuser.email]
      #  )
      #  email.fail_silency = True
      #  email.send()
        
        return redirect('signin') #   return redirect('index')
    return render(request,'cover/signup.html') # return render(request,'signup.html')

def signin(request):
    if request.method == 'POST':
        username = request.POST['username']
        pass1 = request.POST['pass1']
        
        user = authenticate(username = username, password = pass1)
        
        if user is not None:
            login(request, user)
            fname = user.first_name
            messages.success(request, "Your're successfully logged in.")
            
            return redirect('coverPage') # return redirect('dashboard')
            #return render(request, 'index.html',{'fname':fname})
        else:
            messages.error(request, 'You are not logged in')
            #return redirect('signin')
            return redirect('coverPage') #  return redirect('index')
    return render(request,'cover/signin.html') # return render(request,'signin.html')

def signout(request):
    logout(request)
    messages.success(request, 'logged out successfully')
    
    return redirect('/')

def activate(request, uidb64, token):
    try:
        uid = force_str(urlsafe_base64_decode(uidb64))
        myuser = User.objects.get(pk = uid)
    except (TypeError, ValueError, OverflowError, User.DoesNotExist):
        myuser = None
        
        if myuser is not None and generate_token.check_token(myuser, token):
           myuser.is_active = True 
           myuser.save()
           messages.success(request, "Votre compte a ete actived felicitation contecte")
          # login(request, myuser)
           return redirect('/')
        else:
            messages.error(request, 'activation lost!!!!')
            return redirect('/')
    return render(request, 'activate_failed.html')

# CREATE VIEWS FOR SEND EMAIL TO THE NEWSLETTERS SUBSCRIBERS -------------------------

def counter(request):
    text=request.POST['text']
    amount_of_words = len(text.split())
    return render(request, 'counter.html', {'amount':amount_of_words})

def subscribers(request):
    if request.method == 'POST':
        form = SubscribersForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Your Subscription has been successfully')
            return redirect('coverPage')
           # return HttpResponse('<h1>Hello HttpResponse</h1>')   
    else:
        form = SubscribersForm()
        context = {
            'form': form
        }
        return render(request, 'cover/index.html', context)
    
def mail_letter(request):
 #   emails = Subscribers.objects.all()
 #   de = read_frame(emails, fieldnames =['email_subscriber'])
 #   mail_list = df['']
    if request.method == 'POST':
        form = MailMessageForm(request.POST)
        if form.is_valid():
            form.save()
            name=form.cleaned_data.get('title_mail'),
            email='ondonwil@gmail.com',
          #  message += "\n\n{0}".format(form.cleaned_data.get('message_mail'))
    
          #  title = form.cleaned_data.get('title_mail')
          #  message = form.cleaned_data.get('message_mail')
            send_mail(
                   subject=form.cleaned_data.get('message_mail').strip(),
                    message="\n\n{0}".format(form.cleaned_data.get('message_mail')),
                    from_email='ondonwil@gmail.com',
                    recipient_list=[settings.EMAIL_HOST_USER],
        
              #     title,
              #     message,
              #      'ondonwil@gmail.com',
              #      ['ondonwil@gmail.com'],
              #      fail_silently=False
                )
            messages.success(request, 'Message has been send to the Mail list')
            return redirect('mail_letter')
    else:
        form = MailMessageForm()
        
    form=MailMessageForm()
    context = {
        'form':form
        
    }
    return render(request, 'cover/contact.html', context)

# CREATE VIEWS FOR DASHBOARD -------------------------
@user_passes_test(is_visitors)
def dashboard(request, *args, **kwargs):
    operworkplan=Operworkplan.objects.all()
    
  #  dt_fig_trend = Statutworkplan.objects.annotate(
  #      total_planned = Count('statuts_operwork__sub_activity', filter=Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__unit__unit_code__isnull=False))    
  #  ).values('statut_name','statuts_operwork__completion_date','total_planned').order_by('statuts_operwork__completion_date')

    dt_fig_trend = Statutworkplan.objects.annotate(
        total_planned = Count('statuts_operwork__sub_activity', filter=Q(statuts_operwork__sub_activity__isnull=False))    
    ).values('statut_name','statuts_operwork__completion_date','total_planned').order_by('statuts_operwork__completion_date')
    
    
    
# operworkplan_count=Operworkplan.objects.annotate(nb=Count('sub_activity')).values('nb')
  #  operworkplan_count=Operworkplan.objects.annotate(
  #      nb=Case(
  #          When(Q(completion_date__isnull=False)|Q(statut_name__isnull=False), then=Value("1")),
  #          default=Value("N/A")
  #      ))
    
 #    first_act=Operworkplan.objects.aggregate(first_acti_date=Min('completion_date'))['first_acti_date']
 #    last_act=Operworkplan.objects.aggregate(last_acti_date=Max('completion_date'))['last_acti_date']
 #    dates=[]
  #   count=itertools.count()
    
 #    while(dt:= first_act + timezone.timedelta(days=10*next(count))) <=last_act:
 #       dates.append(dt)
        
 #    whens = [
      # When(completion_date__range = (dt, dt+timezone.timedelta(days=10)), then=Value(dt.date()))
      #  for dt in dates
 #    ]
    
#     case=Case(
#         *whens,
#         output_field = CharField()
 #    )
    
 #    operworkplan_count = Operworkplan.objects.annotate(
#         daterange=case
 #    ).values('completion_date').annotate(total_act=Count('sub_activity'))
     
 #   completionDate = Operworkplan.objects.values_list('completion_date', flat=True).distinct()
  #  count_value=Operworkplan.objects.aggregate(total= Count('id'))
  #  avg_value=Operworkplan.objects.aggregate(avg= Avg('vc_amount')) # calculate average
   # avg_wilson=Operworkplan.objects.filter(responsable__startswith='wilson').select_related('statut_name').aggregate(avg_wil= Avg('vc_amount')) # calculate average
  #  result_vc_amount=Operworkplan.objects.aggregate(
   #     total= Count('id'),
   #     min= Min('vc_amount'),
   #     max= Max('vc_amount'),
   #     sum= Sum('vc_amount'),
   #     avg= Avg('vc_amount')
   #     ) 
    one_month_ago = timezone.now() - timezone.timedelta(days=31) # à partir de J-31 qus'au delà
  #  operworkplans= Operworkplan.objects.filter(completion_date__gte=one_month_ago)
    
    #Annotation functions
   # length_subActivity=Operworkplan.objects.annotate(len_subactivity=Length('sub_activity')).filter(len_subactivity__gte=10)
   # anno01=length_subActivity.first().len_subactivity
    
   # anno02=length_subActivity.values('sub_activity','len_subactivity')
   #  anno02=length_subActivity.values('sub_activity','len_subactivity')
    
    ## Unit provient de la table parent GsmWorkplan et vc_amount provient de la table enfant Operworkplan lié à partir de la clé gsmWorkplan_operw 
    
    ##     # number od sub activity by unit
     ## Unit provient de la table parent GsmWorkplan et vc_amount provient de la table enfant Operworkplan lié à partir de la clé gsmWorkplan_operw 
   # concatenation_2 =Concat(
    #    'unit_code', Value('[Amount:'), Count('units_gsmWorkplan__gsmWorkplan_operw__vc_amount'), Value(']'),
    #    output_field=CharField()
    #    )
    
   # operationnals_2=Units.objects.annotate(messages=concatenation_2)
    
    #nombre total sub activity for VID
   # units_subactivity=Units.objects.annotate(total_subactivity=Count('units_gsmWorkplan__gsmWorkplan_operw__sub_activity')).filter(Q(unit_code__gte='VID'))
   # avg_unit_vid=units_subactivity.aggregate(avg_acti=Avg('total_subactivity'))
    
    # PERFORMANCE OF SUB ACTIVITIES IN ONE MONTH AGO-----------------------------------------------------
    totalSubActivities = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False)).count()
    totalSubActivity =int(Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__completion_date__lte=one_month_ago)).count())
    statutworkplans=Statutworkplan.objects.filter(statuts_operwork__completion_date__lte=one_month_ago).annotate(
        total_statut=Count('statuts_operwork__sub_activity'),
        percent_statut=Count('statuts_operwork__sub_activity')*100/totalSubActivity).values(
        'statut_name', 'total_statut','percent_statut'
    ).order_by('-percent_statut')
    
    title_one_month_ago ='Performance of',totalSubActivity,' Sub activities for one month ago'
     
    # SUB ACTIVITY BY COUNTRY SUPORTED---------------------------------
    countries=Country.objects.filter(countries_operwork__isnull=False).annotate(
        num_acti=Count('countries_operwork'),
        percent_acti =Count('countries_operwork')*100/totalSubActivities).values(
        'country_name','num_acti','percent_acti'
    ).order_by('-num_acti')
      
    # number of sub activity by outputs
    # outpts_subactivity=Outputworkplan.objects.filter(Q(outputs_gsmWorkplan__isnull=False)).values('output_code').annotate(num_subactivity=Count('outputs_gsmWorkplan__gsmWorkplan_operw'))
   # Operworkplan['OrderPeriod'].apply(lambda x: x.strftime('%Y-%m'))
    #Operworkplan['OrderPeriod'] = Operworkplan.completion_date.strftime('%Y-%m')
    
    # SUBACTIVITY NOT STARTED-------------------------
    list_not_start =Statutworkplan.objects.filter(Q(statut_name__icontains='Not Started') & Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__completion_date__lte=one_month_ago)).values(
        'statuts_operwork__sub_activity','statuts_operwork__responsable','statuts_operwork__completion_date','statuts_operwork__comments'
    ).order_by('-statuts_operwork__completion_date')
    #list_subactivity=Operworkplan.objects.values('sub_activity','responsable', 'completion_date', 'comments')
    #date_col =pd.DatetimeIndex(operworkplan['completion_date'])
    #operworkplan['year']=date_col.year
      
   # completeStart=Q(statut_name__icontains='Completed')|Q(statut_name__icontains='On Track')
   # notStart=Q(statut_name__icontains='Not Started')|Q(statut_name__icontains='Stalled')
   # recently_completion=Q(statuts_operwork__completion_date__gt=timezone.now()+timezone.timedelta(days=31))
   # activity_notStart=Statutworkplan.objects.filter(notStart|recently_completion|Q(statuts_operwork__sub_activity__isnull=False)).annotate(
   #     total_statut=Count('statuts_operwork__vc_amount')).values(
   #     'statut_name', 'total_statut'
   # )
        
   # activity_notStartIssue=Statutworkplan.objects.filter(notStart|recently_completion|Q(statuts_operwork__sub_activity__isnull=False)).values(
   #     'statuts_operwork__sub_activity',
   #     'statuts_operwork__responsable',
   #     'statuts_operwork__completion_date'
   #     )
     
    # creater fig in bar
   # opert=Operworkplan.objects.all()
   # activities=opert.values_list('sub_activity', flat=True).distinct()
   # goals={}
   # for activitie in activities:
   #     activi_oper=opert.filter(Q(sub_activity=activitie))
   #     total_goals=sum([fix.vc_amount for fix in activi_oper] )
   #     goals[activitie]=total_goals
   # goals = dict(sorted(goals.items(), key=lambda x:x[1], reverse=True))
   # fig2 = px.bar(x=goals.keys(), y=goals.values(), title="Wilson Ondon")
    ## end of fig in bar-------------------------
    
   # sub_activites = Operworkplan.objects.values_list('sub_activity', flat=True).distinct()
   # sub_activity = request.GET.get('sub_activity')
   # units = Units.objects.prefetch_related('units_gsmWorkplan').distinct()
   # units_gsmWorkplan = request.GET.get('units_gsmWorkplan')
    
  #  operworkplans=Operworkplan.objects.filter(units_gsmWorkplan=units_gsmWorkplan).order_by('completion_date')
  #  sub_activity_date =[d.completion_date for d in operworkplans]
  #  sub_activity_oper =[d.gsmWorkplan for d in operworkplans]
    
  #   cds= ColumnDataSource(data=dict(sub_activity_date=sub_activity_date,sub_activity_oper=sub_activity_oper))
    
  #  fig=figure(height=500, title=f"{units_gsmWorkplan} Operworkplan")
    
    # Create chart
    start =  request.GET.get('start')
    end =  request.GET.get('end')
    
    activity_completedStartCount={}
    activity_completedStart={}
    
    if start:
        dt_fig_trend = dt_fig_trend.filter(statuts_operwork__completion_date__gte=start)
        
      # activity_completedStartCount=Statutworkplan.objects.filter(completeStart|recently_completion|Q(statuts_operwork__sub_activity__isnull=False)|Q(statuts_operwork__completion_date__gte=start)).annotate(
      #  total_statut=Count('statuts_operwork__vc_amount')).values(
      #  'statut_name', 'total_statut')
      #  activity_completedStart=Statutworkplan.objects.filter(completeStart|recently_completion|Q(statuts_operwork__sub_activity__isnull=False)|Q(statuts_operwork__completion_date__gte=start)).values(
      #      'statuts_operwork__sub_activity',
      #      'statuts_operwork__responsable',
      #      'statuts_operwork__completion_date')
    
    if end:
       dt_fig_trend = dt_fig_trend.filter(statuts_operwork__completion_date__lte=end)
        
      #  activity_completedStartCount=Statutworkplan.objects.filter(completeStart|recently_completion|Q(statuts_operwork__sub_activity__isnull=False)|Q(statuts_operwork__completion_date__gte=end)).annotate(
      #  total_statut=Count('statuts_operwork__vc_amount')).values(
      #  'statut_name', 'total_statut')
        
      #  activity_completedStart=Statutworkplan.objects.filter(completeStart|recently_completion|Q(statuts_operwork__sub_activity__isnull=False)|Q(statuts_operwork__completion_date__gte=end)).values(
      #      'statuts_operwork__sub_activity',
      #      'statuts_operwork__responsable',
      #      'statuts_operwork__completion_date')
    fig = px.line(dt_fig_trend, x='statuts_operwork__completion_date', y='total_planned', text='total_planned', color='statut_name',labels={'total_planned':'Total of sub activity', 'statuts_operwork__completion_date':'Date', 'statut_name':'Statut name'})
   
        # fig = px.line(
        #     x = [c.statuts_operwork__completion_date for c in dt_fig_trend],
        #     y = [c.total_planned for c in dt_fig_trend],
        #     color=[c.statut_name for c in dt_fig_trend],
        #     text=[c.total_planned for c in dt_fig_trend],
                
        #     title='Amount by date',
        #     labels={'x': 'Date', 'y':'Amount', 'color':'Statut name'}
        # )
   # fig.update_layout(title={
   #     'font_size': 22,
   #     'xanchor': 'center',
   #     'x': 0.5
   # })
    fig.update_traces(textposition="bottom right")
    chart = fig.to_html()
    
    ##0010 count sub activity by month----------------------------------------------------
    subactivityByMonth= Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False)).values('statuts_operwork__completion_date__month','statut_name').annotate(nb_sub_activity=Count('statuts_operwork__sub_activity'))
   # xfig=subactivityByDate.values_list('statuts_operwork__completion_date__month', flat=True)
   # yfig=subactivityByDate.values_list('nb_sub_activity', flat=True)
   # text={f'{nb_sub_activity:.0f}' for nb_sub_activity in yfig}
    fig_statut_month = px.bar(subactivityByMonth, x='statuts_operwork__completion_date__month', y='nb_sub_activity', text='nb_sub_activity', color='statut_name',labels={'nb_sub_activity':'Total of sub activity', 'statuts_operwork__completion_date__month':'Month (value)'})
   # fig03 = px.line(
   #     x = 'xfig',
   #     y = 'yfig', 
   #     text = text
   #     )
    fig_statut_month.update_layout(
          title_text = 'Number of sub activities per month'
          #xaxis_tickangle=90
      )
    fig_statut_month.update_traces(textfont_size=12, textposition="outside", cliponaxis=False)
    chart_statut_month = fig_statut_month.to_html()
    
     ##0011 count sub activity by Unit----------------------------------------------------
    subactivityByUnit= Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False)).values('statuts_operwork__gsmWorkplan__toptask__unit__unit_code','statut_name').annotate(nb_sub_activity=Count('statuts_operwork__sub_activity'))
    fig_statut_unit = px.bar(subactivityByUnit, x='statuts_operwork__gsmWorkplan__toptask__unit__unit_code', y='nb_sub_activity', text='nb_sub_activity', color='statut_name', labels={'nb_sub_activity':'Total of sub activity', 'statuts_operwork__gsmWorkplan__toptask__unit__unit_code':'Unit code', 'statut_name':'Statut name'})
    fig_statut_unit.update_layout(
        title_text = 'Number of sub activities per unit',
         barmode='group',
         xaxis={'categoryorder':'array', 'categoryarray':['CHE','HPD','NUT','TNR','VID','UHU']}
    )
    fig_statut_unit.update_traces(textfont_size=12, textposition="outside", cliponaxis=False)
    chart_statut_unit = fig_statut_unit.to_html()
    
    
     ##0012 count sub activity by Top Task----------------------------------------------------
    subactivityByTopTask= Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False)).values('statuts_operwork__gsmWorkplan__toptask__top_task_short','statuts_operwork__gsmWorkplan__toptask__unit__unit_code','statut_name').annotate(nb_sub_activity=Count('statuts_operwork__sub_activity'))
    fig_statut_topTask = px.bar(subactivityByTopTask, x='statuts_operwork__gsmWorkplan__toptask__top_task_short', y='nb_sub_activity', text_auto=True,hover_data=['statuts_operwork__gsmWorkplan__toptask__unit__unit_code'], color='statut_name',labels={'nb_sub_activity':'Total of sub activity', 'statuts_operwork__gsmWorkplan__toptask__top_task_short':'Top task', 'statut_name':'Statut name'})
    fig_statut_topTask.update_layout(
        title_text = 'Number of sub activities per Top task',
         xaxis_tickangle=90
    )
    fig_statut_topTask.update_traces(textfont_size=12, textposition="outside", cliponaxis=False)
    chart_statut_topTask = fig_statut_topTask.to_html()
    
    ## LIST BY UNIT FOR OUTPUT,KPI, TOP TASK, LOWEST TASK and SUB ACTIVITIES-----------------------
    table_outputKpiTask = Units.objects.exclude(units_toptask__toptask_gsmWorkplan__lowest_task_short__isnull=True).annotate(
        total_output = Count('units_toptask__output__output_code',distinct=('units_toptask__output__output_code'), filter=Q(units_toptask__output__isnull=False)),
        total_kpi = Count('units_toptask__kpi',distinct=('units_toptask__kpi'), filter=Q(units_toptask__kpi__isnull=False)),
        total_TopTask = Count('units_toptask__top_task_short',distinct=('units_toptask__top_task_short'), filter=Q(units_toptask__top_task_short__isnull=False)),
        total_lowest = Count('units_toptask__toptask_gsmWorkplan__lowest_task_short',distinct=('units_toptask__toptask_gsmWorkplan__lowest_task_short'), filter=Q(units_toptask__toptask_gsmWorkplan__lowest_task_short__isnull=False)),
        total_subactivities = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity',distinct=('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity'), filter=Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False))
    ).values('unit_code','total_output','total_kpi','total_TopTask','total_lowest','total_subactivities').distinct()
   
    
    # SUB ACTIVITIES DISTRIBUTED BY UNITS AND IMPLEMENTATION STATUS-----------------------------------
  #  testTble2 = Statutworkplan.objects.exclude(statuts_operwork__gsmWorkplan__isnull=False).annotate(
  #      total_lowest = Count('statuts_operwork__gsmWorkplan__lowest_task_short', filter=Q(statuts_operwork__gsmWorkplan__isnull=False)),
        
 #       total_planned = Count('statuts_operwork__sub_activity', filter=Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__unit__unit_code__isnull=False)),
 #       total_completed = Count('statuts_operwork__sub_activity', filter=Q(statut_name__icontains='Completed') & Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__unit__unit_code__isnull=False)),
 #       total_OnTrack = Count('statuts_operwork__sub_activity', filter=Q(statut_name__icontains='On Track') & Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__unit__unit_code__isnull=False)),
 #       total_Not_Started = Count('statuts_operwork__sub_activity', filter=Q(statut_name__icontains='Not Started') & Q(statuts_operwork__sub_activity__isnull=False)& Q(statuts_operwork__gsmWorkplan__unit__unit_code__isnull=False)),
 #       total_Issues = Count('statuts_operwork__sub_activity', filter=Q(statut_name__icontains='Stalled') & Q(statuts_operwork__sub_activity__isnull=False)& Q(statuts_operwork__gsmWorkplan__unit__unit_code__isnull=False))
 #   ).values('statuts_operwork__gsmWorkplan__unit__unit_code','total_lowest','total_planned','total_completed','total_OnTrack','total_Not_Started','total_Issues').distinct()
   
    tables_subactivities = Units.objects.exclude(units_toptask__toptask_gsmWorkplan__lowest_task_short__isnull=True).annotate(  
        total_lowest = Count('units_toptask__toptask_gsmWorkplan__lowest_task_short', distinct=('units_toptask__toptask_gsmWorkplan__lowest_task_short'),filter=Q(units_toptask__toptask_gsmWorkplan__lowest_task_short__isnull=False)),
        total_planned = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity',distinct=('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity'),  filter=Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False)),
        total_completed = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity',distinct=('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity'),  filter=Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name__icontains='Completed') & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False) & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False)),
        total_OnTrack = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', distinct=('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity'), filter=Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name__icontains='On Track') & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False) & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False)),
        total_Not_Started = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity',distinct=('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity'),  filter=Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name__icontains='Not Started') & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False)& Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False)),
        total_Issues = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', distinct=('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity'), filter=Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name__icontains='Stalled') & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False)& Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False))
    ).values('unit_code','total_lowest','total_planned','total_completed','total_OnTrack','total_Not_Started','total_Issues').distinct()
   
      ##0012 count sub activity by Countries and statut name----------------------------------------------------
    subactivityByCountries= Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False)).values('statut_name','statuts_operwork__country__country_name').annotate(nb_sub_activity=Count('statuts_operwork__sub_activity')).exclude(statuts_operwork__country__country_name__in ='All Countries')
    fig_statut_countries = px.bar(subactivityByCountries, x='statuts_operwork__country__country_name', y='nb_sub_activity', text='nb_sub_activity', color='statut_name',labels={'nb_sub_activity':'Total of sub activity', 'statut_name':'Name of statut', 'statuts_operwork__country__country_name':'Country'})
    fig_statut_countries.update_layout(
        title_text = 'Number of sub activities per country and statut',
        uniformtext_minsize=8, 
        uniformtext_mode='hide',
        barmode='stack',
        xaxis_tickangle=90,
        legend=dict(
        x=0,
        y=1.0,
        bgcolor='rgba(255, 255, 255, 0)',
        bordercolor='rgba(255, 255, 255, 0)'
    ),
   
    bargap=0.15, # gap between bars of adjacent location coordinates.
    bargroupgap=0.1 # gap between bars of the same location coordinate.
    )
    fig_statut_countries.update_traces(textfont_size=12, textangle=0, textposition="auto", cliponaxis=False)
    chart_statut_countries = fig_statut_countries.to_html() # Put on the dahsboard html the graph with title : PERFORMANCE BY COUNTRY SUPORTED
    
    ## Value for inbox 
    vbox_totalSubActivity = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False)).count()
    vbox_notStartIssue = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='Stalled')).count()
    vbox_notStart = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='Not Started')).count()
    vbox_onTrack = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='On Track')).count()
    vbox_completed = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='Completed')).count()
    vbox_completedPercent=(vbox_completed/vbox_totalSubActivity)*100
    vbox_onTrackPercent=(vbox_onTrack/vbox_totalSubActivity)*100
    vbox_vbox_notStartPercent=(vbox_notStart/vbox_totalSubActivity)*100
    vbox_notStartIssuePercent=(vbox_notStartIssue/vbox_totalSubActivity)*100
    # VID
    vbox_total_VID = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains='VID')).count()
    vbox_completed_VID = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='Completed')& Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains='VID')).count()
    vbox_completed_VIDPercent=round((vbox_completed_VID/vbox_total_VID)*100,1) 
    #CHE
    vbox_total_CHE = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains='CHE')).count()
    vbox_completed_CHE = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='Completed') & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains='CHE')).count()
    vbox_OnTrack_CHE = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='On Track') & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains='CHE')).count()
    vbox_completed_CHEPercent=round(((vbox_completed_CHE+vbox_OnTrack_CHE)/vbox_total_CHE)*100,1) 
    
    # HPD
    vbox_total_HPD = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains='HPD')).count()
    vbox_completed_HPD = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='Completed') & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains='HPD')).count()
    vbox_OnTrack_HPD = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='On Track') & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains='HPD')).count()
    vbox_completed_HPDPercent=("{:0.2f}".format(((vbox_completed_HPD+vbox_OnTrack_HPD)/vbox_total_HPD)*100))    
    
     # NUT
    vbox_total_NUT = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains='NUT')).count()
    vbox_completed_NUT = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='Completed') & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains='NUT')).count()
    vbox_OnTrack_NUT = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='On Track') & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains='NUT')).count()
    vbox_completed_NUTPercent=round(((vbox_completed_NUT+vbox_OnTrack_NUT)/vbox_total_NUT)*100,1) 
    
    # TNR
    vbox_total_TNR = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains='TNR')).count()
    vbox_completed_TNR = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='Completed') & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains='TNR')).count()
    vbox_OnTrack_TNR = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='On Track') & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains='TNR')).count()
    vbox_completed_TNRPercent=round(((vbox_completed_TNR+vbox_OnTrack_TNR)/vbox_total_TNR)*100,1) 
    
     # UHU
    vbox_total_UHU = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains='UHU')).count()
    vbox_completed_UHU = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='Completed') & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains='UHU')).count()
    vbox_OnTrack_UHU = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='On Track') & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains='UHU')).count()
    vbox_completed_UHUPercent=round((vbox_completed_UHU+vbox_OnTrack_UHU)/vbox_total_UHU*100,1) 
    
     # OUTPUT 1.1.1
    total_111 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='1.1.1')).count()
    vbox_completed_111 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='Completed') & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='1.1.1')).count()
    vbox_OnTrack_111 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='On Track') & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='1.1.1')).count()
    output_111=vbox_completed_111+vbox_OnTrack_111
    if total_111 ==0:
            vbox_completed_111Percent=(output_111/vbox_totalSubActivity)*100  
    else: 
            vbox_completed_111Percent=(output_111/total_111)*100
   
    
     # OUTPUT 3.1.1
    total_311 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.1.1')).count()
    vbox_completed_311 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='Completed') & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.1.1')).count()
    vbox_OnTrack_311 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='On Track') & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.1.1')).count()
    output_311=vbox_completed_311+vbox_OnTrack_311
    if total_311 ==0:
            vbox_completed_311Percent=(output_311/vbox_totalSubActivity)*100  
    else: 
            vbox_completed_311Percent=(output_311/total_311)*100
    
    # OUTPUT 3.1.2
    total_312 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.1.2')).count()
    vbox_completed_312 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='Completed') & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.1.2')).count()
    vbox_OnTrack_312 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='On Track') & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.1.2')).count()
    output_312=vbox_completed_312+vbox_OnTrack_312
    if total_312 ==0:
            vbox_completed_312Percent=(output_312/vbox_totalSubActivity)*100  
    else: 
            vbox_completed_312Percent=(output_312/total_312)*100
 
    
     # OUTPUT 3.2.1
    total_321 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.2.1')).count()
    vbox_completed_321 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='Completed') & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.2.1')).count()
    vbox_OnTrack_321 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='On Track') & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.2.1')).count()
    output_321=vbox_completed_321+vbox_OnTrack_321
    if total_321 ==0:
            vbox_completed_321Percent=(output_321/vbox_totalSubActivity)*100  
    else: 
            vbox_completed_321Percent=(output_321/total_321)*100
   
    
      # OUTPUT 3.2.2
    total_322 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.2.2')).count()
    vbox_completed_322 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='Completed') & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.2.2')).count()
    vbox_OnTrack_322 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='On Track') & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.2.2')).count()
    output_322=vbox_completed_322+vbox_OnTrack_322
    if total_322 ==0:
            vbox_completed_322Percent=(output_322/vbox_totalSubActivity)*100  
    else: 
            vbox_completed_322Percent=(output_322/total_322)*100
  
    
     # OUTPUT 3.3.1
    total_331 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.3.1')).count()
    vbox_completed_331 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='Completed') & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.3.1')).count()
    vbox_OnTrack_331 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='On Track') & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.3.1')).count()
    output_331=vbox_completed_331+vbox_OnTrack_331
    if total_331 ==0:
            vbox_completed_331Percent=(output_331/vbox_totalSubActivity)*100  
    else: 
            vbox_completed_331Percent=(output_331/total_331)*100
             
      # OUTPUT 3.3.2
    total_332= Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.3.2')).count()
    vbox_completed_332 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='Completed') & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.3.2')).count()
    vbox_OnTrack_332 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='On Track') & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.3.2')).count()
    output_332=vbox_completed_332+vbox_OnTrack_332
    if total_332 ==0:
            vbox_completed_332Percent=(output_332/vbox_totalSubActivity)*100  
    else: 
            vbox_completed_332Percent=(output_332/total_332)*100
    
        # OUTPUT 4.1.1
    total_411= Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='4.1.1')).count()
    vbox_completed_411 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='Completed') & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='4.1.1')).count()
    vbox_OnTrack_411 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='On Track') & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='4.1.1')).count()
    output_411=vbox_completed_411+vbox_OnTrack_411
    if total_411 ==0:
            vbox_completed_411Percent=(output_411/vbox_totalSubActivity)*100  
    else: 
            vbox_completed_411Percent=(output_411/total_411)*100
    
        # OUTPUT 4.1.2
    total_421= Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='4.2.1')).count()
    vbox_completed_421 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='Completed') & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='4.2.1')).count()
    vbox_OnTrack_421 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='On Track') & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='4.2.1')).count()
    output_421=vbox_completed_421+vbox_OnTrack_421
    if total_421 ==0:
            vbox_completed_421Percent=(output_421/vbox_totalSubActivity)*100  
    else: 
            vbox_completed_421Percent=(output_421/total_421)*100
    
    
          # OUTPUT 4.2.2
    total_422= Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='4.2.2')).count()
    vbox_completed_422 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='Completed') & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='4.2.2')).count()
    vbox_OnTrack_422 = Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__icontains='On Track') & Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='4.2.2')).count()
    output_422=vbox_completed_422+vbox_OnTrack_422
    if total_422 ==0:
            vbox_completed_422Percent=(output_422/vbox_totalSubActivity)*100  
    else: 
            vbox_completed_422Percent=(output_422/total_422)*100
   
    
    context = {'chart': chart, 
               'form': DateForm(),
               'totalSubActivity':totalSubActivity,
               'title_one_month_ago':title_one_month_ago,
               'statutworkplans': statutworkplans,
               'countries':countries,
               'list_not_start':list_not_start,
               'activity_completedStart': activity_completedStart,
               'activity_completedStartCount':activity_completedStartCount,
               'subactivityByUnit':subactivityByUnit,
               'chart_statut_month': chart_statut_month,
               'chart_statut_topTask':chart_statut_topTask,
               'chart_statut_unit':chart_statut_unit,
               'chart_statut_countries':chart_statut_countries,
               'vbox_totalSubActivity':vbox_totalSubActivity,
               'vbox_notStart':vbox_notStart,
               'vbox_onTrack':vbox_onTrack,
               'vbox_completed':vbox_completed,
               'vbox_completedPercent':vbox_completedPercent,
               'vbox_notStartIssue':vbox_notStartIssue,
               'vbox_onTrackPercent':vbox_onTrackPercent,
               'vbox_vbox_notStartPercent':vbox_vbox_notStartPercent,
               'vbox_notStartIssuePercent':vbox_notStartIssuePercent,
               # UNIT
               'vbox_completed_VID':vbox_completed_VID,
               'vbox_completed_VIDPercent':vbox_completed_VIDPercent,
               'vbox_completed_CHE':vbox_completed_CHE,
               'vbox_completed_CHEPercent': vbox_completed_CHEPercent,
               'vbox_completed_HPDPercent':vbox_completed_HPDPercent,
               'vbox_completed_NUTPercent':vbox_completed_NUTPercent,
               'vbox_completed_TNRPercent':vbox_completed_TNRPercent,
               'vbox_completed_UHUPercent':vbox_completed_UHUPercent,
               #output
               'total_111':total_111,
               'output_111':output_111,
               'vbox_completed_111Percent':vbox_completed_111Percent,
               'total_311':total_311,
               'output_311':output_311,
               'vbox_completed_311Percent':vbox_completed_311Percent,
               'total_312':total_312,
               'output_312':output_312,
               'vbox_completed_312Percent':vbox_completed_312Percent,
               'total_321':total_321,
               'output_321':output_321,
               'vbox_completed_321Percent':vbox_completed_321Percent,
               'total_322':total_322,
               'output_322':output_322,
               'vbox_completed_322Percent':vbox_completed_322Percent,
               'total_331':total_331,
               'output_331':output_331,
               'vbox_completed_331Percent':vbox_completed_331Percent,
               'output_332':output_332,
               'output_332':output_332,
               'vbox_completed_332Percent':vbox_completed_332Percent,
               'total_411':total_411,
               'output_411':output_411,
               'vbox_completed_411Percent':vbox_completed_411Percent,
               'total_421':total_421,
               'output_421':output_421,
               'vbox_completed_421Percent':vbox_completed_421Percent,
               'total_422':total_422,
               'output_422':output_422,
               'vbox_completed_422Percent':vbox_completed_422Percent,
               #TEST
               'tables_subactivities':tables_subactivities,
           #    'dt_fig_trend':dt_fig_trend,
               'table_outputKpiTask':table_outputKpiTask
               
               }
    return render(request, 'dashboard.html', context)

# CREATE VIEWS FOR GENERAL INFO SETTING -------------------------- 
@login_required(login_url='general_info')
@lockdown()
def general_info(request, *args, **kwargs):
    data = Outputworkplan.objects.all()
    kpi = Kpi.objects.all()
    gsmWorkplan=GsmWorkplan.objects.all()
    form_gsm = GsmwpForm()
    
    toptasks=Toptask.objects.all()
    form_toptask = ToptaskForm()
    
    
    form = KpiForm()
    form_output = OutputForm()
    countries=Country.objects.all()
    countryForm = CountryForm()
    statutworkplans=Statutworkplan.objects.all()
    units=Units.objects.all()
    page = request.GET.get('page')
    num_of_items = 10
    # paginator of country 
    paginator = Paginator(countries, num_of_items)
       
    try:
        countries = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        countries = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        countries = paginator.page(page)
        
        
    #print(data)
    context = {'data': data, 'kpis': kpi,'gsmWorkplans': gsmWorkplan,'form_gsm':form_gsm, 'toptasks':toptasks,'form_toptask':form_toptask,'form' : form,'form_output' : form_output, 'countries' : countries, 'countryForm' : countryForm, 'paginator': paginator, 'statutworkplans':statutworkplans, 'units':units}
    return render(request, 'pages/forms/general_info.html',context)

# CREATE VIEWS FOR COUNTRY-------------------------
@permission_required('myuhp.view_country',raise_exception=True)
def countryList(request):
    if request.method == 'POST':
        form = CountryForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Country has been created')
            return redirect('general_info')
           # return HttpResponse('<h1>Hello HttpResponse</h1>')   
    else:
        form = CountryForm()
        context = {
            'form': form
        }
    return render(request, 'pages/forms/general_info.html', context)

# Method to update Country name
@permission_required('myuhp.change_country',raise_exception=True)
def edit_country(request, pk):
    country=Country.objects.get(pk=pk)
    #output = get_object_or_404(Outputworkplan, id = pk)
    if request.method == 'POST':
        form = CountryForm(request.POST, instance = country)
        if form.is_valid():
            form.save()
            messages.success(request, 'Country name updated successfully')
            return redirect('general_info')
    else:
        form = CountryForm(instance=country)
        return render(request, 'pages/forms/country/country_edit.html', {
            'form': form,
            'country': country
        } )

# Methof to delete country name
@permission_required('myuhp.delete_country',raise_exception=True)
def delete_country(request,pk):
    country=Country.objects.get(id=pk)
    country.delete()
    messages.info(request, 'country deleted successfully')
    return redirect('general_info')

# CREATE VIEWS FOR STATUT------------------------------
@permission_required('myuhp.view_statutworkplan',raise_exception=True)
def statutList(request):
    if request.method == 'POST':
        form = StatutForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Statut has been created')
            return redirect('general_info')
           # return HttpResponse('<h1>Hello HttpResponse</h1>')   
    else:
        form = StatutForm()
        context = {
            'form': form
        }
    return render(request, 'pages/forms/general_info.html', context)
  
  
  # Method to update statut rubrique

@permission_required('myuhp.change_statutworkplan',raise_exception=True)
def edit_statut(request, pk):
    statutworkplan=Statutworkplan.objects.get(pk=pk)
    #output = get_object_or_404(Outputworkplan, id = pk)
    if request.method == 'POST':
        form = StatutForm(request.POST, instance = statutworkplan)
        if form.is_valid():
            form.save()
            messages.success(request, 'Statut name updated successfully')
            return redirect('general_info')
    else:
        form = StatutForm(instance=statutworkplan)
        return render(request, 'pages/forms/statutWp/statut_edit.html', {
            'form': form,
            'statutworkplan': statutworkplan
        } )

# Methof to delete statut name
@permission_required('myuhp.delete_statutworkplan',raise_exception=True)
def delete_statut(request,pk):
    statutworkplan=Statutworkplan.objects.get(id=pk)
    statutworkplan.delete()
    messages.info(request, 'Statut workplan deleted successfully')
    return redirect('general_info')

# CREATE VIEWS FOR UNITS-------------------------
@permission_required('myuhp.add_units',raise_exception=True)
def add_units(request):
    if request.method == 'POST':
        form_unit = UnitsForm(request.POST)
        if form_unit.is_valid():
            form_unit.save()
            messages.success(request, 'Unit has been created')
            return redirect('general_info')
           # return HttpResponse('<h1>Hello HttpResponse</h1>')   
    else:
        form_unit = UnitsForm()
        context = {
            'form_unit': form_unit
        }
    return render(request, 'pages/forms/general_info.html', context)

# Edit unit
@permission_required('myuhp.change_units',raise_exception=True)
def edit_unit(request, pk):
    unit=Units.objects.get(pk=pk)
    #output = get_object_or_404(Outputworkplan, id = pk)
    if request.method == 'POST':
        form = UnitsForm(request.POST, instance = unit)
        if form.is_valid():
            form.save()
            messages.success(request, 'Unit name updated successfully')
            return redirect('general_info')
    else:
        form = UnitsForm(instance=unit)
        return render(request, 'pages/forms/unitCluster/unit_edit.html', {
            'form': form,
            'unit': unit
        } )

# Methof to delete Unit name
@permission_required('myuhp.delete_units',raise_exception=True)
def delete_unit(request,pk):
    unit=Units.objects.get(id=pk)
    unit.delete()
    messages.info(request, 'Unit from workplan deleted successfully')
    return redirect('general_info')

# CREATE VIEWS FOR OUTPUT-----------------------------------------
# Methof to add new output
@permission_required('myuhp.view_outputworkplan',raise_exception=True)
def add_output(request):
    if request.method == 'POST':
        form = OutputForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Output has been created')
            return redirect('general_info')
           # return HttpResponse('<h1>Hello HttpResponse</h1>')   
    else:
        form = OutputForm()
        context = {
            'form': form
        }
    return render(request, 'pages/forms/outputs/output_add.html', context)

# Method to update output
@permission_required('myuhp.change_outputworkplan',raise_exception=True)
def edit_output(request, pk):
    output = Outputworkplan.objects.get(pk=pk)
    #output = get_object_or_404(Outputworkplan, id = pk)
    if request.method == 'POST':
        form = OutputForm(request.POST, instance = output)
        if form.is_valid():
            form.save()
            messages.success(request, 'Output updated successfully')
            return redirect('general_info')
    else:
        form = OutputForm(instance=output)
        return render(request, 'pages/forms/outputs/output_edit.html', {
            'form': form,
            'output': output
        } )
# Methof to delete output
@permission_required('myuhp.delete_outputworkplan',raise_exception=True)
def delete_output(request,pk):
    output=Outputworkplan.objects.get(id=pk)
    output.delete()
    messages.info(request, 'Output deleted successfully')
    return redirect('general_info')

# CREATE VIEWS FOR KPI------------------------------------
# Method to add new kpi
@permission_required('myuhp.add_kpi',raise_exception=True)
def add_kpi(request):
    if request.method == 'POST':
        form = KpiForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'KPI has been created')
            return redirect('general_info')
           # return HttpResponse('<h1>Hello HttpResponse</h1>')   
    else:
        form = KpiForm()
        context = {
            'form': form
        }
    return render(request, 'pages/forms/kpi/kpi_add.html', context)

# Method to update kpi
@permission_required('myuhp.change_kpi',raise_exception=True)
def edit_kpi(request, pk):
    kpi=Kpi.objects.get(pk=pk)
    #output = get_object_or_404(Outputworkplan, id = pk)
    if request.method == 'POST':
        form = KpiForm(request.POST, instance = kpi)
        if form.is_valid():
            form.save()
            messages.success(request, 'KPI updated successfully')
            return redirect('general_info')
    else:
        form = KpiForm(instance=kpi)
        return render(request, 'pages/forms/kpi/kpi_edit.html', {
            'form': form,
            'output': kpi
        } )

# Method to delete kpi
@permission_required('myuhp.delete_kpi',raise_exception=True)
def delete_kpi(request,pk):
    Kpi=Kpi.objects.get(id=pk)
    Kpi.delete()
    messages.warning(request, 'KPI deleted successfully')
    return redirect('general_info')

# Method to load kpi
@lockdown()
def load_kpi(request):
    options_kpi = Kpi.objects.all()
    options_kpi1=options_kpi[0]
    data={'options_kpi': options_kpi,'options_kpi1':options_kpi1}
    return render(request,'pages/forms/droplists/kpi_options.html', data )

#@permission_required('Myuhp.delete_KpiAchieve', raise_exception=False)
#@permission_required('myuhp.view_kpi achieve',raise_exception=True)
@login_required
@group_required('Technical officer')
def kpi_achieve_index(request):
    context={}
    form_unit = SelectUnitForm(request.GET or None)
    by_unit =  request.GET.get('by_unit')
    end_date =  request.GET.get('end')
    form = KpiAchieveForm()
    kpiAchieves = KpiAchieve.objects.all().order_by("kpi")
    
    if request.method=='GET':
      #  by_unit =  request.GET.get('by_unit')
        
        st=request.GET.get('kpi_code')
        if st!=None:
            kpiAchieves= KpiAchieve.objects.select_related('kpi').filter(kpi__kpi_code__icontains=st).order_by("kpi")
        
        if by_unit!=None:
            kpiAchieves = KpiAchieve.objects.select_related('kpi').filter(Q(kpi__unit__unit_code__icontains=by_unit)).order_by("kpi")
    
   ## page = request.GET.get('page')
   ## num_of_items = 10
   # # paginator of country 
   # paginator = Paginator(kpiAchieves, num_of_items)
       
   # try:
   #     kpiAchieves = paginator.page(page)
   # except PageNotAnInteger:
   #     page = 1
   #     kpiAchieves = paginator.page(page)
   # except EmptyPage:
   #     page = paginator.num_pages
   #     kpiAchieves = paginator.page(page)
        
   # context['paginator']=paginator 
    context['kpiAchieves'] = kpiAchieves
    context['title'] ='home'
    if request.method == 'POST':
        if 'save' in request.POST:
            pk = request.POST.get('save')
            if not pk:
                form=KpiAchieveForm(request.POST)
            else:
                kpiAchieve = KpiAchieve.objects.get(id = pk)
                form = KpiAchieveForm(request.POST ,instance = kpiAchieve) 
            form.save()
            form = KpiAchieveForm()
            messages.success(request, 'KPI result saved or updated successfully')
        elif 'delete' in request.POST:
            pk = request.POST.get('delete')
            kpiAchieve = KpiAchieve.objects.get(id = pk)
            kpiAchieve.delete()
            messages.success(request, 'KPI result deleted successfully')
           
            
        elif 'edit' in request.POST:
            pk = request.POST.get('edit')
            kpiAchieve = KpiAchieve.objects.get(id = pk)
            form = KpiAchieveForm(instance = kpiAchieve)
          
    context['form'] = form
    context['form_unit'] = form_unit
    context['by_unit'] = by_unit
    context['end_date'] = end_date
   
    return render(request, 'pages/forms/kpi/kpi_achieve_index.html', context)

# Method for list all KPI submitted
#@permission_required('myuhp.change_kpi achieve',raise_exception=True)
@login_required
@group_required('Technical officer')
def kpi_report(request):
    kpis = Kpi.objects.all().order_by("kpi_code") 
    if request.method=='GET':
        st=request.GET.get('kpi_code')
        if st!=None:
            kpis= Kpi.objects.filter(kpi_code__icontains=st)
                        
   # kpiAchieves = KpiAchieve.objects.all().order_by("kpi") 
   # if request.method=='GET':
   #     st=request.GET.get('kpi_code')
   #     if st!=None:
   #         kpiAchieves= KpiAchieve.objects.filter(kpi__kpi_code__icontains=st)
    
    page = request.GET.get('page')
    num_of_items = 10
    # paginator of country 
    paginator = Paginator(kpis, num_of_items)
       
    try:
        kpis = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        kpis = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        kpis = paginator.page(page)
            
        
    data={'kpis': kpis, 'paginator':paginator}
    
    return render(request, 'pages/forms/kpi/kpi_report.html', data)

# Method for view single KPI page
@login_required
@group_required('Technical officer')
def single_kpi_page(request, pk):
    
    try:
        kpis = Kpi.objects.get(id=pk) 
        kpiAchieves = KpiAchieve.objects.select_related('kpi').filter(Q(kpi__id__icontains=kpis.id) & Q(kpi__kpi_code__icontains=kpis.kpi_code) & Q(report_resut__isnull=False)).values(
            
            'kpi__expected_result',
            'kpi__kpi_code',
            'kpi__kpi_description',
            'kpi__kpi_country_level',
            'kpi__data_description',
            'kpi__data_source',
            'kpi__collection_methods',
            'kpi__data_frequency',
            'kpi__kpi_link',
            'kpi_baseline',
            'kpi_target',
            'report_date',
            'report_resut',
            'report_comment').order_by('kpi')
         
   #     kpiAchieves = Kpi.objects.filter(Q(id__icontains=pk) & Q(kpis_kpiAchive__report_resut__isnull=False)).values(
   #         'expected_result',
   #         'kpi_code',
   #         'kpi_description',
   #         'kpi_country_level',
    #        'data_description',
    #        'data_source',
    #        'collection_methods',
   #         'data_frequency',
  #          'kpi_link',
  #          'kpis_kpiAchive__kpi_baseline',
  #          'kpis_kpiAchive__kpi_target',
  #          'kpis_kpiAchive__report_date',
  #          'kpis_kpiAchive__report_resut',
  #          'kpis_kpiAchive__report_comment').order_by('kpi_code')
        data={'kpis':kpis, 'kpiAchieves':kpiAchieves}   
    except:
        data={'message':'The KPI you requested doesn t\' exist or isn t in the submitted results'}
       
    return render(request,'pages/forms/kpi/single_kpi_page.html',data)


# CREATE VIEWS FOR TOP TASK fROM GSM WORKPLAN-------------------------------------

# Method for add new Top task
@permission_required('myuhp.add_gsmworkplan',raise_exception=True)
def add_toptask(request):
    options_kpi = Kpi.objects.all()
    
    if request.method == 'POST':
        form = ToptaskForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Top Task has been created')
            return redirect('general_info')
           # return HttpResponse('<h1>Hello HttpResponse</h1>')   
    else:
        form = ToptaskForm()
        context = {
            'form': form,
            'options_kpi': options_kpi
        }
    return render(request, 'pages/forms/toptask/toptask_add.html', context)

# Method for update Top task 
@permission_required('myuhp.change_gsmworkplan',raise_exception=True)
def edit_toptask(request, pk):
    topTasks=Toptask.objects.get(pk=pk)
   
    #output = get_object_or_404(Outputworkplan, id = pk)
    if request.method == 'POST':
        form = ToptaskForm(request.POST, instance = topTasks)
        if form.is_valid():
            form.save()
            messages.success(request, 'Top Task updated successfully')
            return redirect('general_info')
    else:
        form = ToptaskForm(instance=topTasks)
        return render(request, 'pages/forms/toptask/toptask_edit.html', {
            'form': form,
            'output': topTasks
        } )
        
# Method for delete lowest task
@permission_required('myuhp.delete_gsmworkplan',raise_exception=True)
def delete_toptask(request,pk):
    topTasks=Toptask.objects.get(id=pk)
    topTasks.delete()
    messages.warning(request, 'Top Task deleted successfully')
    return redirect('general_info')


# CREATE VIEWS FOR LOWEST TASK fROM GSM WORKPLAN-------------------------------------

# Method for add new lowest task
@permission_required('myuhp.add_gsmworkplan',raise_exception=True)
def add_lowest(request):
    options_kpi = Kpi.objects.all()
    if request.method == 'POST':
        form = GsmwpForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Task has been created')
            return redirect('general_info')
           # return HttpResponse('<h1>Hello HttpResponse</h1>')   
    else:
        form = GsmwpForm()
        context = {
            'form': form,
            'options_kpi': options_kpi
        }
    return render(request, 'pages/forms/gsm/lowest_add.html', context)

# Method for update lowest task 
@permission_required('myuhp.change_gsmworkplan',raise_exception=True)
def edit_lowest(request, pk):
    gsmWp=GsmWorkplan.objects.get(pk=pk)
    #output = get_object_or_404(Outputworkplan, id = pk)
    if request.method == 'POST':
        form = GsmwpForm(request.POST, instance = gsmWp)
        if form.is_valid():
            form.save()
            messages.success(request, 'Task updated successfully')
            return redirect('general_info')
    else:
        form = GsmwpForm(instance=gsmWp)
    return render(request, 'pages/forms/gsm/lowest_edit.html', {
        'form': form,
        'output': gsmWp
    } )
    
# Method for delete lowest task
@permission_required('myuhp.delete_gsm workplan',raise_exception=True)
def delete_lowest(request,pk):
    gsmWp=GsmWorkplan.objects.get(id=pk)
    gsmWp.delete()
    messages.warning(request, 'Task deleted successfully')
    return redirect('general_info')

# CREATE VIEWS FOR SUB ACTIVITY fROM OPERATIONNAL WORKPLAN------------------------------------
# Method for load sub activity
@permission_required('myuhp.view_operworkplan',raise_exception=False)
def sub_activity_view(request):
    form_unit = SelectUnitForm(request.GET or None)
    by_unit =  request.GET.get('by_unit')
    end_date =  request.GET.get('end')
    # paginator of workplan 
    workplans = Operworkplan.objects.all().order_by("gsmWorkplan")
    form = WorkplanForm()
    if request.method=='GET':
        st=request.GET.get('sub_activity')
        if st!=None:
            workplans= Operworkplan.objects.filter(sub_activity__icontains=st)
        if by_unit!=None:
            workplans = Operworkplan.objects.all().filter(Q(gsmWorkplan__toptask__unit__unit_code__icontains=by_unit)& Q(completion_date__lte=end_date)).order_by("gsmWorkplan")
    
   # page = request.GET.get('page')
   # num_of_items = 10
  
 #   paginator = Paginator(workplans, num_of_items)
       
  #  try:
  #      workplans = paginator.page(page)
  #  except PageNotAnInteger:
  #      page = 1
  #      workplans = paginator.page(page)
  #  except EmptyPage:
  #      page = paginator.num_pages
  #      workplans = paginator.page(page)
            
    context={'form_unit': form_unit,
             'form':form,
             'workplans': workplans,
             'by_unit':by_unit,
             'end_date':end_date
            # 'paginator':paginator
             }
    
        
    return render(request, 'pages/forms/workplans/workplans.html', context)


class workplansView(ListView):
    context_object_name = "workplans"
    model = Operworkplan
    template_name = "pages/forms/workplans/workplans.html"
    form = WorkplanForm()
    paginate_by = 10
    paginate_orphans = 5
    paginator = Paginator(Operworkplan, paginate_by)
    queryset = Operworkplan.objects.all().order_by("gsmWorkplan")  # Default: Model.objects.all()

    def get_context_data(self, **kwargs):
        context = super(workplansView, self).get_context_data(**kwargs)
        context["workplans"] = Operworkplan.objects.all().order_by("gsmWorkplan")
        context['current_page'] = context.pop('page_obj', None)
        context["form"] = self.form
        
            
        return context

# Method for add new sub activity
#@permission_required('myuhp.add_operworkplan', raise_exception=False)
@login_required
@group_required('Technical officer')
def add_new_workplan_view(request):
    if request.method == "POST":
        form = WorkplanForm(request.POST)
        if form.is_valid():
            form.save(commit=True)
          #  form.save_m2m
         #   return JsonResponse({'msg': 'Data saved'})
            messages.success(request, "Sub activity Added Successfully")
            return redirect('workplans')
        else:
            print("ERROR FORM INVALID")
            return JsonResponse({'msg': 'ERROR FORM INVALID'})
    else:
   #     form = WorkplanForm()
  #  return JsonResponse({'form': form})
        context = {
            'form': form
        }
    return render(request, 'pages/forms/workplans/add_workplan_modal.html', context)

# Method-1 for edit sub activity
#@permission_required('myuhp.change_operworkplan', raise_exception=True)
@login_required
@group_required('Technical officer')
def edit_workplan_view(request, pk):
    countries=Country.objects.get(pk=pk)
    operworkplan = Operworkplan.objects.get(pk=pk)
  #  instance = get_object_or_404(models.Operworkplan, pk=pk)
    form = WorkplanForm(instance=operworkplan)
    if request.method == "POST":
        form = WorkplanForm(request.POST, instance=operworkplan)
        if form.is_valid():
            form.save(commit=True)
            return JsonResponse({'msg': 'Data Updated'})
        else:
            print("ERROR FORM INVALID")
            return JsonResponse({'msg': 'ERROR FORM INVALID'})
    
    else:
        form = WorkplanForm(instance=operworkplan)
        # print("ERROR FORM INVALID")
    return render(request, 'pages/forms/workplans/edit_workplan_modal.html', {'form': form, 'countries':countries})
          
# Method-2 for edit sub activity
@login_required
@group_required('Technical officer')
def edit_op_wp(request, pk): 
    operworkplan = Operworkplan.objects.get(id=pk)
    form = WorkplanForm(instance=operworkplan)
    
    if request.method == "POST":
        form = WorkplanForm(request.POST, instance=operworkplan)
        if form.is_valid():
            form.save()
            messages.success(request, "Sub activity updated Successfully")
            return redirect('workplans')
           
   
   # form = WorkplanForm(request.POST or None, instance=operworkplan)
    messages.warning(request, "ERROR FORM INVALID")
  
    return render(request, 'pages/forms/workplans/workplans.html', {'form': form})

# Method for delete sub activity
@permission_required('myuhp.delete_operworkplan', raise_exception=True)
def delete_workplan_view(request,pk):
    data=Operworkplan.objects.filter(id=pk)
    data.delete()
    messages.warning(request, 'Sub activity deleted successfully')
   
    return redirect('workplans')

# Fiche synthese workplan
@login_required
@group_required('Technical officer')
def fiche_workplan(request):
    selected = "operworkplans"
    workplan_list = Operworkplan.objects.all()
    if request.method =="POST":
        form = OperworkplanSearchForm(request.POST)
        if form.is_valid():
            base_url = reverse('operworkplans')
            query_string = urlencode(form.cleaned_data)
            url = '{}?{}'.format(base_url, query_string)
            return redirect(url)
    else:
        form = OperworkplanSearchForm()
        sreach_name = request.GET.get('gsmWorkplan','')
        if sreach_name is not None:
            workplan_list = workplan_list.filter(gsmWorkplan = sreach_name)
            form.fields['gsmWorkplan'].initial = sreach_name 
              
    # Pagination:10 elements par page
    paginator = Paginator(workplan_list.order_by('gsmWorkplan'),10)
    try:
        page = request.GET.get("page")
        if not page:
            page = 1
            workplan_list = paginator.page(page)
    except EmptyPage:
        # si on depasse la limite de pages, on prend la dernière
        workplan_list = paginator.page(paginator.num_pages())      
   # units=Units.objects.all().order_by("unit_code")
   # data={'units':units}
    return render(request,'pages/forms/workplans/workplan_fiche.html',{'workplan_list':workplan_list})

# Method for list all UNIT that submitted sub activity
#@user_passes_test(is_visitors)
#@permission_required('Myuhp.view_operworkplan', raise_exception=True)
#@permission_required('myuhp.view_operworkplan')
#@group_required('Administrative Assistance')
@login_required
@group_required('Technical officer')
def sub_activity_report(request, *args, **kwargs):
    
  # units = Units.objects.filter(unit_code__istartswith='VID').prefetch_related('units_kpi','units_gsmWorkplan')
    units = Units.objects.prefetch_related('units_kpi','units_toptask').distinct()
   # units = Units.objects.prefetch_related('units_kpi','units_gsmWorkplan')\
   #     .filter(units_kpi__kpi_code='AFR KPI 1.1.1.a')
    
   # operworkplans = Operworkplan.objects.select_related('gsmWorkplan','country','statut_name')
    operworkplans = Operworkplan.objects.only( 'completion_date','gsmWorkplan__toptask__top_task_description').select_related('units_toptask').distinct()
    gsmWorkplans = GsmWorkplan.objects.all().order_by("lowest_task_short").distinct() 

    page = request.GET.get('page')
    num_of_items = 10
    # paginator of country 
    paginator = Paginator(operworkplans, num_of_items)
       
    try:
        operworkplans = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        operworkplans = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        operworkplans = paginator.page(page)
           
    #unit_id = request.GET.get('unit_id')
 #   gsmWorkplans = GsmWorkplan.objects.filter(unit_id=unit_id).all()
            
    data={'operworkplans': operworkplans, 'paginator':paginator,'gsmWorkplans':gsmWorkplans,  'units': units}
   # if request.user.groups.filter(name='Technical officer').exists():
   # if request.user.has_perm('Myuhp.view_operworkplan'):
    return render(request, 'pages/forms/workplans/sub_activity_report.html', data)
   # else:
       # return render(request, 'pages/examples/403.html')
   #     return HttpResponse('<h4> You don\'t have access to the operational plan </h4>')
        

# Method for view single Sub activityI page
@login_required
@group_required('Technical officer')
def single_sub_activity_page(request,unit_code):
    try:

        data_unit =list(Units.objects.prefetch_related('units_toptask').filter(Q(unit_code__icontains=unit_code)& Q(units_toptask__toptask_gsmWorkplan__lowest_task_description__isnull=False)).distinct())
        data_outputs = Outputworkplan.objects.prefetch_related('outputs_toptask').filter(Q(outputs_toptask__unit__unit_code__isnull=False) & Q(outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False) & Q(outputs_toptask__unit__unit_code__icontains=unit_code)).distinct().order_by("outputs_toptask__output__output_code") 
      #  data_outputs = Operworkplan.objects.exclude(sub_activity__isnull=True).filter(Q(gsmWorkplan__toptask__unit__unit_code__icontains=unit_code)).distinct().order_by("gsmWorkplan__toptask__output__output_code") 
        kpis = Kpi.objects.filter(Q(unit__unit_code__icontains=unit_code)).prefetch_related('kpis_kpiAchive','kpis_toptask').distinct()
        context = {
                'data_unit':data_unit,
                'data_outputs':data_outputs,
                'kpis': kpis
                }
    except:
        context={'message':'The sub-activity you requested doesn t\' exist or isn t in the submitted results'}
       
    return render(request,'pages/forms/workplans/single_sub_activity_page.html',context)

# CREATE VIEW FOR INDIVIDUAL ANNUAL REPORT-----------------------------

# Method for list all Individual annual report submitted
@permission_required('myuhp.view_individualreport', raise_exception=True)
def individual_report(request):
    individualReports = IndividualReport.objects.all().order_by("staff_name") 
    page = request.GET.get('page')
    num_of_items = 10
    # paginator of country 
    paginator = Paginator(individualReports, num_of_items)
       
    try:
        individualReports = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        individualReports = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        individualReports = paginator.page(page)
            
    data={'individualReports': individualReports, 'paginator':paginator}
    
    return render(request, 'pages/forms/individual_report/index_individual_report.html', data)

# Method for add Individual annual report
@permission_required('myuhp.add_individualreport', raise_exception=True)
def add_indivReport(request):
     if request.method == 'POST':
        form = IndivReportForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Individual repport has been created')
            return redirect('indiv_report')
           # return HttpResponse('<h1>Hello HttpResponse</h1>')   
     else:
        form = IndivReportForm()
        context = {
            'form': form
        }
     return render(request, 'pages/forms/individual_report/add_individual_report.html',context)

# Method for view single personnal Individual annual report
def single_report_page(request, name):
    try:
        individualReport=IndividualReport.objects.get(staff_name=name) 
        data={'individualReport':individualReport}   
    except:
        data={'message':'l\'article que vous avez demandé n\'existe pas'}
       
    return render(request,'pages/forms/individual_report/single_report_page.html',data)

# Method for edit Individual annual report
@permission_required('myuhp.change_individualreport', raise_exception=True)
def edit_indivReport(request, pk): 
    individualReport = IndividualReport.objects.get(id=pk)
    form = IndivReportForm(instance=individualReport)   
    if request.method == "POST":
        form = IndivReportForm(request.POST, instance=individualReport)
        if form.is_valid():
            form.save()
            messages.success(request, "Individual annual report updated Successfully")
         #   return render(request, 'pages/forms/individual_report/index_individual_report.html')
            return redirect('indiv_report')
        else:
            print("ERROR FORM INVALID")
            return JsonResponse({'msg': 'ERROR FORM INVALID'})
    else:       
     print("ERROR FORM INVALID")
    return render(request, 'pages/forms/individual_report/edit_individual_report.html', {'form': form, 'individualReport':individualReport})

def events_cluster(request):
    return render(request, 'pages/eventsWorks/events_calendar.html')

# Methon for SURVEY PROJECT
              
# Method to add new data into the SuverDataset
#@lockdown
#@permission_required('myuhp.add_surveyproject', raise_exception=True)
@login_required
@group_required('Technical officer')
def survey_add_project(request):
    context={}
    form = SurveyProjectForm()
    surveyProjects = SurveyProject.objects.all()
    
    if request.method=='GET':
        st=request.GET.get('title_surv')
        if st!=None:
            surveyProjects= SurveyProject.objects.filter(title_surv__icontains=st)
    
    page = request.GET.get('page')
    num_of_items = 10
    # paginator of country 
    paginator = Paginator(surveyProjects, num_of_items)
       
    try:
        surveyProjects = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        surveyProjects = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        surveyProjects = paginator.page(page)
        
        
    context['surveyProjects'] = surveyProjects
    context['paginator']=paginator
    context['title'] ='home'
    if request.method == 'POST':
        if 'save' in request.POST:
            pk = request.POST.get('save')
            if not pk:
                form=SurveyProjectForm(request.POST)
            else:
                surveyProject = SurveyProject.objects.get(id = pk)
                form = SurveyProjectForm(request.POST ,instance = surveyProject) 
            form.save()
            form = SurveyProjectForm()
            messages.success(request, 'Survey project saved or updated successfully')
        elif 'delete' in request.POST:
            pk = request.POST.get('delete')
            surveyProject = SurveyProject.objects.get(id = pk)
            surveyProject.delete()
            messages.success(request, 'Survey project deleted successfully')
           
            
        elif 'edit' in request.POST:
            pk = request.POST.get('edit')
            surveyProject = SurveyProject.objects.get(id = pk)
            form = SurveyProjectForm(instance = surveyProject)
          
    context['form'] = form
   
    return render(request, 'pages/survey/survey_project.html', context)

# Method for import and export Table SurveyDataset
@lockdown()
@permission_required('myuhp.add_surveydataset', raise_exception=True)
def simple_upload(request):
    if request.method =='POST':
        survey_resource = SurveyResource()
        dataset = Dataset()
        new_survey = request.FILES['myfile']
        
        if not new_survey.name.endswith('xlsx'):
            messages.info(request, 'wrong format')
            return render(request, 'upload.html')
        
        import_data = dataset.load(new_survey.read(),format='xlsx')
        for data in import_data:
            value = SurveyDataset(
                data[0],
                data[1],
                data[2],
                data[3],
                data[4],
                data[5],
                data[6],
                data[7]
            )
            value.save()
    return render(request, 'pages/survey/upload.html')
              
# Method to add new data into the SuverDataset
#@login_required
#@group_required('Technical officer')
def survey_add_data(request):
  #  ct=ContentType.objects.get_for_model(SurveyDataset)
  #  if request.user.permissions.filter(codename="viw_surveydataset", contentype=ct).exists():
        
    context={}
    form_survey_title = SelectSurveyForm(request.GET or None)
    by_survey =  request.GET.get('by_survey')
    end_day =  request.GET.get('end')
 
    form = SurveyDatasetForm()
    surveyDatasets = SurveyDataset.objects.all().order_by("-surveyProject__end_date")
    if request.method=='GET':
        st=request.GET.get('quest_code')
        if st!=None:
            surveyDatasets= SurveyDataset.objects.filter(quest_code__icontains=st).order_by("-surveyProject__end_date")
            
        if by_survey!=None:
            surveyDatasets = SurveyDataset.objects.filter(Q(surveyProject__id__icontains=by_survey)& Q(surveyProject__end_date__lte=end_day)).order_by("quest_code")
            
    #page = request.GET.get('page')
    #num_of_items = 10
  
  #  paginator = Paginator(surveyDatasets, num_of_items)
       
  #  try:
  #      surveyDatasets = paginator.page(page)
  #  except PageNotAnInteger:
  #      page = 1
  #      surveyDatasets = paginator.page(page)
  #  except EmptyPage:
  #      page = paginator.num_pages
  #      surveyDatasets = paginator.page(page)
        

    context['surveyDatasets'] = surveyDatasets
   # context['paginator']=paginator
    context['title'] ='home'
    if request.user.groups.filter(name='Technical officer').exists():
    
        if request.method == 'POST':
            if 'save' in request.POST:
                pk = request.POST.get('save')
                if not pk:
                    form=SurveyDatasetForm(request.POST)
                else:
                    surveyDataset = SurveyDataset.objects.get(id = pk)
                    form = SurveyDatasetForm(request.POST ,instance = surveyDataset) 
                form.save()
                form = SurveyDatasetForm()
                messages.success(request, 'Survey dataset saved or updated successfully')
            elif 'delete' in request.POST:
                pk = request.POST.get('delete')
                surveyDataset = SurveyDataset.objects.get(id = pk)
                surveyDataset.delete()
                messages.success(request, 'Survey dataset deleted successfully')
            
                
            elif 'edit' in request.POST:
                pk = request.POST.get('edit')
                surveyDataset = SurveyDataset.objects.get(id = pk)
                form = SurveyDatasetForm(instance = surveyDataset)
            
        context['form'] = form
        
        context['form_survey_title'] = form_survey_title 
        context['by_survey'] = by_survey
        context['end_day'] = end_day
    else:
        context['form'] = form
        
        context['form_survey_title'] = form_survey_title 
        context['by_survey'] = by_survey
        context['end_day'] = end_day
        return render(request, 'pages/survey/survey_dataset.html', context)
   
 #   if request.user.has_perm('myuhp.view_surveydataset'):
    return render(request, 'pages/survey/survey_dataset.html', context)
   # else:
   #     return HttpResponse("Not allowed")

# Method for list all Survey dataset submitted
#@permission_required('myuhp.view_surveydataset',raise_exception=True)

def survey_report(request):
    surveyProjects = SurveyProject.objects.all().order_by("title_surv") 
    if request.method=='GET':
        st=request.GET.get('responsible')
        if st!=None:
            surveyProjects= SurveyProject.objects.filter(responsible__icontains=st)
    
    page = request.GET.get('page')
    num_of_items = 10
    # paginator of country 
    paginator = Paginator(surveyProjects, num_of_items)
       
    try:
        surveyProjects = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        surveyProjects = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        surveyProjects = paginator.page(page)
            
    data={'surveyProjects': surveyProjects, 'paginator':paginator}
    
    return render(request, 'pages/survey/survey_report.html', data)

# Method for view single Survey DataSet page

def single_survey_page(request, pk):
    try:
        surveyProject=SurveyProject.objects.get(id=pk) 
        surveyDatasets =SurveyProject.objects.filter(Q(id__icontains=pk) & Q(project_surveyData__quest_code__isnull=False)).values(
        'responsible',
        'title_surv',
        'start_date',
        'end_date',
        'location_survey',
        'project_surveyData__quest_code',
        'project_surveyData__question',
        'project_surveyData__response_text',
        'project_surveyData__response_num',
        'project_surveyData__level_1',
        'project_surveyData__level_2'
    ).order_by('-start_date')
        
        
        page = request.GET.get('page')
        num_of_items = 10
        # paginator of country 
        paginator = Paginator(surveyDatasets, num_of_items)
        
        try:
            surveyDatasets = paginator.page(page)
        except PageNotAnInteger:
            page = 1
            surveyDatasets = paginator.page(page)
        except EmptyPage:
            page = paginator.num_pages
            surveyDatasets = paginator.page(page)
        
        data={'surveyProject':surveyProject, 'surveyDatasets':surveyDatasets, 'paginator':paginator}   
    except:
        data={'message':'The Survey DataSet you requested doesn t\' exist or isn t in the submitted results'}
    
       
    return render(request,'pages/survey/single_survey_page.html',data)
    
# Method for upload document or image
@lockdown()
@permission_required('myuhp.add_docsave',raise_exception=True)
def docSave_upload(request):
    form =None
    if request.method == "POST":
        form = DocSaveForm(request.POST, request.FILES)
       # print(form)
        if form.is_valid():
            form.save()
            # Get the current instance object to display in the template
            img_obj = form.instance
        return render(request, 'pages/docSave/doc_upload.html', {'form': form, 'img_obj': img_obj})
    else:
        form = DocSaveForm()
        img_obj_all = DocSave.objects.all()
     
    return render(request,'pages/docSave/doc_upload.html', {'form': form, 'img_obj_all':img_obj_all})

# CREATE VIEWS FOR SUBSCRIBERS -------------------------- 
@lockdown()
def subscribers_views(request, *args, **kwargs):
    subscribers = Subscribers.objects.all()
    page = request.GET.get('page')
    num_of_items = 10
    # paginator of subscribers 
    paginator = Paginator(subscribers, num_of_items)
       
    try:
        subscribers = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        subscribers = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        subscribers = paginator.page(page)

    context = {'subscribers': subscribers, 'paginator': paginator}
    return render(request, 'pages/forms/subscribers/subscribers_list.html',context)

# Method for add subscribers
def add_subscribers(request):
     if request.method == 'POST':
        form = SubscribersForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'New subscriber has been created')
            return redirect('subscribers_views') 
     else:
        form = SubscribersForm()
        context = {
            'form': form
        }
     return render(request, 'pages/forms/subscribers/subscribers_add.html',context)
 
# Method for edit subscribers
def edit_subscribers(request, pk): 
    subscriber = Subscribers.objects.get(id=pk)
    form = SubscribersForm(instance=subscriber)
    
    if request.method == "POST":
        form = SubscribersForm(request.POST, instance=subscriber)
        if form.is_valid():
            form.save()
            messages.success(request, "Subscribers updated Successfully")
            return redirect('subscribers_views')

    messages.warning(request, "ERROR FORM INVALID")
  
    return render(request, 'pages/forms/subscribers/subscribers_edit.html', {'form': form})

# Method for delete subscribers
def delete_subscribers(request,pk):
    data=Subscribers.objects.filter(id=pk)
    data.delete()
    messages.warning(request, 'Subscribers deleted successfully')
   
    return redirect('subscribers_views')


#Method to export ALL SURVEY PROJECT on Excel file
def export_project_survey(request):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="Project_survey.xlsx"'

    wb = Workbook()
    ws = wb.active
    ws.title = "Project Survey"

    # Add headers
    headers = ["id","responsible", "title_surv", "start_date", "end_date", "location_survey"]
    ws.append(headers)

    # Add data from the model
    surveyProjects = SurveyProject.objects.all()
    for surveyProject in surveyProjects:
        ws.append([surveyProject.pk, surveyProject.responsible, surveyProject.title_surv, surveyProject.start_date, surveyProject.end_date, surveyProject.location_survey])

    # Save the workbook to the HttpResponse
    wb.save(response)
    return response

#Method to export SURVEY DATASET ALL DATA on Excel file
def export_to_excel(request):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="survey.xlsx"'

    wb = Workbook()
    ws = wb.active
    ws.title = "Survey"

    # Add headers
    headers = ["responsible", "title_surv", "start_date", "end_date", "location_survey", "quest_code", "question", "response_text", "response_num","level_1","level_2"]
    ws.append(headers)

    # Add data from the model
    surveyDatasets = SurveyDataset.objects.all()
    for s in surveyDatasets:
        ws.append([s.surveyProject.responsible, s.surveyProject.title_surv, s.surveyProject.start_date, s.surveyProject.end_date, s.surveyProject.location_survey, 
                   s.quest_code,  s.question,  s.response_text, s.response_num, s.level_1, s.level_2])

    # Save the workbook to the HttpResponse
    wb.save(response)
    return response

#Method to export DATASET on Excel file ONE BY ONE PROJECT

def export_survey_page(request,pk):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="survey_dataset.xls"'

    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet('Survey')

    # Sheet header, first row
    row_num = 0

    font_style = xlwt.XFStyle()
    font_style.font.bold = True

    columns = ['responsible', 'title_surv', 'start_date', 'end_date', 'location_survey', 'quest_code', 'question', 'response_text', 'response_num','level_1','level_2' ]

    for col_num in range(len(columns)):
        ws.write(row_num, col_num, columns[col_num], font_style)

    # Sheet body, remaining rows
    font_style = xlwt.XFStyle()

    rows = SurveyProject.objects.filter(Q(id__icontains=pk) & Q(project_surveyData__quest_code__isnull=False)).all().values_list(
        'responsible',
        'title_surv',
        'start_date',
        'end_date',
        'location_survey',
        'project_surveyData__quest_code',
        'project_surveyData__question',
        'project_surveyData__response_text',
        'project_surveyData__response_num',
        'project_surveyData__level_1',
        'project_surveyData__level_2'
    ).order_by('-start_date')
    for row in rows:
        row_num += 1
        for col_num in range(len(row)):
            ws.write(row_num, col_num, row[col_num], font_style)

    wb.save(response)
    return response

# Method to export on Excel file
def export_kpi_excel(request,pk):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="kpi_results.xls"'

    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet('KPI')

    # Sheet header, first row
    row_num = 0

    font_style = xlwt.XFStyle()
    font_style.font.bold = True

    columns = ['kpi_code', 'kpi_description', 'kpi_baseline', 'kpi_target', 'report_date', 'report_resut', 'report_comment' ]

    for col_num in range(len(columns)):
        ws.write(row_num, col_num, columns[col_num], font_style)

    # Sheet body, remaining rows
    font_style = xlwt.XFStyle()

    rows = Kpi.objects.filter(Q(id__icontains=pk) & Q(kpis_kpiAchive__kpi_target__isnull=False)).values_list(
        'kpi_code',
        'kpi_description',
        'kpis_kpiAchive__kpi_baseline',
        'kpis_kpiAchive__kpi_target',
        'kpis_kpiAchive__report_date',
        'kpis_kpiAchive__report_resut',
        'kpis_kpiAchive__report_comment'
    ).order_by('-kpis_kpiAchive__report_date')
    for row in rows:
        row_num += 1
        for col_num in range(len(row)):
            ws.write(row_num, col_num, row[col_num], font_style)

    wb.save(response)
    return response


#Method to export on Excel file the list of KPI
def export_listKpi_excel(request):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="Kpi.xlsx"'

    wb = Workbook()
    ws = wb.active
    ws.title = "Kpi"

    # Add headers
    headers = ["Unit","Output", "kpi_code", "kpi_description", "kpi_country_level", "data_description", "data_source", "collection_methods", "data_frequency", "kpi_link"]
    ws.append(headers)

    # Add data from the model
    kpis = Kpi.objects.all()
    for kpi in kpis:
        ws.append([kpi.unit.unit_code, kpi.expected_result, kpi.kpi_code, kpi.kpi_description, kpi.kpi_country_level, 
                   kpi.data_description,  kpi.data_source,  kpi.collection_methods, kpi.data_frequency, kpi.kpi_link])

    # Save the workbook to the HttpResponse
    wb.save(response)
    return response




#Method to export on Excel file the list of Output
def export_listOutput_excel(request):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="Output.xlsx"'

    wb = Workbook()
    ws = wb.active
    ws.title = "Output"

    # Add headers
    headers = ["output_code","output_description"]
    ws.append(headers)

    # Add data from the model
    outputs = Outputworkplan.objects.all()
    for output in outputs:
        ws.append([output.output_code, output.output_description])

    # Save the workbook to the HttpResponse
    wb.save(response)
    return response




#Method to export on Excel file the list of KPI
def export_subactivity_excel(request):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="Workplan.xlsx"'

    wb = Workbook()
    ws = wb.active
    ws.title = "Workplan"

    # Add headers
    headers = ["Unit","Output", "kpi description", "top task num", "top task description", "lowest task num", "lowest task description", "sub activity", "responsable", "completion date", "statut name", "comments"]
    ws.append(headers)

    # Add data from the model
    operworkplans = Operworkplan.objects.all()
    for operworkplan in operworkplans:
        ws.append([operworkplan.gsmWorkplan.toptask.unit.unit_code, 
                   operworkplan.gsmWorkplan.toptask.output.output_description,
                   operworkplan.gsmWorkplan.toptask.kpi.kpi_description, 
                   operworkplan.gsmWorkplan.toptask.top_task,
                   operworkplan.gsmWorkplan.toptask.top_task_description,
                   operworkplan.gsmWorkplan.lowest_task, 
                   operworkplan.gsmWorkplan.lowest_task_description,
                   operworkplan.sub_activity,  
                   operworkplan.responsable, 
                   operworkplan.completion_date,
                   operworkplan.statut_name.statut_name, 
                   operworkplan.comments])

    # Save the workbook to the HttpResponse
    wb.save(response)
    return response

#METHOD TO GENERATE SUMMARY OF KPI ON PDF FORMAT
@login_required
def render_pdf_view(request):
    template_path= 'pages/report.html'
    #units_plan = Units.objects.prefetch_related('units_gsmWorkplan','units_kpi').order_by('unit_code')
    #user_profiles=[item.name_subscriber for item in Subscribers.objects.all()]
    # user_profiles=list(Subscribers.objects.values())
    #user_profiles=list(Subscribers.objects.values('name_subscriber','email_subscriber'))
    #lower_name_subsc = Subscribers.objects.annotate(lower_name=Lower('name_subscriber'))
    
    vid_Units=list(Units.objects.filter(unit_code__icontains='VID').values_list('unit_description', flat=True))
    vid_kpis = Kpi.objects.prefetch_related('kpis_kpiAchive').filter(Q(unit__unit_code__icontains='VID')).distinct()
    #Q(kpis_kpiAchive__report_resut__isnull=True) & 
    
    che_Units=list(Units.objects.filter(unit_code__icontains='CHE').values_list('unit_description', flat=True))
    che_kpis = Kpi.objects.prefetch_related('kpis_kpiAchive').filter(Q(unit__unit_code__icontains='CHE')).distinct()
    
    hpd_Units=list(Units.objects.filter(unit_code__icontains='HPD').values_list('unit_description', flat=True))
    hpd_kpis = Kpi.objects.prefetch_related('kpis_kpiAchive').filter(Q(unit__unit_code__icontains='HPD')).distinct()
    
    nut_Units=list(Units.objects.filter(unit_code__icontains='NUT').values_list('unit_description', flat=True))
    nut_kpis = Kpi.objects.prefetch_related('kpis_kpiAchive').filter(Q(unit__unit_code__icontains='NUT')).distinct()
    
    tnr_Units=list(Units.objects.filter(unit_code__icontains='TNR').values_list('unit_description', flat=True))
    tnr_kpis = Kpi.objects.prefetch_related('kpis_kpiAchive').filter(Q(unit__unit_code__icontains='TNR')).distinct()
    
    uhu_Units=list(Units.objects.filter(unit_code__icontains='UHU').values_list('unit_description', flat=True))
    uhu_kpis = Kpi.objects.prefetch_related('kpis_kpiAchive').filter(Q(unit__unit_code__icontains='UHU')).distinct()
    context = {
               'vid_Units':vid_Units,
               'vid_kpis':vid_kpis,
               'che_Units':che_Units,
               'che_kpis':che_kpis,
               'hpd_Units':hpd_Units,
               'hpd_kpis':hpd_kpis,
               'nut_Units':nut_Units,
               'nut_kpis':nut_kpis,
               'tnr_Units':tnr_Units,
               'tnr_kpis':tnr_kpis,
               'uhu_Units':uhu_Units,
               'uhu_kpis':uhu_kpis
               }
    #create a Django response object , and specify content_type as pdf
    response = HttpResponse(content_type='application/pdf')
    # if download:
    # response['Content-Disposition']='attachment; filename="report.pdf"'
    # if display:
    response['Content-Disposition']='filename="Kpi_achieve_AfroUHP.pdf"'
    #find the temple and render it
    template = get_template(template_path)
    html = template.render(context)
    #create pdf
    pisa_status = pisa.CreatePDF(
        html, dest=response
    )
    # if error then show  some funy view
    if pisa_status.err:
        return HttpResponse('We had some errors <pre>'+html+'</pre>')
    return response

#Method to generate the pdf for activities file from html
@login_required
def render_pdf_activities(request):
    template_path= 'pages/pdfTemplate/pdf_activity.html'
   
    data_unit =list(Units.objects.prefetch_related('units_toptask').filter(Q(units_toptask__toptask_gsmWorkplan__lowest_task_description__isnull=False)).distinct())
    data_outputs = Outputworkplan.objects.prefetch_related('outputs_toptask').distinct().order_by("outputs_toptask__output__output_code") 
    
    context = {
               'data_unit':data_unit,
               'data_outputs':data_outputs,
               'code_unit':'UHP Cluster'
               }
    #create a Django response object , and specify content_type as pdf
    response = HttpResponse(content_type='application/pdf')
    # if download:
    # response['Content-Disposition']='attachment; filename="report.pdf"'
    # if display:
    response['Content-Disposition']='filename="UHP_subactivity.pdf"'
    #find the temple and render it
    template = get_template(template_path)
    html = template.render(context)
    #create pdf
    pisa_status = pisa.CreatePDF(
        html, dest=response
    )
    # if error then show  some funy view
    if pisa_status.err:
        return HttpResponse('We had some errors <pre>'+html+'</pre>')
    return response


#Method to generate the pdf for CHE activities file from html
@login_required
def CHE_pdf_activities(request):
    template_path= 'pages/pdfTemplate/che_activity.html'
   
    code_unit='CHE'
    if code_unit:
       data_unit =list(Units.objects.prefetch_related('units_toptask').filter(Q(unit_code__icontains=code_unit)& Q(units_toptask__toptask_gsmWorkplan__lowest_task_description__isnull=False)).distinct())
       data_outputs = Outputworkplan.objects.filter(Q(outputs_toptask__unit__unit_code__icontains=code_unit)).prefetch_related('outputs_toptask').distinct().order_by("outputs_toptask__output__output_code") 
     
    context = {
               'data_unit':data_unit,
               'data_outputs':data_outputs,
               'code_unit':code_unit
               }
    #create a Django response object , and specify content_type as pdf
    response = HttpResponse(content_type='application/pdf')
    # if download:
    # response['Content-Disposition']='attachment; filename="report.pdf"'
    # if display:
    response['Content-Disposition']='filename="CHE_subactivity.pdf"'
    #find the temple and render it
    template = get_template(template_path)
    html = template.render(context)
    #create pdf
    pisa_status = pisa.CreatePDF(
        html, dest=response
    )
    # if error then show  some funy view
    if pisa_status.err:
        return HttpResponse('We had some errors <pre>'+html+'</pre>')
    return response

# Method for upload pdf document or report in pdf format
@lockdown()
@permission_required('myuhp.add_reportsave',raise_exception=True)
def report_upload(request):
    form =None
    if request.method == "POST":
        form = ReportSaveForm(request.POST, request.FILES)
       # print(form)
        if form.is_valid():
            form.save()
            messages.success(request, 'pdf file added successfully')
            # Get the current instance object to display in the template
            img_report_obj = form.instance
        return render(request, 'pages/docSave/report_upload.html', {'form': form, 'img_report_obj': img_report_obj})
    else:
        form = ReportSaveForm()
        img_report_obj_all = ReportSave.objects.all()
     
    return render(request,'pages/docSave/report_upload.html', {'form': form, 'img_report_obj_all':img_report_obj_all})

# Method for views all reports upload in the system
def views_report(request):
    reportSaves = ReportSave.objects.all().order_by("title_rep") 
    if request.method=='GET':
        st=request.GET.get('title_rep')
        if st!=None:
            reportSaves= ReportSave.objects.filter(title_rep__icontains=st)
    
    page = request.GET.get('page')
    num_of_items = 10
    # paginator of country 
    paginator = Paginator(reportSaves, num_of_items)
       
    try:
        reportSaves = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        reportSaves = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        reportSaves = paginator.page(page)
            
    data={'reportSaves': reportSaves, 'paginator':paginator}
    
    return render(request, 'pages/docSave/report_list.html', data)

# Method to put  into the table all reports upload in the system
@login_required
@permission_required('myuhp.view_reportsave',raise_exception=True)
def index_report(request):
    reportSaves = ReportSave.objects.all().order_by("title_rep") 
    if request.method=='GET':
        st=request.GET.get('title_rep')
        if st!=None:
            reportSaves= ReportSave.objects.filter(title_rep__icontains=st)
    
    page = request.GET.get('page')
    num_of_items = 10
    # paginator of country 
    paginator = Paginator(reportSaves, num_of_items)
       
    try:
        reportSaves = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        reportSaves = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        reportSaves = paginator.page(page)
            
    data={'reportSaves': reportSaves, 'paginator':paginator}
    
    return render(request, 'pages/docSave/report_index.html', data)

# Edit report thad added in the system
@login_required
@permission_required('myuhp.add_reportsave',raise_exception=True)
def editReport(request, pk):
    reportSaves = ReportSave.objects.get(id=pk)
    form = ReportSaveForm(instance=reportSaves)
    
    if request.method == "POST":
        form = ReportSaveForm(request.POST, request.FILES, instance=reportSaves)
        if form.is_valid():
            form.save()
            messages.success(request, "Report updated Successfully")
            return redirect('index_report')

    messages.warning(request, "ERROR FORM INVALID")
    return render(request, 'pages/docSave/report_edit.html', {'form': form, 'reportSaves':reportSaves})

# DISPLAY DASHORD 01 FROM SHINY APP---------------------
def shiny_dashboard_ind(request):
    
 return render(request,'pages/dataBank/dashboard_indicator.html')

# DISPLAY PAGE OF ERROR 404
def handel404(request, exception):
    return render(request, 'pages/examples/404.html')

# DISPLAY PAGE OF ERROR 403
def handel403(request, exception):
    return render(request, 'pages/examples/403.html')

# DISPLAY PAGE OF ERROR 500
def custom_500(request):
    return render(request, 'pages/examples/500.html', status=500)

# DISPLAY MONITORING PERFORMANCE OF WORKPLAN BY UNIT AND BUY COMPLETION DATE 
def unit_sub_activity_view(request):
        units={}
        dataset_byUnit={}
        kpis={}
        line_chart_unit={}
        dt_fig_unit={}
        date_completion__max=timezone.now()
        notStartedByCountries = []
        df_subActByCountries = {}
        dt_pie_status={}
        dta_NumberTask = {} 
    
        form = SelectUnitForm(request.GET or None)
        
        try:
            one_month_ago = timezone.now() + timezone.timedelta(days=15) # under ToDay untill next 15 days
            one_month_next = timezone.now() + timezone.timedelta(days=31) # under ToDay untill next 31 days
            by_unit =  request.GET.get('by_unit')
           # start =  request.GET.get('start')
            end_date =  request.GET.get('end')
          
           # kpis = Kpi.objects.prefetch_related('kpis_kpiAchive','kpis_gsmWorkplan').distinct()
            units =list(Units.objects.filter(Q(unit_code__icontains=by_unit)).prefetch_related('units_kpi','units_toptask').values_list('unit_description', flat=True).distinct())
            
            #1. DATASET 01
            dt_fig_unit = Statutworkplan.objects.exclude(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__isnull=True).filter(Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains=by_unit)).annotate(
            total_planned = Count('statuts_operwork__sub_activity', filter=Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__isnull=False)),
            until_thisDay_planned = Count('statuts_operwork__sub_activity', filter=Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__isnull=False) & Q(statuts_operwork__completion_date__lte=one_month_ago)),
            select_day_planned = Count('statuts_operwork__sub_activity', filter=Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__isnull=False) & Q(statuts_operwork__completion_date__lte=end_date)),
           # total_activitiy=Case(
           #             statuts_operwork__completion_date__gt=start,
           #             statuts_operwork__completion_date__lt=end,
           #             then="total_planned",
           #             default=Value("0"),
           #             ), 
            ).values('statuts_operwork__gsmWorkplan__toptask__unit__unit_code','statut_name','statuts_operwork__completion_date','total_planned', 'until_thisDay_planned','select_day_planned').order_by('statuts_operwork__completion_date')

           #2. DATASET 02
            date_completion__max=max(list(dt_fig_unit.values_list('statuts_operwork__completion_date')))
           
           #3. DATASET 03  
            fig_unit = px.line(dt_fig_unit, x='statuts_operwork__completion_date', y='total_planned', text='total_planned', color='statut_name',labels={'total_planned':'Total of sub activity', 'statuts_operwork__completion_date':'Date', 'statut_name':'Statut name'})
            fig_unit.update_traces(textposition="bottom right")
            line_chart_unit = fig_unit.to_html(full_html=False, include_plotlyjs=False)
            
            #4. DATASET 04
            dt_pie_status= dt_fig_unit.values('statut_name','total_planned','until_thisDay_planned','select_day_planned').aggregate(        
                planed_notStart=Sum('total_planned', filter=Q(statut_name__isnull=False) & Q(statut_name__icontains='Not Started')),
                planed_issue=Sum('total_planned', filter=Q(statut_name__isnull=False) & Q(statut_name__icontains='Stalled')),
                planed_onTrack=Sum('total_planned', filter=Q(statut_name__isnull=False) & Q(statut_name__icontains='On Track')),
                planed_completed=Sum('total_planned', filter=Q(statut_name__isnull=False) & Q(statut_name__icontains='Completed')),
                one_month_ago_activity=Sum('until_thisDay_planned', filter=Q(statut_name__isnull=False)),
                this_period=Sum('select_day_planned', filter=Q(statut_name__isnull=False))
            )
            dta_NumberTask = Units.objects.exclude(units_toptask__toptask_gsmWorkplan__lowest_task_short__isnull=True).filter(unit_code__icontains=by_unit).annotate(
                total_output = Count('units_toptask__output__output_code',distinct=('units_toptask__output__output_code'), filter=Q(units_toptask__output__isnull=False)),
                total_kpi = Count('units_toptask__kpi',distinct=('units_toptask__kpi'), filter=Q(units_toptask__kpi__isnull=False)),
                total_TopTask = Count('units_toptask__top_task_short',distinct=('units_toptask__top_task_short'), filter=Q(units_toptask__top_task_short__isnull=False)),
                total_lowest = Count('units_toptask__toptask_gsmWorkplan__lowest_task_short', filter=Q(units_toptask__toptask_gsmWorkplan__lowest_task_short__isnull=False)),
                total_subactivities = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', filter=Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False))
                ).values('unit_code','total_output','total_kpi','total_TopTask','total_lowest','total_subactivities').distinct()
        
            tables_subactivities = Units.objects.exclude(units_toptask__toptask_gsmWorkplan__lowest_task_short__isnull=True).filter(unit_code__icontains=by_unit).annotate(  
                total_lowest = Count('units_toptask__toptask_gsmWorkplan__lowest_task_short', filter=Q(units_toptask__toptask_gsmWorkplan__lowest_task_short__isnull=False)),
                total_planned = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', filter=Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False)),
                total_completed = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', filter=Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name__icontains='Completed') & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False) & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False)),
                total_OnTrack = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', filter=Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name__icontains='On Track') & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False) & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False)),
                total_Not_Started = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', filter=Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name__icontains='Not Started') & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False)& Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False)),
                total_Issues = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', filter=Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name__icontains='Stalled') & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False)& Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False))
                ).values('unit_code','total_lowest','total_planned','total_completed','total_OnTrack','total_Not_Started','total_Issues').distinct()
            
            completedByCountries= Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__isnull=False) &  Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains=by_unit)& Q(statuts_operwork__completion_date__lte=end_date) & Q(statut_name__icontains='Completed')).annotate(nb_sub_activity=Count('statuts_operwork__sub_activity')).values(
               # 'statut_name',
                'statuts_operwork__country__country_name',
                'nb_sub_activity'
                ).distinct()
            notStartedByCountries= Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__isnull=False)  & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains=by_unit)& Q(statuts_operwork__completion_date__lte=end_date) & Q(statut_name__icontains='Not Started')).annotate(nb_sub_activity=Count('statuts_operwork__sub_activity')).values(
                #'statut_name',
                'statuts_operwork__country__country_name',
                'nb_sub_activity'
                ).distinct()
            
            subActByCountries= Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__isnull=False)  & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains=by_unit)& Q(statuts_operwork__completion_date__lte=end_date)).values(
               # 'statut_name',
                'statuts_operwork__country__country_name'
                ).annotate(nb_sub_activity=Count('statuts_operwork__sub_activity'))
      
            notStartedByCountries = pd.DataFrame(notStartedByCountries)
            # parsing the DataFrame in json format. 
            json_records = notStartedByCountries.reset_index().to_json(orient ='records') 
            #data_NumberTask = [] 
            notStartedByCountries = json.loads(json_records) 
            
            #5. DATASET 5
            kpi_results = Kpi.objects.prefetch_related('kpis_kpiAchive').filter(Q(unit__unit_code__icontains=by_unit)& Q(kpis_kpiAchive__report_date__lte=end_date)).values(
                'kpi_description', 'kpis_kpiAchive__report_date', 'kpis_kpiAchive__kpi_baseline', 'kpis_kpiAchive__kpi_target', 'kpis_kpiAchive__report_resut', 'kpis_kpiAchive__report_comment'
            )
         #   datasetTask = GsmWorkplan.objects.exclude(unit__unit_code__isnull=True).filter(Q(unit__unit_code__icontains=by_unit) & Q(gsmWorkplan_operw__completion_date__lte=end_date) & Q(gsmWorkplan_operw__sub_activity__isnull=False) & Q(unit__unit_code__isnull=False)).prefetch_related('gsmWorkplan_operw').distinct().order_by('-gsmWorkplan_operw__completion_date')
            dataset_byUnit = Statutworkplan.objects.exclude(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__isnull=True).filter(Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains=by_unit)& Q(statuts_operwork__completion_date__lte=end_date)).values(
                'statuts_operwork__gsmWorkplan__toptask__output__output_code','statuts_operwork__gsmWorkplan__toptask__top_task_description', 'statuts_operwork__gsmWorkplan__lowest_task_description','statuts_operwork__sub_activity','statuts_operwork__responsable','statuts_operwork__completion_date','statut_name','statuts_operwork__country__country_code','statuts_operwork__comments'
            )
            dataset_byUnit_Next = Statutworkplan.objects.exclude(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__isnull=True).filter(Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains=by_unit) & Q(statuts_operwork__completion_date__lte=one_month_next) & Q(statut_name__icontains='Not Started')).values(
                'statuts_operwork__gsmWorkplan__toptask__output__output_code','statuts_operwork__gsmWorkplan__toptask__top_task_description', 'statuts_operwork__gsmWorkplan__lowest_task_description','statuts_operwork__sub_activity','statuts_operwork__responsable','statuts_operwork__completion_date','statut_name','statuts_operwork__country__country_code','statuts_operwork__comments'
            )
               
            # METHOD TO GENERATE WORD FILE
          
            if form.is_valid():
                    
                data={'form': SelectUnitForm(),
                    'units':units, 
                    # 'gsmWorkplans':gsmWorkplans, 
                    'dt_fig_unit':dt_fig_unit,
                    'one_month_ago':one_month_ago, 
                    'end_date':end_date,
                    'date_completion__max':date_completion__max,
                    'line_chart_unit':line_chart_unit,
                    'dt_pie_status':dt_pie_status,
                    'dataset_byUnit':dataset_byUnit,
                    'notStartedByCountries':notStartedByCountries,
                    'df_subActByCountries':df_subActByCountries
                    } 
        except:
                 messages.success(request, 'Please select UNIT')
                    
                 data={'form': SelectUnitForm(),
                        'units':units, 
                        # 'gsmWorkplans':gsmWorkplans, 
                        'dt_fig_unit':dt_fig_unit,
                        'one_month_ago':one_month_ago, 
                        'end_date':end_date,
                        'date_completion__max':date_completion__max,
                        'line_chart_unit':line_chart_unit,
                        'dt_pie_status':dt_pie_status,
                        'dataset_byUnit':dataset_byUnit,
                        'notStartedByCountries':notStartedByCountries,
                        'df_subActByCountries':df_subActByCountries
                        } 
                    
        return render(request,'pages/forms/workplans/sub_activity_level.html',data)

#Method to generate the pdf for VID unit sub activities views from html
@login_required
def VID_pdf_activities(request):
    template_path= 'pages/pdfTemplate/vid_activity.html'
    
    code_unit='VID'
    if code_unit:
       data_unit =list(Units.objects.prefetch_related('units_toptask').filter(Q(unit_code__icontains=code_unit)& Q(units_toptask__toptask_gsmWorkplan__lowest_task_description__isnull=False)).distinct())
       data_outputs = Outputworkplan.objects.filter(Q(outputs_toptask__unit__unit_code__icontains=code_unit)).prefetch_related('outputs_toptask').distinct().order_by("outputs_toptask__output__output_code") 
     
    context = {
               'data_unit':data_unit,
               'data_outputs':data_outputs,
               'code_unit':code_unit
               }
    #create a Django response object , and specify content_type as pdf
    response = HttpResponse(content_type='application/pdf')
    # if download:
    # response['Content-Disposition']='attachment; filename="report.pdf"'
    # if display:
    response['Content-Disposition']='filename="VID_subactivity.pdf"'
    #find the temple and render it
    template = get_template(template_path)
    html = template.render(context)
    #create pdf
    pisa_status = pisa.CreatePDF(
        html, dest=response
    )
    # if error then show  some funy view
    if pisa_status.err:
        return HttpResponse('We had some errors <pre>'+html+'</pre>')
    return response

#Method to generate the pdf for HPD unit sub activities views from html
@login_required
def HPD_pdf_activities(request):
    template_path= 'pages/pdfTemplate/hpd_activity.html'
    
    code_unit='HPD'
    if code_unit:
       data_unit =list(Units.objects.prefetch_related('units_toptask').filter(Q(unit_code__icontains=code_unit)& Q(units_toptask__toptask_gsmWorkplan__lowest_task_description__isnull=False)).distinct())
       data_outputs = Outputworkplan.objects.filter(Q(outputs_toptask__unit__unit_code__icontains=code_unit) & Q(outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False)).prefetch_related('outputs_toptask').distinct().order_by("outputs_toptask__output__output_code") 
       data_tasks = Toptask.objects.filter(Q(unit__unit_code__icontains=code_unit)).select_related('output').prefetch_related('toptask_gsmWorkplan').distinct().order_by("output__output_code") 
      
    context = {
               'data_unit':data_unit,
               'data_outputs':data_outputs,
               'code_unit':code_unit,
               'data_tasks':data_tasks
               }
    #create a Django response object , and specify content_type as pdf
    response = HttpResponse(content_type='application/pdf')
    # if download:
    # response['Content-Disposition']='attachment; filename="report.pdf"'
    # if display:
    response['Content-Disposition']='filename="HPD_subactivity.pdf"'
    #find the temple and render it
    template = get_template(template_path)
    html = template.render(context)
    #create pdf
    pisa_status = pisa.CreatePDF(
        html, dest=response
    )
    # if error then show  some funy view
    if pisa_status.err:
        return HttpResponse('We had some errors <pre>'+html+'</pre>')
    return response


#Method to generate the pdf for HPD unit sub activities views from html
@login_required
def UHU_pdf_activities(request):
    template_path= 'pages/pdfTemplate/uhu_activity.html'
   
    code_unit='UHU'
    if code_unit:
       data_unit =list(Units.objects.prefetch_related('units_toptask').filter(Q(unit_code__icontains=code_unit)& Q(units_toptask__toptask_gsmWorkplan__lowest_task_description__isnull=False)).distinct())
       data_outputs = Outputworkplan.objects.filter(Q(outputs_toptask__unit__unit_code__icontains=code_unit)).prefetch_related('outputs_toptask').distinct().order_by("outputs_toptask__output__output_code") 
     
    context = {
               'data_unit':data_unit,
               'data_outputs':data_outputs,
               'code_unit':code_unit
               }
    #create a Django response object , and specify content_type as pdf
    response = HttpResponse(content_type='application/pdf')
    # if download:
    # response['Content-Disposition']='attachment; filename="report.pdf"'
    # if display:
    response['Content-Disposition']='filename="UHU_subactivity.pdf"'
    #find the temple and render it
    template = get_template(template_path)
    html = template.render(context)
    #create pdf
    pisa_status = pisa.CreatePDF(
        html, dest=response
    )
    # if error then show  some funy view
    if pisa_status.err:
        return HttpResponse('We had some errors <pre>'+html+'</pre>')
    return response

#Method to generate the pdf for HPD unit sub activities views from html
@login_required
def NUT_pdf_activities(request):
    template_path= 'pages/pdfTemplate/nut_activity.html'
    
    code_unit='NUT'
    if code_unit:
       data_unit =list(Units.objects.prefetch_related('units_toptask').filter(Q(unit_code__icontains=code_unit)& Q(units_toptask__toptask_gsmWorkplan__lowest_task_description__isnull=False)).distinct())
       data_outputs = Outputworkplan.objects.filter(Q(outputs_toptask__unit__unit_code__icontains=code_unit)).prefetch_related('outputs_toptask').distinct().order_by("outputs_toptask__output__output_code") 
     
    context = {
               'data_unit':data_unit,
               'data_outputs':data_outputs,
               'code_unit':code_unit
               }
    #create a Django response object , and specify content_type as pdf
    response = HttpResponse(content_type='application/pdf')
    # if download:
    # response['Content-Disposition']='attachment; filename="report.pdf"'
    # if display:
    response['Content-Disposition']='filename="NUT_subactivity.pdf"'
    #find the temple and render it
    template = get_template(template_path)
    html = template.render(context)
    #create pdf
    pisa_status = pisa.CreatePDF(
        html, dest=response
    )
    # if error then show  some funy view
    if pisa_status.err:
        return HttpResponse('We had some errors <pre>'+html+'</pre>')
    return response

#Method to generate the pdf for HPD unit sub activities views from html
@login_required
def TNR_pdf_activities(request):
    template_path= 'pages/pdfTemplate/tnr_activity.html'
    
    code_unit='TNR'
    if code_unit:
       data_unit =list(Units.objects.prefetch_related('units_toptask').filter(Q(unit_code__icontains=code_unit)& Q(units_toptask__toptask_gsmWorkplan__lowest_task_description__isnull=False)).distinct())
       data_outputs = Outputworkplan.objects.filter(Q(outputs_toptask__unit__unit_code__icontains=code_unit)).prefetch_related('outputs_toptask').distinct().order_by("outputs_toptask__output__output_code") 
     
    context = {
               'data_unit':data_unit,
               'data_outputs':data_outputs,
               'code_unit':code_unit
               }
    #create a Django response object , and specify content_type as pdf
    response = HttpResponse(content_type='application/pdf')
    # if download:
    # response['Content-Disposition']='attachment; filename="report.pdf"'
    # if display:
    response['Content-Disposition']='filename="TNR_subactivity.pdf"'
    #find the temple and render it
    template = get_template(template_path)
    html = template.render(context)
    #create pdf
    pisa_status = pisa.CreatePDF(
        html, dest=response
    )
    # if error then show  some funy view
    if pisa_status.err:
        return HttpResponse('We had some errors <pre>'+html+'</pre>')
    return response

# Method to add Meeting project
@login_required
@group_required('Administrative Assistance')
#@permission_required('myuhp.add_meetingproject', raise_exception=True)
def meeting_add_project(request):
    context={}
    form = MeetingProjectForm()
    meetingProjects = MeetingProject.objects.all()
    
    if request.method=='GET':
        st=request.GET.get('chair_name')
        if st!=None:
            meetingProjects= MeetingProject.objects.filter(chair_name__icontains=st)
    
    page = request.GET.get('page')
    num_of_items = 50
    # paginator of country 
    paginator = Paginator(meetingProjects, num_of_items)
       
    try:
        meetingProjects = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        meetingProjects = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        meetingProjects = paginator.page(page)
  
    context['meetingProjects'] = meetingProjects
    context['paginator']=paginator
    context['title'] ='home'
    if request.method == 'POST':
        if 'save' in request.POST:
            pk = request.POST.get('save')
            if not pk:
                form=MeetingProjectForm(request.POST)
            else:
                meetingProject = MeetingProject.objects.get(id = pk)
                form = MeetingProjectForm(request.POST ,instance = meetingProject) 
            form.save()
            form = MeetingProjectForm()
            messages.success(request, 'Meeting project saved or updated successfully')
        elif 'delete' in request.POST:
            pk = request.POST.get('delete')
            meetingProject = MeetingProject.objects.get(id = pk)
            meetingProject.delete()
            messages.success(request, 'Meeting project deleted successfully')
           
            
        elif 'edit' in request.POST:
            pk = request.POST.get('edit')
            meetingProject = MeetingProject.objects.get(id = pk)
            form = MeetingProjectForm(instance = meetingProject)
          
    context['form'] = form
   
    return render(request, 'pages/meeting/meeting_project.html', context)
             
# Method to add Meeting Discussion 
#@permission_required('myuhp.add_surveydataset',raise_exception=True)
@login_required
@group_required('Administrative Assistance')
def meeting_add_data(request):
    context={}
    by_name_meeting={}
    end_date={}
    form_meeting = SelectMeetingForm(request.GET or None)
    form = MeetingDiscussionForm()
    meetingDiscussions = MeetingDiscussion.objects.all()
  
    if request.method=='GET':
        by_name_meeting =  request.GET.get('by_name_meeting')
        end_date =  request.GET.get('end')
        st=request.GET.get('responsible')
        if st!=None:
            meetingDiscussions= MeetingDiscussion.objects.filter(responsible__icontains=st)
            
        if by_name_meeting!=None:
            meetingDiscussions = MeetingDiscussion.objects.all().filter(Q(meetingProject__id__icontains=by_name_meeting)& Q(meetingProject__date_meeting__lte=end_date)).order_by("meetingProject")    
            
    
#    page = request.GET.get('page')
#    num_of_items = 100
#    # paginator of country 
#    paginator = Paginator(meetingDiscussions, num_of_items)
       
#    try:
#        meetingDiscussions = paginator.page(page)
#    except PageNotAnInteger:
#        page = 1
#        meetingDiscussions = paginator.page(page)
#    except EmptyPage:
#        page = paginator.num_pages
#        meetingDiscussions = paginator.page(page)
            
    context['meetingDiscussions'] = meetingDiscussions
#    context['paginator']=paginator
    #context['form']=form
    context['title'] ='home'
    if request.method == 'POST':
        
        if 'save' in request.POST:
            pk = request.POST.get('save')
            if not pk:
                form=MeetingDiscussionForm(request.POST)
            else:
                meetingDiscussion = MeetingDiscussion.objects.get(id = pk)
                form = MeetingDiscussionForm(request.POST ,instance = meetingDiscussion) 
            form.save()
            form = MeetingDiscussionForm()
            messages.success(request, 'Meeting discussion saved or updated successfully')
        elif 'delete' in request.POST:
            pk = request.POST.get('delete')
            meetingDiscussion = MeetingDiscussion.objects.get(id = pk)
            meetingDiscussion.delete()
            messages.success(request, 'Meeting discussion deleted successfully')
       
        elif 'edit' in request.POST:
            pk = request.POST.get('edit')
            meetingDiscussion = MeetingDiscussion.objects.get(id = pk)
            form = MeetingDiscussionForm(instance = meetingDiscussion)

    context['form'] = form
    context['form_meeting'] = form_meeting
    context['by_name_meeting'] = by_name_meeting
    context['end_date'] = end_date
   
    return render(request, 'pages/meeting/meeting_discussion.html', context)

#Method to export MEETING DATASET ALL DATA on Excel file
def export_meeting_to_excel(request):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="meeting.xlsx"'

    wb = Workbook()
    ws = wb.active
    ws.title = "Meeting"

    # Add headers
    headers = [#"output",
             #  "kpi",
               "date_meeting", 
               "name_meeting", 
            #   "type_meeting", 
               "objective_meeting", "taking_place", "chair_name", "note_taker","participants_list","topic_discussion","summary_discussion","recommandation","actions_points","responsible","action_deadlines","action_status","feedback_discussion"]
    ws.append(headers)

    # Add data from the model
    meetingDiscussions = MeetingDiscussion.objects.all()
    for s in meetingDiscussions:
        ws.append([#s.meetingProject.output,
                #   s.meetingProject.kpi, 
                   s.meetingProject.date_meeting, 
                   s.meetingProject.name_meeting, 
                 #  s.meetingProject.type_meeting,
                   s.meetingProject.objective_meeting,
                   s.meetingProject.taking_place, 
                   s.meetingProject.chair_name,
                   s.meetingProject.note_taker, 
                   s.meetingProject.participants_list, 
                   s.topic_discussion,
                   s.summary_discussion, 
                   s.recommandation, 
                   s.actions_points, 
                   s.responsible,s.action_deadlines,s.action_status,s.feedback_discussion])

    # Save the workbook to the HttpResponse
    wb.save(response)
    return response

 ## METHOD TO DISPLAY ALL MEETING REPORT SUBMITTED
#@permission_required('myuhp.view_meetingproject',raise_exception=True)
@login_required
@group_required('Administrative Assistance')
def meeting_report(request):
    meetingProjects = MeetingProject.objects.all().order_by("date_meeting") 
    if request.method=='GET':
        st=request.GET.get('name_meeting')
        if st!=None:
            meetingProjects= MeetingProject.objects.filter(name_meeting__icontains=st)
    
    page = request.GET.get('page')
    num_of_items = 10
    # paginator of country 
    paginator = Paginator(meetingProjects, num_of_items)
       
    try:
        meetingProjects = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        meetingProjects = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        meetingProjects = paginator.page(page)
            
    data={'meetingProjects': meetingProjects, 'paginator':paginator}
    
    return render(request, 'pages/meeting/meeting_report.html', data)

#METHOD TO VIEW SINGLE PAGE FOR  MEETING REPORT  
@login_required
@group_required('Administrative Assistance')
def single_meeting_page(request, pk):
    try:
        meetingProject=MeetingProject.objects.get(id=pk) 
        meetingProjects =MeetingProject.objects.filter(Q(id__icontains=pk) & Q(meetingProject_discussion__summary_discussion__isnull=False)).values(
        'date_meeting',
        'name_meeting',
     #   'type_meeting',
        'objective_meeting',
        'taking_place',
        'chair_name',
        'note_taker',
        'participants_list',
        'meetingProject_discussion__topic_discussion',
        'meetingProject_discussion__summary_discussion',
        'meetingProject_discussion__recommandation',
        'meetingProject_discussion__actions_points',
        'meetingProject_discussion__responsible',
        'meetingProject_discussion__action_deadlines',
        'meetingProject_discussion__action_status',
        'meetingProject_discussion__feedback_discussion'
    ).order_by('-date_meeting')
        
        
        page = request.GET.get('page')
        num_of_items = 10
        # paginator of country 
        paginator = Paginator(meetingProjects, num_of_items)
        
        try:
            meetingProjects = paginator.page(page)
        except PageNotAnInteger:
            page = 1
            meetingProjects = paginator.page(page)
        except EmptyPage:
            page = paginator.num_pages
            meetingProjects = paginator.page(page)
        
        data={'meetingProject':meetingProject, 'meetingProjects':meetingProjects, 'paginator':paginator}   
    except:
        data={'message':'The Meeting Report you requested doesn t\' exist or isn t in the submitted results'}
    
       
    return render(request,'pages/meeting/single_report_page.html',data)

#METHOD TO EXPORT MEETING REPORT AND DETAILS TO EXCEL FILE 
@login_required
@group_required('Administrative Assistance')
def export_meeting_page(request,pk):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="meeting_report.xls"'

    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet('Meeting')

    # Sheet header, first row
    row_num = 0

    font_style = xlwt.XFStyle()
    font_style.font.bold = True

    columns = [
        #"output",
             #  "kpi",
               "date_meeting", 
               "name_meeting", 
               "type_meeting", 
               "objective_meeting", "taking_place", "chair_name", "note_taker","participants_list","topic_discussion","summary_discussion","recommandation","actions_points","responsible","action_deadlines","action_status","feedback_discussion"
        ]

    for col_num in range(len(columns)):
        ws.write(row_num, col_num, columns[col_num], font_style)

    # Sheet body, remaining rows
    font_style = xlwt.XFStyle()

    rows = MeetingProject.objects.filter(Q(id__icontains=pk) & Q(meetingProject_discussion__summary_discussion__isnull=False)).all().values_list(
        'date_meeting',
        'name_meeting',
        'type_meeting__group_meeting',
        'objective_meeting',
        'taking_place',
        'chair_name',
        'note_taker',
        'participants_list',
        'meetingProject_discussion__topic_discussion',
        'meetingProject_discussion__summary_discussion',
        'meetingProject_discussion__recommandation',
        'meetingProject_discussion__actions_points',
        'meetingProject_discussion__responsible',
        'meetingProject_discussion__action_deadlines',
        'meetingProject_discussion__action_status',
        'meetingProject_discussion__feedback_discussion'
    ).order_by('-date_meeting')
    for row in rows:
        row_num += 1
        for col_num in range(len(row)):
            ws.write(row_num, col_num, row[col_num], font_style)

    wb.save(response)
    return response


# METHODE TO EXPORT MEETING REPORT TO WORD FILE
@login_required
@group_required('Administrative Assistance')
def meeting_doc_report(request, pk):   
   # document.add_picture('static/img/logos.png', width=Inches(1.25))
    meetingProject=MeetingProject.objects.get(id=pk)
  #  meetingProject=MeetingProject.objects.all().values(
  #      'date_meeting',
  #      'name_meeting',
  #      'type_meeting',
  #      'objective_meeting',
  #      'taking_place',
  #      'chair_name',
  #      'note_taker',
  #      'participants_list'
  #  )
    meetingProjects =MeetingProject.objects.filter(Q(id__icontains=pk) & Q(meetingProject_discussion__summary_discussion__isnull=False)).values(
        'meetingProject_discussion__topic_discussion',
        'meetingProject_discussion__summary_discussion',
        'meetingProject_discussion__recommandation',
        'meetingProject_discussion__actions_points',
        'meetingProject_discussion__responsible',
        'meetingProject_discussion__action_deadlines',
        'meetingProject_discussion__action_status',
        'meetingProject_discussion__feedback_discussion'
    ).order_by('meetingProject_discussion__topic_discussion')
    
    if meetingProject is not None:
        document = Document()
        section = document.sections
        section.page_height = 8
        section.page_width = 11
        section = document.sections[-1]
        
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height=section.page_height, section.page_width
        
        section.page_width = new_width
        section.page_height = new_height
        
       
        document.add_heading('Meeting Report submitted', 0)
    #  p = document.add_paragraph('A plain paragraph having some ')
    #  p.add_run('bold').bold = True
    #  p.add_run(' and some ')
    #   p.add_run('italic.').italic = True

    #  document.add_heading('Heading, level 1', level=1)
    #  document.add_paragraph('Intense quote', style='IntenseQuote')

    #  document.add_paragraph(
    #  'first item in unordered list', style='ListBullet'
    #  )
    #  document.add_paragraph(
    #  'first item in ordered list', style='ListNumber'
    #  )
        DateOfMeeting = document.add_paragraph('Date of meeting: ' , style='ListBullet')
        DateOfMeeting.add_run(str(meetingProject.date_meeting)).bold = True
        
        NameMeeting = document.add_paragraph('Title of meeting: ' , style='ListBullet')
        NameMeeting.add_run(str(meetingProject.name_meeting)).bold = True
            
        ObjectiveMeeting = document.add_paragraph('Objective: ' , style='ListBullet')
        ObjectiveMeeting.add_run(str(meetingProject.objective_meeting)).bold = True
            
        TakingPlace = document.add_paragraph('Taking place: ' , style='ListBullet')
        TakingPlace.add_run(str(meetingProject.taking_place)).bold = True
            
        ChairName = document.add_paragraph('Chair Name: ' , style='ListBullet')
        ChairName.add_run(str(meetingProject.chair_name)).bold = True
                
        NoteTaker = document.add_paragraph('NoteTaker Name: ' , style='ListBullet')
        NoteTaker.add_run(str(meetingProject.note_taker)).bold = True
            
        Participants = document.add_paragraph('Participants list: ' , style='ListBullet')
        Participants.add_run(str(meetingProject.participants_list)).bold = True
                      
    #  for j in range(len(meetingProject)):
    #       DateOfMeeting = document.add_paragraph('Date of meeting: ' , style='ListBullet')
    #       DateOfMeeting.add_run(str(meetingProject[j]['date_meeting'])).bold = True
            
    #       NameMeeting = document.add_paragraph('Title of meeting: ' , style='ListBullet')
    #       NameMeeting.add_run(str(meetingProject[j]['name_meeting'])).bold = True
            
    #       ObjectiveMeeting = document.add_paragraph('Objective: ' , style='ListBullet')
    #       ObjectiveMeeting.add_run(str(meetingProject[j]['objective_meeting'])).bold = True
            
    #       TakingPlace = document.add_paragraph('Taking place: ' , style='ListBullet')
    #       TakingPlace.add_run(str(meetingProject[j]['taking_place'])).bold = True
            
    #       ChairName = document.add_paragraph('Chair Name: ' , style='ListBullet')
    #       ChairName.add_run(str(meetingProject[j]['chair_name'])).bold = True
                
    #       NoteTaker = document.add_paragraph('NoteTaker Name: ' , style='ListBullet')
    #       NoteTaker.add_run(str(meetingProject[j]['note_taker'])).bold = True
            
    #       Participants = document.add_paragraph('Participants list: ' , style='ListBullet')
    #       Participants.add_run(str(meetingProject[j]['participants_list'])).bold = True
                                
        document.add_paragraph('Discussion:', style='ListBullet' )                     

        table = document.add_table(rows=1, cols=8)
        table.style='Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Topics for discussion'
        hdr_cells[1].text = 'Summary of discussion'
        hdr_cells[2].text = 'Recommandations'
        hdr_cells[3].text = 'Actions points'
        hdr_cells[4].text = 'Responsible'
        hdr_cells[5].text = 'Deadlines'
        hdr_cells[6].text = 'Status'
        hdr_cells[7].text = 'Feedback of discussion'
        
        for i in range(len(meetingProjects)):
            row_cells = table.add_row().cells
            row_cells[0].text = str(meetingProjects[i]['meetingProject_discussion__topic_discussion'])
            row_cells[1].text = str(meetingProjects[i]['meetingProject_discussion__summary_discussion'])
            row_cells[2].text = str(meetingProjects[i]['meetingProject_discussion__recommandation'])
            row_cells[3].text = str(meetingProjects[i]['meetingProject_discussion__actions_points'])
            row_cells[4].text = str(meetingProjects[i]['meetingProject_discussion__responsible'])
            row_cells[5].text = str(meetingProjects[i]['meetingProject_discussion__action_deadlines'])
            row_cells[6].text = str(meetingProjects[i]['meetingProject_discussion__action_status'])
            row_cells[7].text = str(meetingProjects[i]['meetingProject_discussion__feedback_discussion'])
            
        document.add_page_break()

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=Meeting_report_'+str(meetingProject.note_taker)+'.docx'
        document.save(response)

        return response
    else:
        return HttpResponse("Failed to Download Meeting report")
    
    
# METHOD TO CONVERT HTML TO PDF FILE
def pdf(request):
   # pdf=htmlTopdf(dashboard.html)
    return HttpResponse(pdf, content_type='application/pdf')

# METHOD TO CREATE BRIEFING PROJECT


#@permission_required('myuhp.add_briefingproject', raise_exception=True)
@login_required
@group_required('Technical officer')
def briefing_add_project(request):
    context={}
    form = BriefingProjectForm()
    briefingProjects = BriefingProject.objects.all()
    
    if request.method=='GET':
        st=request.GET.get('briefing_title')
        if st!=None:
            briefingProjects= BriefingProject.objects.filter(briefing_title__icontains=st)
    
    page = request.GET.get('page')
    num_of_items = 100
    # paginator of country 
    paginator = Paginator(briefingProjects, num_of_items)
       
    try:
        briefingProjects = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        briefingProjects = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        briefingProjects = paginator.page(page)
  
    context['briefingProjects'] = briefingProjects
    context['paginator']=paginator
    context['title'] ='home'
    if request.method == 'POST':
        if 'save' in request.POST:
            pk = request.POST.get('save')
            if not pk:
                form=BriefingProjectForm(request.POST)
            else:
                briefingProject = MeetingProject.objects.get(id = pk)
                form = BriefingProjectForm(request.POST ,instance = briefingProject) 
            form.save()
            form = BriefingProjectForm()
            messages.success(request, 'Briefing project saved or updated successfully')
        elif 'delete' in request.POST:
            pk = request.POST.get('delete')
            briefingProject = BriefingProject.objects.get(id = pk)
            briefingProject.delete()
            messages.success(request, 'Briefing project deleted successfully')
           
            
        elif 'edit' in request.POST:
            pk = request.POST.get('edit')
            briefingProject = BriefingProject.objects.get(id = pk)
            form = BriefingProjectForm(instance = briefingProject)
          
    context['form'] = form
   
    return render(request, 'pages/briefing/briefing_project.html', context)


# METHOD TO CREATE BACKGROUND FOR BRIEFING
#@permission_required('myuhp.add_briefingbackground',raise_exception=True)
@login_required
@group_required('Technical officer')
def briefing_add_data(request):
    context={}
    form = BriefingBackgroundForm()
    briefingBackgrounds = BriefingBackground.objects.all()
  
    if request.method=='GET':
        st=request.GET.get('subject_background')
        if st!=None:
            briefingBackgrounds= BriefingBackground.objects.filter(subject_background__icontains=st)
    
    page = request.GET.get('page')
    num_of_items = 100
    # paginator of country 
    paginator = Paginator(briefingBackgrounds, num_of_items)
       
    try:
        briefingBackgrounds = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        briefingBackgrounds = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        briefingBackgrounds = paginator.page(page)
            
    context['briefingBackgrounds'] = briefingBackgrounds
    context['paginator']=paginator
    context['form']=form
    context['title'] ='home'
    if request.method == 'POST':
        if 'save' in request.POST:
            pk = request.POST.get('save')
            if not pk:
                form=BriefingBackgroundForm(request.POST)
            else:
                briefingBackground = BriefingBackground.objects.get(id = pk)
                form = BriefingBackgroundForm(request.POST ,instance = briefingBackground) 
            form.save()
            form = BriefingBackgroundForm()
            messages.success(request, 'Brefing background saved or updated successfully')
        elif 'delete' in request.POST:
            pk = request.POST.get('delete')
            briefingBackground = BriefingBackground.objects.get(id = pk)
            briefingBackground.delete()
            messages.success(request, 'Brefing background deleted successfully')
        elif 'edit' in request.POST:
            pk = request.POST.get('edit')
            briefingBackground = BriefingBackground.objects.get(id = pk)
            form = BriefingBackgroundForm(instance = briefingBackground)
          
    context['form'] = form
   
    return render(request, 'pages/briefing/briefing_background.html', context)


#METHOD TO EXPORT BRIEFING DATASET ALL DATA ON EXCEL FILE
@login_required
@group_required('Technical officer')
def export_briefingNote_to_excel(request):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="briefingNote.xlsx"'
    wb = Workbook()
    ws = wb.active
    ws.title = "BriefingNote"
    # Add headers
    headers = [
               "briefing_title", 
               "unit", 
               "start_date", 
               "end_date", 
               "reporting_date", 
               "subject_background",
               "specific_topic",
               "accomplished_last_period",
               "planned_next_steps",
            #   "output",
            #   "kpi",
               "comment_background"
              ]
    ws.append(headers)

    # Add data from the model
    briefingBackgrounds = BriefingBackground.objects.all()
    for s in briefingBackgrounds:
        ws.append([
                   s.briefingProject.briefing_title, 
                   s.briefingProject.unit.unit_code, 
                   s.briefingProject.start_date,
                   s.briefingProject.end_date, 
                   s.briefingProject.reporting_date,
                   s.subject_background, 
                   s.specific_topic, 
                   s.accomplished_last_period, 
                   s.planned_next_steps, 
                   s.comment_background])
    # Save the workbook to the HttpResponse
    wb.save(response)
    return response

## METHOD TO DISPLAY ALL BRIEFING NOTES SUBMITTED
#@permission_required('myuhp.view_brefingproject',raise_exception=True)
@login_required
@group_required('Technical officer')
def briefing_report(request):
    briefingProjects = BriefingProject.objects.all().order_by("reporting_date") 
    if request.method=='GET':
        st=request.GET.get('briefing_title')
        if st!=None:
            briefingProjects= BriefingProject.objects.filter(briefing_title__icontains=st)
    
    page = request.GET.get('page')
    num_of_items = 10
    # paginator of country 
    paginator = Paginator(briefingProjects, num_of_items)
       
    try:
        briefingProjects = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        briefingProjects = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        briefingProjects = paginator.page(page)
            
    data={'briefingProjects': briefingProjects, 'paginator':paginator}
    
    return render(request, 'pages/briefing/briefing_report.html', data)

@login_required
@group_required('Technical officer')
def single_briefing_page(request, pk):
    try:
        briefingProject = BriefingProject.objects.get(id=pk) 
        briefingProjects = BriefingProject.objects.filter(Q(id__icontains=pk)).values(
        'briefing_title',
        'unit',
     #   'type_meeting',
        'start_date',
        'end_date',
        'reporting_date',
        'briefingProject_background__subject_background',
        'briefingProject_background__specific_topic',
        'briefingProject_background__accomplished_last_period',
        'briefingProject_background__planned_next_steps',
        'briefingProject_background__output__output_code',
        'briefingProject_background__kpi__kpi_code',
        'briefingProject_background__comment_background'
    ).distinct()
        page = request.GET.get('page')
        num_of_items = 20
        # paginator of country 
        paginator = Paginator(briefingProjects, num_of_items)
        
        try:
            briefingProjects = paginator.page(page)
        except PageNotAnInteger:
            page = 1
            briefingProjects = paginator.page(page)
        except EmptyPage:
            page = paginator.num_pages
            briefingProjects = paginator.page(page)
        
        data={'briefingProject':briefingProject, 'briefingProjects':briefingProjects, 'paginator':paginator}   
    except:
        data={'message':'The Briefing note you requested doesn t\' exist or isn t in the submitted results'}
    
       
    return render(request,'pages/briefing/single_briefing_note.html',data)

#METHOD TO EXPORT MEETING REPORT AND DETAILS TO EXCEL FILE 
@login_required
@group_required('Technical officer')
def export_briefing_page(request,pk):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="briefing_report.xls"'

    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet('Meeting')

    # Sheet header, first row
    row_num = 0

    font_style = xlwt.XFStyle()
    font_style.font.bold = True

    columns = [
               "briefing_title", 
               "unit", 
               "start_date", 
               "end_date", 
               "reporting_date", 
               "subject_background",
               "specific_topic",
               "accomplished_last_period",
               "planned_next_steps",
               "output",
               "kpi",
               "comment_background"
        ]

    for col_num in range(len(columns)):
        ws.write(row_num, col_num, columns[col_num], font_style)

    # Sheet body, remaining rows
    font_style = xlwt.XFStyle()

    rows = BriefingProject.objects.select_related('unit').filter(Q(id__icontains=pk)).all().values_list(
       'briefing_title',
        'unit__unit_code',
     #   'type_meeting',
        'start_date',
        'end_date',
        'reporting_date',
        'briefingProject_background__subject_background',
        'briefingProject_background__specific_topic',
        'briefingProject_background__accomplished_last_period',
        'briefingProject_background__planned_next_steps',
        'briefingProject_background__output__output_code',
        'briefingProject_background__kpi__kpi_code',
        'briefingProject_background__comment_background'
    ).order_by('-reporting_date')
    for row in rows:
        row_num += 1
        for col_num in range(len(row)):
            ws.write(row_num, col_num, row[col_num], font_style)

    wb.save(response)
    return response


# METHODE TO EXPORT MEETING REPORT TO WORD FILE
@login_required
@group_required('Technical officer')
def briefing_doc_report(request, pk):   
   # document.add_picture('static/img/logos.png', width=Inches(1.25))
    briefingProject=BriefingProject.objects.get(id=pk)
  #  meetingProject=MeetingProject.objects.all().values(
  #      'date_meeting',
  #      'name_meeting',
  #      'type_meeting',
  #      'objective_meeting',
  #      'taking_place',
  #      'chair_name',
  #      'note_taker',
  #      'participants_list'
  #  )
    briefingProjects =BriefingProject.objects.filter(Q(id__icontains=pk) & Q(briefingProject_background__subject_background__isnull=False)).values(
        'briefingProject_background__subject_background',
        'briefingProject_background__specific_topic',
        'briefingProject_background__accomplished_last_period',
        'briefingProject_background__planned_next_steps',
        'briefingProject_background__output',
        'briefingProject_background__kpi',
        'briefingProject_background__comment_background'
    ).order_by('briefingProject_background__subject_background')
    
    if briefingProject is not None:
        document = Document()
        section = document.sections
        section.page_height = 8
        section.page_width = 11
        section = document.sections[-1]
        
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height=section.page_height, section.page_width
        
        section.page_width = new_width
        section.page_height = new_height
        
       
       # document.add_heading('Briefing note background submitted', 0)
        Title_report = document.add_heading('Briefing note background submitted by '+str(briefingProject.unit), 0)
        Title_report.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # document.add_heading('Performance monitoring report', 0)
        EditReport= document.add_paragraph('Printed on: ')
        EditReport.add_run("%s" % date.today().strftime('%B %d, %Y')).bold = True
        EditReport.add_run(' From UHP website').italic = True
        EditReport.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # PUT THE NAME OF UNIT IN THE DOCUMENT
      #  AuthorReport = document.add_heading(str(briefingProject.unit), level=1)
      #  AuthorReport.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #  p = document.add_paragraph('A plain paragraph having some ')
    #  p.add_run('bold').bold = True
    #  p.add_run(' and some ')
    #   p.add_run('italic.').italic = True

    #  document.add_heading('Heading, level 1', level=1)
    #  document.add_paragraph('Intense quote', style='IntenseQuote')

    #  document.add_paragraph(
    #  'first item in unordered list', style='ListBullet'
    #  )
    #  document.add_paragraph(
    #  'first item in ordered list', style='ListNumber'
    #  )
        BriefingTitle = document.add_paragraph('Title of briefing note: ' , style='ListBullet')
        BriefingTitle.add_run(str(briefingProject.briefing_title)).bold = True
        
        StartDate = document.add_paragraph('Start date: ' , style='ListBullet')
        StartDate.add_run(str(briefingProject.start_date)).bold = True
            
        EndDate = document.add_paragraph('End date: ' , style='ListBullet')
        EndDate.add_run(str(briefingProject.end_date)).bold = True
            
        SumbissionDate = document.add_paragraph('Submission date: ' , style='ListBullet')
        SumbissionDate.add_run(str(briefingProject.reporting_date)).bold = True
            
        
                      
    #  for j in range(len(meetingProject)):
    #       DateOfMeeting = document.add_paragraph('Date of meeting: ' , style='ListBullet')
    #       DateOfMeeting.add_run(str(meetingProject[j]['date_meeting'])).bold = True
            
    #       NameMeeting = document.add_paragraph('Title of meeting: ' , style='ListBullet')
    #       NameMeeting.add_run(str(meetingProject[j]['name_meeting'])).bold = True
            
    #       ObjectiveMeeting = document.add_paragraph('Objective: ' , style='ListBullet')
    #       ObjectiveMeeting.add_run(str(meetingProject[j]['objective_meeting'])).bold = True
            
    #       TakingPlace = document.add_paragraph('Taking place: ' , style='ListBullet')
    #       TakingPlace.add_run(str(meetingProject[j]['taking_place'])).bold = True
            
    #       ChairName = document.add_paragraph('Chair Name: ' , style='ListBullet')
    #       ChairName.add_run(str(meetingProject[j]['chair_name'])).bold = True
                
    #       NoteTaker = document.add_paragraph('NoteTaker Name: ' , style='ListBullet')
    #       NoteTaker.add_run(str(meetingProject[j]['note_taker'])).bold = True
            
    #       Participants = document.add_paragraph('Participants list: ' , style='ListBullet')
    #       Participants.add_run(str(meetingProject[j]['participants_list'])).bold = True
                                
        document.add_paragraph('Briefing note background:', style='ListBullet' )                     
        
    # records = (
    #     (3, '101', 'Spam'),
    #     (7, '422', 'Eggs'),
    #     (4, '631', 'Spam, spam, eggs, and spam')
    #)
        table = document.add_table(rows=1, cols=7)
        table.style='Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Subject'
        hdr_cells[1].text = 'Specific topic'
        hdr_cells[2].text = 'Activities accomplished last period'
        hdr_cells[3].text = 'Activities planned current month'
        hdr_cells[4].text = 'Link to output'
        hdr_cells[5].text = 'Link to Kpi'
        hdr_cells[6].text = 'Comments'
       
        
        for i in range(len(briefingProjects)):
            row_cells = table.add_row().cells
            row_cells[0].text = str(briefingProjects[i]['briefingProject_background__subject_background'])
            row_cells[1].text = str(briefingProjects[i]['briefingProject_background__specific_topic'])
            row_cells[2].text = str(briefingProjects[i]['briefingProject_background__accomplished_last_period'])
            row_cells[3].text = str(briefingProjects[i]['briefingProject_background__planned_next_steps'])
            row_cells[4].text = str(briefingProjects[i]['briefingProject_background__output'])
            row_cells[5].text = str(briefingProjects[i]['briefingProject_background__kpi'])
            row_cells[6].text = str(briefingProjects[i]['briefingProject_background__comment_background'])
           
            
        document.add_page_break()

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=Briefing_note_'+str(briefingProject.unit) +str(briefingProject.reporting_date)+'.docx'
        document.save(response)

        return response
    else:
        return HttpResponse("Failed to Download Briefing note background")


# METHOD TO CREATE RISK REGISTER 
#@permission_required('myuhp.add_riskidentification', raise_exception=True)
@login_required
@group_required('Team Leader')
def risk_identification_add(request):
    context={}
    form_unit = SelectUnitForm(request.GET or None)
    by_unit =  request.GET.get('by_unit')
    end_date =  request.GET.get('end')
    
    
    form = RiskIdentificationForm()
    riskIdentifications = RiskIdentification.objects.all()
    
    if request.method=='GET':
        st=request.GET.get('risk_name')
        if st!=None:
            riskIdentifications= RiskIdentification.objects.filter(risk_name__icontains=st)
        
        if by_unit!=None:
            riskIdentifications = RiskIdentification.objects.all().filter(Q(unit__unit_code__icontains=by_unit)& Q(response_date__lte=end_date)).order_by("toptask")
    
   # page = request.GET.get('page')
   # num_of_items = 10
   
  #  paginator = Paginator(riskIdentifications, num_of_items)
       
 #   try:
  #      riskIdentifications = paginator.page(page)
  #  except PageNotAnInteger:
  #      page = 1
  #      riskIdentifications = paginator.page(page)
  #  except EmptyPage:
  #      page = paginator.num_pages
  #      riskIdentifications = paginator.page(page)
        
   # context['paginator']=paginator 
    context['riskIdentifications'] = riskIdentifications
    context['title'] ='home'
    if request.method == 'POST':
        if 'save' in request.POST:
            pk = request.POST.get('save')
            if not pk:
                form=RiskIdentificationForm(request.POST)
            else:
                riskIdentification = RiskIdentification.objects.get(id = pk)
                form = RiskIdentificationForm(request.POST ,instance = riskIdentification) 
            form.save()
            form = RiskIdentificationForm()
            messages.success(request, 'Risk name saved or updated successfully')
        elif 'delete' in request.POST:
            pk = request.POST.get('delete')
            riskIdentification = RiskIdentification.objects.get(id = pk)
            riskIdentification.delete()
            messages.success(request, 'Risk name deleted successfully')
                
        elif 'edit' in request.POST:
            pk = request.POST.get('edit')
            riskIdentification = RiskIdentification.objects.get(id = pk)
            form = RiskIdentificationForm(instance = riskIdentification)
          
    context['form'] = form
    context['form_unit'] = form_unit 
    context['by_unit'] = by_unit
    context['end_date'] = end_date

    return render(request, 'pages/riskRegister/risk_identification.html', context)


#METHOD TO EXPORT BRIEFING DATASET ALL DATA ON EXCEL FILE
def export_risk_to_excel(request):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="riskIdentification.xlsx"'

    wb = Workbook()
    ws = wb.active
    ws.title = "Risk"

    # Add headers
    headers = [
               "unit", 
               "toptask", 
               "risk_cause", 
               "risk_impact", 
               "risk_category", 
               "risk_status",
               "risk_occuring",
               "risk_rating",
               "risk_criticality",
               "response_decision",
               "risk_action",
               "budget",
               "response_date",
               "review_plan",
               "comments",
               "responsible",
               "monitoring",
               "risk_criticality_after"
              ]
    ws.append(headers)

    # Add data from the model
    riskIdentification = RiskIdentification.objects.all()
    for s in riskIdentification:
        ws.append([
                   s.unit.unit_code, 
                   s.toptask.top_task_description, 
                   s.risk_cause,
                   s.risk_impact, 
                   s.risk_category,
                   s.risk_status, 
                   s.risk_occuring, 
                   s.risk_rating, 
                   s.risk_criticality,
                   s.response_decision,
                   s.risk_action,
                   s.budget,
                   s.response_date,
                   s.review_plan,
                   s.comments,
                   s.responsible,
                   s.monitoring,
                   s.risk_criticality_after
                 
                 ])

    # Save the workbook to the HttpResponse
    wb.save(response)
    return response


#@permission_required('myuhp.view_riskidentification',raise_exception=True)
@login_required
@group_required('Team Leader')
def risk_report(request):
    units = Units.objects.all().order_by("unit_code") 
    if request.method=='GET':
        st=request.GET.get('unit_code')
        if st!=None:
            units= Units.objects.filter(unit_code__icontains=st)
    
    page = request.GET.get('page')
    num_of_items = 10
    # paginator of country 
    paginator = Paginator(units, num_of_items)
       
    try:
        units = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        units = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        units = paginator.page(page)
            
    data={'units': units, 'paginator':paginator}
    
    return render(request, 'pages/riskRegister/risk_report.html', data)


#METHOD TO VIEW SINGLE PAGE FOR  MEETING REPORT  
@login_required
@group_required('Team Leader')
def unit_risk_page(request, pk):
    try:
        unit=Units.objects.get(id=pk) 
        units =Units.objects.filter(Q(id__icontains=pk) & Q(units_risk__risk_name__isnull=False)).values(
        'unit_code',
        'unit_description',
        'units_risk__toptask__top_task_description',
        'units_risk__risk_name',
        'units_risk__risk_cause',
        'units_risk__risk_impact', 
        'units_risk__risk_category',
        'units_risk__risk_status', 
        'units_risk__risk_occuring', 
        'units_risk__risk_rating', 
        'units_risk__risk_criticality',
        'units_risk__response_decision',
        'units_risk__risk_action',
        'units_risk__budget',
        'units_risk__response_date',
        'units_risk__review_plan',
        'units_risk__comments',
        'units_risk__responsible',
        'units_risk__monitoring',
        'units_risk__risk_criticality_after'
    ).order_by('unit_code')
        
        page = request.GET.get('page')
        num_of_items = 10
        # paginator of country 
        paginator = Paginator(units, num_of_items)
        
        try:
            units = paginator.page(page)
        except PageNotAnInteger:
            page = 1
            units = paginator.page(page)
        except EmptyPage:
            page = paginator.num_pages
            units = paginator.page(page)
        
        data={'unit':unit, 'units':units, 'paginator':paginator}   
    except:
        data={'message':'The Risk register you requested doesn t\' exist or isn t in the submitted results'}
      
    return render(request,'pages/riskRegister/unit_report_risk.html',data)


#METHOD TO EXPORT UNIT RISK EGISTER TO EXCEL FILE 
@login_required
@group_required('Team Leader')
def export_unit_risk(request,pk):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="unit_riskRegister.xls"'

    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet('Risk_register')

    # Sheet header, first row
    row_num = 0

    font_style = xlwt.XFStyle()
    font_style.font.bold = True

    columns = [
        "unit code",
        "unit description",
        "Top task",
        "Risk name",
        "Risk cause",
        "Risk impact", 
        "Rrisk category",
        "Risk status", 
        "Risk occuring", 
        "Risk rating", 
        "Risk criticality",
        "Response decision",
        "Action",
        "Budget",
        "Rresponse date",
        "Review plan",
        "Comments",
        "Responsible",
        "Monitoring",
        "Criticality after"
        ]

    for col_num in range(len(columns)):
        ws.write(row_num, col_num, columns[col_num], font_style)

    # Sheet body, remaining rows
    font_style = xlwt.XFStyle()

    rows = Units.objects.filter(Q(id__icontains=pk) & Q(units_risk__risk_name__isnull=False)).all().values_list(
        'unit_code',
        'unit_description',
        'units_risk__toptask__top_task_description',
        'units_risk__risk_name',
        'units_risk__risk_cause',
        'units_risk__risk_impact', 
        'units_risk__risk_category',
        'units_risk__risk_status', 
        'units_risk__risk_occuring', 
        'units_risk__risk_rating', 
        'units_risk__risk_criticality',
        'units_risk__response_decision',
        'units_risk__risk_action',
        'units_risk__budget',
        'units_risk__response_date',
        'units_risk__review_plan',
        'units_risk__comments',
        'units_risk__responsible',
        'units_risk__monitoring',
        'units_risk__risk_criticality_after'
    ).order_by('units_risk__toptask__top_task_description')
    for row in rows:
        row_num += 1
        for col_num in range(len(row)):
            ws.write(row_num, col_num, row[col_num], font_style)

    wb.save(response)
    return response


# METHODE TO EXPORT MEETING REPORT TO WORD FILE
@login_required
@group_required('Team Leader')
def riskRegister_doc_report(request, pk):   
   # document.add_picture('static/img/logos.png', width=Inches(1.25))
    unit=Units.objects.get(id=pk)
    units =Units.objects.filter(Q(id__icontains=pk) & Q(units_risk__risk_name__isnull=False)).values(
        'unit_code',
        'unit_description',
        'units_risk__toptask__top_task_description',
        'units_risk__risk_name',
        'units_risk__risk_cause',
        'units_risk__risk_impact', 
        'units_risk__risk_category',
        'units_risk__risk_status', 
        'units_risk__risk_occuring', 
        'units_risk__risk_rating', 
        'units_risk__risk_criticality',
        'units_risk__response_decision',
        'units_risk__risk_action',
        'units_risk__budget',
        'units_risk__response_date',
        'units_risk__review_plan',
        'units_risk__comments',
        'units_risk__responsible',
        'units_risk__monitoring',
        'units_risk__risk_criticality_after'
    ).order_by('units_risk__toptask__top_task_description')
    
    if unit is not None:
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name ='Times New Roman'
        # FORMAT OF DOCUMENT
        section = document.sections
        section.page_height = Inches(9)
        section.page_width = Inches(12)
        section = document.sections[-1]
        
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height=section.page_height, section.page_width
        
        section.page_width = new_width
        section.page_height = new_height
                
        # TITLE FOR DOCUMENT
        Title_report = document.add_heading('Risk Register submitted by '+str(unit.unit_code), 0)
        Title_report.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # document.add_heading('Performance monitoring report', 0)
        EditReport= document.add_paragraph('Printed on: ')
        EditReport.add_run("%s" % date.today().strftime('%B %d, %Y')).bold = True
        EditReport.add_run(' From UHP website').italic = True
        EditReport.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # PUT THE NAME OF UNIT IN THE DOCUMENT
        AuthorReport = document.add_heading(str(unit.unit_description), level=1)
        AuthorReport.alignment = WD_ALIGN_PARAGRAPH.CENTER
                           
    #  p = document.add_paragraph('A plain paragraph having some ')
    #  p.add_run('bold').bold = True
    #  p.add_run(' and some ')
    #   p.add_run('italic.').italic = True

    #  document.add_heading('Heading, level 1', level=1)
    #  document.add_paragraph('Intense quote', style='IntenseQuote')

    #  document.add_paragraph(
    #  'first item in unordered list', style='ListBullet'
    #  )
    #  document.add_paragraph(
    #  'first item in ordered list', style='ListNumber'
    #  )
       
                                
       # document.add_paragraph('Discussion:', style='ListBullet' )                     
        
   
        table = document.add_table(rows=1, cols=13)
        table.style='Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Top task'
        hdr_cells[1].text = 'Risk name'
        hdr_cells[2].text = 'Risk Cause'
        hdr_cells[3].text = 'Risk Impact'
        hdr_cells[4].text = 'Category'
        hdr_cells[5].text = 'Criticality'
        hdr_cells[6].text = 'Decision'
        hdr_cells[7].text = 'Action'
        hdr_cells[8].text = 'Budget'
        hdr_cells[9].text = 'Date'
        hdr_cells[10].text = 'Review plan'
        hdr_cells[11].text = 'Comments'
        hdr_cells[12].text = 'After action'
        
        for i in range(len(units)):
            row_cells = table.add_row().cells
            row_cells[0].text = str(units[i]['units_risk__toptask__top_task_description'])
            row_cells[1].text = str(units[i]['units_risk__risk_name'])
            row_cells[2].text = str(units[i]['units_risk__risk_cause'])
            row_cells[3].text = str(units[i]['units_risk__risk_impact'])
            row_cells[4].text = str(units[i]['units_risk__risk_category'])
            row_cells[5].text = str(units[i]['units_risk__risk_criticality'])
            row_cells[6].text = str(units[i]['units_risk__response_decision'])
            row_cells[7].text = str(units[i]['units_risk__risk_action'])
            row_cells[8].text = str(units[i]['units_risk__budget'])
            row_cells[9].text = str(units[i]['units_risk__response_date'])
            row_cells[10].text = str(units[i]['units_risk__review_plan'])
            row_cells[11].text = str(units[i]['units_risk__comments'])
            row_cells[12].text = str(units[i]['units_risk__risk_criticality_after'])
            
        document.add_page_break()

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=Risk_Register_'+str(unit.unit_code)+'.docx'
        document.save(response)

        return response
    else:
        return HttpResponse("Failed to Download Risk Register")
    
     

# METHOD TO CREATE SUB ACTIVITIES IN THE SYSTEM
#@permission_required('myuhp.add_operworkplan', raise_exception=True)
@login_required
@group_required('Team Leader')
def wpkl_subActivities_add(request):
    context={}
    form_unit = SelectUnitForm(request.GET or None)
    by_unit =  request.GET.get('by_unit')
    end_date =  request.GET.get('end')
    form = WorkplanForm()
    operworkplans = Operworkplan.objects.all()
    
    if request.method=='GET':
        st=request.GET.get('sub_activity')
        if st!=None:
            operworkplans= Operworkplan.objects.filter(sub_activity__icontains=st)
            
        if by_unit!=None:
            operworkplans = Operworkplan.objects.all().filter(Q(gsmWorkplan__toptask__unit__unit_code__icontains=by_unit)& Q(completion_date__lte=end_date)).order_by("gsmWorkplan")

    
  #  page = request.GET.get('page')
  #  num_of_items = 200
   
 #   paginator = Paginator(operworkplans, num_of_items)
       
 #   try:
 #       operworkplans = paginator.page(page)
 #   except PageNotAnInteger:
 #       page = 1
 #       operworkplans = paginator.page(page)
 #   except EmptyPage:
 #       page = paginator.num_pages
 #       operworkplans = paginator.page(page)
        
 #   context['paginator']=paginator 
    context['operworkplans'] = operworkplans
    context['title'] ='home'
    if request.method == 'POST':
        if 'save' in request.POST:
            pk = request.POST.get('save')
            if not pk:
                form=WorkplanForm(request.POST)
            else:
                operworkplan = Operworkplan.objects.get(id = pk)
                form = WorkplanForm(request.POST ,instance = operworkplan) 
            form.save()
            form = WorkplanForm()
            messages.success(request, 'Sub activity saved or updated successfully')
        elif 'delete' in request.POST:
            pk = request.POST.get('delete')
            operworkplan = Operworkplan.objects.get(id = pk)
            operworkplan.delete()
            messages.success(request, 'Sub activity deleted successfully')
                
        elif 'edit' in request.POST:
            pk = request.POST.get('edit')
            operworkplan = Operworkplan.objects.get(id = pk)
            form = WorkplanForm(instance = operworkplan)
          
    context['form'] = form
    context['form_unit'] = form_unit 
    context['by_unit'] = by_unit
    context['end_date'] = end_date

    return render(request, 'pages/forms/workplans/sub_activity_workplan.html', context)

@login_required
@group_required('Technical officer')
def subActivities_doc_report(request, unit_code): 
    
    data_unit =list(Units.objects.prefetch_related('units_toptask').filter(Q(unit_code__icontains=unit_code)& Q(units_toptask__toptask_gsmWorkplan__lowest_task_description__isnull=False)).distinct())
    data_outputs = Outputworkplan.objects.exclude(outputs_toptask__unit__unit_code__isnull=True).filter(Q(outputs_toptask__unit__unit_code__icontains=unit_code)).prefetch_related('outputs_toptask').distinct().order_by("outputs_toptask__output__output_code").values(
        'outputs_toptask__output__output_description',
        'outputs_toptask__top_task_description',
        'outputs_toptask__toptask_gsmWorkplan__lowest_task_description',
        'outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity',
        'outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__country__country_code',
        'outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__responsable',
        'outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__completion_date',
        'outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name',
        'outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__comments'
    )
    kpis = Kpi.objects.filter(Q(unit__unit_code__icontains=unit_code)).prefetch_related('kpis_kpiAchive','kpis_toptask').distinct()
    
    dta_NumberTask = Units.objects.filter(Q(unit_code__icontains=unit_code)).annotate(
        total_output = Count('units_toptask__output__output_code',distinct=('units_toptask__output__output_code'), filter=Q(units_toptask__output__isnull=False)),
        total_kpi = Count('units_toptask__kpi',distinct=('units_toptask__kpi'), filter=Q(units_toptask__kpi__isnull=False)),
        total_TopTask = Count('units_toptask__top_task_short',distinct=('units_toptask__top_task_short'), filter=Q(units_toptask__top_task_short__isnull=False)),
        total_lowest = Count('units_toptask__toptask_gsmWorkplan__lowest_task_short',distinct=('units_toptask__toptask_gsmWorkplan__lowest_task_short'), filter=Q(units_toptask__toptask_gsmWorkplan__lowest_task_short__isnull=False)),
        total_subactivities = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity',distinct=('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity'), filter=Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False))
        ).values('unit_code','total_output','total_kpi','total_TopTask','total_lowest','total_subactivities').distinct()

    if data_outputs is not None:
        document = Document()
        section = document.sections
        section.page_height = 8
        section.page_width = 11
        section = document.sections[-1]
        
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height=section.page_height, section.page_width
        
        section.page_width = new_width
        section.page_height = new_height
        
       
       # document.add_heading('Briefing note background submitted', 0)
        Title_report = document.add_heading('Operational workplan with the sub-activity submitted by '+str(unit_code), 0)
        Title_report.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # document.add_heading('Performance monitoring report', 0)
        EditReport= document.add_paragraph('Printed on: ')
        EditReport.add_run("%s" % date.today().strftime('%B %d, %Y')).bold = True
        EditReport.add_run(' From UHP website').italic = True
        EditReport.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # PUT THE NAME OF UNIT IN THE DOCUMENT
      #  AuthorReport = document.add_heading(str(briefingProject.unit), level=1)
      #  AuthorReport.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #  p = document.add_paragraph('A plain paragraph having some ')
    #  p.add_run('bold').bold = True
    #  p.add_run(' and some ')
    #   p.add_run('italic.').italic = True

    #  document.add_heading('Heading, level 1', level=1)
    #  document.add_paragraph('Intense quote', style='IntenseQuote')

    #  document.add_paragraph(
    #  'first item in unordered list', style='ListBullet'
    #  )
    #  document.add_paragraph(
    #  'first item in ordered list', style='ListNumber'
    #  )
    #    BriefingTitle = document.add_paragraph('Title of briefing note: ' , style='ListBullet')
    #    BriefingTitle.add_run(str(dta_NumberTask.model.)).bold = True
       
    #    StartDate = document.add_paragraph('Start date: ' , style='ListBullet')
    #    StartDate.add_run(str(briefingProject.start_date)).bold = True
            
    #    EndDate = document.add_paragraph('End date: ' , style='ListBullet')
    #    EndDate.add_run(str(briefingProject.end_date)).bold = True
            
    #    SumbissionDate = document.add_paragraph('Submission date: ' , style='ListBullet')
    #    SumbissionDate.add_run(str(briefingProject.reporting_date)).bold = True
            
        
                      
        for j in range(len(dta_NumberTask)):
            OutputNumber = document.add_paragraph('Number of output: ' , style='ListBullet')
            OutputNumber.add_run(str(dta_NumberTask[j]['total_output'])).bold = True
            
            KpiNumber = document.add_paragraph('Number of Kpis integrated in the workplan: ' , style='ListBullet')
            KpiNumber.add_run(str(dta_NumberTask[j]['total_kpi'])).bold = True
            
            TopTaskNumber = document.add_paragraph('Number of top tasks: ' , style='ListBullet')
            TopTaskNumber.add_run(str(dta_NumberTask[j]['total_TopTask'])).bold = True
            
            LowestTaskNumber = document.add_paragraph('Number of lowest tasks: ' , style='ListBullet')
            LowestTaskNumber.add_run(str(dta_NumberTask[j]['total_lowest'])).bold = True
            
            SubactivityNumber = document.add_paragraph('Number of sub-activity: ' , style='ListBullet')
            SubactivityNumber.add_run(str(dta_NumberTask[j]['total_subactivities'])).bold = True
                     
                     
        table = document.add_table(rows=1, cols=9)
        table.style='Table Grid'
        table.autofit = False
        table.allow_autofit = False
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Output'
        hdr_cells[1].text = 'Top task'
        hdr_cells[2].text = 'Lowest task'
        hdr_cells[3].text = 'Sub activity'
        hdr_cells[4].text = 'Country where supported'
        hdr_cells[5].text = 'Responsible'
        hdr_cells[6].text = 'Completion date'
        hdr_cells[7].text = 'Status'
        hdr_cells[8].text = 'Comments'
       
        for i in range(len(data_outputs)):
            row_cells = table.add_row().cells
            row_cells[0].text = str(data_outputs[i]['outputs_toptask__output__output_description'])
            row_cells[1].text = str(data_outputs[i]['outputs_toptask__top_task_description'])
            row_cells[2].text = str(data_outputs[i]['outputs_toptask__toptask_gsmWorkplan__lowest_task_description'])
            row_cells[3].text = str(data_outputs[i]['outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity'])
           # row_cells[4].text = set(data_outputs[i]['outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__country__country_code'])
            
            for countries in str(data_outputs[i]['outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__country__country_code']):
                if countries not in ('AFRO') and len(countries)>4:
                    row_cells[4].text +=str(countries[0])+str(countries[1])
                else:
                    row_cells[4].text +=str(countries)
   
            row_cells[5].text = str(data_outputs[i]['outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__responsable'])
            row_cells[6].text = str(data_outputs[i]['outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__completion_date'])
            row_cells[7].text = str(data_outputs[i]['outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name'])
            row_cells[8].text = str(data_outputs[i]['outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__comments'])
        
        document.add_page_break()

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=Workplan_'+str(unit_code) +'.docx'
        document.save(response)

        return response
    else:
        return HttpResponse("Failed to Download Operational workplan")
      

# DOWNLOAD MONITORING PERFORMANCE OF WORKPLAN BY UNIT AND BUY COMPLETION DATE 
@login_required
@group_required('Technical officer')
def download_monitoring_progress(request):
    #initialisation des nom
        units={} # Liste des unités
        dataset_byUnit={}
        kpis={}
        line_chart_unit={}
        dt_fig_unit={}
        date_completion__max=timezone.now()
        month_completion__max={}
        notStartedByCountries = []
        df_subActByCountries = {}
        dt_pie_status={}
        dta_NumberTask = {} 
        dta_PercentTask={}
        dt_fig_status_output={}
        dt_fig_status_compl={}
        dt_fig_risk_total={}
        total_risks ={}
    
        form = SelectUnitForm(request.GET or None)
        
        try:
            one_month_ago = timezone.now() + timezone.timedelta(days=15) # under ToDay untill next 15 days
            one_month_next = timezone.now() + timezone.timedelta(days=31) # under ToDay untill next 31 days
            by_unit =  request.GET.get('by_unit') # get the unit name form formular
           # start =  request.GET.get('start')
            end_date =  request.GET.get('end') # get the date putng in the formular
          
           # kpis = Kpi.objects.prefetch_related('kpis_kpiAchive','kpis_gsmWorkplan').distinct()
            units =list(Units.objects.filter(Q(unit_code__icontains=by_unit)).prefetch_related('units_kpi','units_toptask').values_list('unit_description', flat=True).distinct()) # obtain the decription name for the unite code selected
            
            #1. DATASET 01
            dt_fig_unit = Statutworkplan.objects.exclude(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__isnull=True).filter(Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains=by_unit)).annotate(
            total_planned = Count('statuts_operwork__sub_activity', filter=Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__isnull=False)),
            until_thisDay_planned = Count('statuts_operwork__sub_activity', filter=Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__isnull=False) & Q(statuts_operwork__completion_date__lte=one_month_ago)),
            select_day_planned = Count('statuts_operwork__sub_activity', filter=Q(statuts_operwork__sub_activity__isnull=False) & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__isnull=False) & Q(statuts_operwork__completion_date__lte=end_date)),
           # total_activitiy=Case(
           #             statuts_operwork__completion_date__gt=start,
           #             statuts_operwork__completion_date__lt=end,
           #             then="total_planned",
           #             default=Value("0"),
           #             ), 
            ).values('statuts_operwork__gsmWorkplan__toptask__unit__unit_code','statut_name','statuts_operwork__completion_date','total_planned', 'until_thisDay_planned','select_day_planned').order_by('statuts_operwork__completion_date')
            

           #2. DATASET 02
            date_completion__max=max(list(dt_fig_unit.values_list('statuts_operwork__completion_date')))
           
           #3. DATASET 03  
            fig_unit = px.line(dt_fig_unit, x='statuts_operwork__completion_date', y='total_planned', text='total_planned', color='statut_name',labels={'total_planned':'Total of sub activity', 'statuts_operwork__completion_date':'Date', 'statut_name':'Statut name'})
            fig_unit.update_traces(textposition="bottom right")
            line_chart_unit = fig_unit.to_html(full_html=False, include_plotlyjs=False)
            
            #4. DATASET 04
            dt_pie_status= dt_fig_unit.values('statut_name','total_planned','until_thisDay_planned','select_day_planned').aggregate(        
                planed_notStart=Sum('total_planned', filter=Q(statut_name__isnull=False) & Q(statut_name__icontains='Not Started')),
                planed_issue=Sum('total_planned', filter=Q(statut_name__isnull=False) & Q(statut_name__icontains='Stalled')),
                planed_onTrack=Sum('total_planned', filter=Q(statut_name__isnull=False) & Q(statut_name__icontains='On Track')),
                planed_completed=Sum('total_planned', filter=Q(statut_name__isnull=False) & Q(statut_name__icontains='Completed')),
                one_month_ago_activity=Sum('until_thisDay_planned', filter=Q(statut_name__isnull=False)),
                this_period=Sum('select_day_planned', filter=Q(statut_name__isnull=False))
            )
            
            dta_NumberTask = Units.objects.exclude(units_toptask__toptask_gsmWorkplan__lowest_task_short__isnull=True).filter(Q(unit_code__icontains=by_unit) & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__completion_date__lte=end_date)).annotate(
                total_output = Count('units_toptask__output__output_code',distinct=('units_toptask__output__output_code'), filter=Q(units_toptask__output__isnull=False)),
                total_kpi = Count('units_toptask__kpi',distinct=('units_toptask__kpi'), filter=Q(units_toptask__kpi__isnull=False)),
                total_TopTask = Count('units_toptask__top_task_short',distinct=('units_toptask__top_task_short'), filter=Q(units_toptask__top_task_short__isnull=False)),
                total_lowest = Count('units_toptask__toptask_gsmWorkplan__lowest_task_short', filter=Q(units_toptask__toptask_gsmWorkplan__lowest_task_short__isnull=False)),
                total_subactivities = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', filter=Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False)),
                total_subactivity_completed = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', filter=(Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False) & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name__icontains='Completed'))),
                total_subactivity_onTrack = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', filter=(Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False) & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name__icontains='On Track'))),
                total_subactivity_notStart = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', filter=(Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False) & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name__icontains='Not Started'))),
                total_subactivity_issue = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', filter=(Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False) & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name__icontains='Stalled')))
                ).values('unit_code','total_output','total_kpi','total_TopTask','total_lowest','total_subactivities','total_subactivity_completed','total_subactivity_onTrack','total_subactivity_notStart','total_subactivity_issue').distinct()
        
            total_subactivities_period= dta_NumberTask.values('total_subactivities','total_subactivity_completed','total_subactivity_onTrack','total_subactivity_notStart','total_subactivity_issue').aggregate(       
                planed_sub_activity=Sum('total_subactivities')              
            )
            
            dta_PercentTask= dta_NumberTask.values('total_subactivities','total_subactivity_completed','total_subactivity_onTrack','total_subactivity_notStart','total_subactivity_issue').aggregate(       
                planed_notStart=(Sum('total_subactivity_notStart')/Sum('total_subactivities'))*100,
                planed_issue=(Sum('total_subactivity_issue')/Sum('total_subactivities'))*100,
                planed_onTrack=(Sum('total_subactivity_onTrack')/Sum('total_subactivities'))*100,
                planed_completed=(Sum('total_subactivity_completed')/Sum('total_subactivities'))*100
            )
            
          #  dta_PercentTask = pd.DataFrame(dta_PercentTask)
            
        
         
            tables_subactivities = Units.objects.exclude(units_toptask__toptask_gsmWorkplan__lowest_task_short__isnull=True).filter(unit_code__icontains=by_unit).annotate(  
                total_lowest = Count('units_toptask__toptask_gsmWorkplan__lowest_task_short', filter=Q(units_toptask__toptask_gsmWorkplan__lowest_task_short__isnull=False)),
                total_planned = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', filter=Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False)),
                total_completed = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', filter=Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name__icontains='Completed') & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False) & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False)),
                total_OnTrack = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', filter=Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name__icontains='On Track') & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False) & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False)),
                total_Not_Started = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', filter=Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name__icontains='Not Started') & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False)& Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False)),
                total_Issues = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', filter=Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name__icontains='Stalled') & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False)& Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False))
                ).values('unit_code','total_lowest','total_planned','total_completed','total_OnTrack','total_Not_Started','total_Issues').distinct()
          
            completedByCountries= Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__isnull=False) &  Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains=by_unit)& Q(statuts_operwork__completion_date__lte=end_date) & Q(statut_name__icontains='Completed')).annotate(nb_sub_activity=Count('statuts_operwork__sub_activity')).values(
                'statut_name',
                'statuts_operwork__country__country_name',
                'nb_sub_activity'
                ).distinct()
            notStartedByCountries= Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__isnull=False)  & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains=by_unit)& Q(statuts_operwork__completion_date__lte=end_date) & Q(statut_name__icontains='Not Started')).annotate(nb_sub_activity=Count('statuts_operwork__sub_activity')).values(
                #'statut_name',
                'statuts_operwork__country__country_name',
                'nb_sub_activity'
                ).distinct()
            
            subActByCountries= Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False) & Q(statut_name__isnull=False)  & Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains=by_unit)& Q(statuts_operwork__completion_date__lte=end_date)).values(
               # 'statut_name',
                'statuts_operwork__country__country_name'
                ).annotate(nb_sub_activity=Count('statuts_operwork__sub_activity'))
    
      
            notStartedByCountries = pd.DataFrame(notStartedByCountries)
            # parsing the DataFrame in json format. 
            json_records = notStartedByCountries.reset_index().to_json(orient ='records') 
            #data_NumberTask = [] 
            notStartedByCountries = json.loads(json_records) 
            
            #5. DATASET 5
            kpi_results = Kpi.objects.prefetch_related('kpis_kpiAchive').filter(Q(unit__unit_code__icontains=by_unit)& Q(kpis_kpiAchive__report_date__lte=end_date)).values(
                'kpi_description', 'kpis_kpiAchive__report_date', 'kpis_kpiAchive__kpi_baseline', 'kpis_kpiAchive__kpi_target', 'kpis_kpiAchive__report_resut', 'kpis_kpiAchive__report_comment'
            )
         #   datasetTask = GsmWorkplan.objects.exclude(unit__unit_code__isnull=True).filter(Q(unit__unit_code__icontains=by_unit) & Q(gsmWorkplan_operw__completion_date__lte=end_date) & Q(gsmWorkplan_operw__sub_activity__isnull=False) & Q(unit__unit_code__isnull=False)).prefetch_related('gsmWorkplan_operw').distinct().order_by('-gsmWorkplan_operw__completion_date')
            dataset_byUnit = Statutworkplan.objects.exclude(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__isnull=True).filter(Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains=by_unit)& Q(statuts_operwork__completion_date__lte=end_date)& Q(statut_name__in =['Completed','On Track'])).values(
                'statuts_operwork__gsmWorkplan__toptask__output__output_code','statuts_operwork__gsmWorkplan__toptask__top_task_description', 'statuts_operwork__gsmWorkplan__lowest_task_description','statuts_operwork__sub_activity','statuts_operwork__responsable','statuts_operwork__completion_date','statut_name','statuts_operwork__country__country_code','statuts_operwork__comments'
            )
            
            
            dataset_byUnit_Next = Statutworkplan.objects.exclude(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__isnull=True).filter(Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains=by_unit) & Q(statuts_operwork__completion_date__lte=one_month_next) & Q(statut_name__icontains='Not Started')).values(
                'statuts_operwork__gsmWorkplan__toptask__output__output_code','statuts_operwork__gsmWorkplan__toptask__top_task_description', 'statuts_operwork__gsmWorkplan__lowest_task_description','statuts_operwork__sub_activity','statuts_operwork__responsable','statuts_operwork__completion_date','statut_name','statuts_operwork__country__country_code','statuts_operwork__comments'
            )
            
            dt_fig_status_Completed = Statutworkplan.objects.exclude(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__isnull=True).filter(Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains=by_unit)& Q(statut_name__icontains='Completed') & Q(statuts_operwork__completion_date__isnull=False) & Q(statuts_operwork__completion_date__lte=end_date)).annotate(
                nb_sub_activity=Count('statuts_operwork__sub_activity')
            ).values('statuts_operwork__completion_date','nb_sub_activity').order_by('statuts_operwork__completion_date')
            
            
            dt_fig_status_output = Statutworkplan.objects.exclude(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__isnull=True).filter(Q(statuts_operwork__gsmWorkplan__toptask__unit__unit_code__icontains=by_unit) & Q(statuts_operwork__completion_date__isnull=False) & Q(statuts_operwork__completion_date__lte=end_date)).annotate(
                nb_output_111=Count('statuts_operwork__sub_activity',filter=Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='1.1.1')),
                nb_output_311=Count('statuts_operwork__sub_activity',filter=Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.1.1')),
                nb_output_312=Count('statuts_operwork__sub_activity',filter=Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.1.2')),
                nb_output_321=Count('statuts_operwork__sub_activity',filter=Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.2.1')),
                nb_output_322=Count('statuts_operwork__sub_activity',filter=Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.2.2')),
                nb_output_331=Count('statuts_operwork__sub_activity',filter=Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.3.1')),
                nb_output_332=Count('statuts_operwork__sub_activity',filter=Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='3.3.2')),
                nb_output_411=Count('statuts_operwork__sub_activity',filter=Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='4.1.1')),
                nb_output_421=Count('statuts_operwork__sub_activity',filter=Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='4.2.1')),
                nb_output_422=Count('statuts_operwork__sub_activity',filter=Q(statuts_operwork__gsmWorkplan__toptask__output__output_code__icontains='4.2.2'))
            ).values('statut_name','nb_output_111','nb_output_311','nb_output_312','nb_output_321','nb_output_322','nb_output_331','nb_output_332','nb_output_411','nb_output_421','nb_output_422').order_by('-statut_name')
                    
            dt_fig_status_compl = Units.objects.filter(Q(unit_code__icontains=by_unit) & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__completion_date__lte=end_date)).annotate(
                nb_Completed=Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', filter=(Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False) & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name__icontains='Completed'))),
                nb_On_Track = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', filter=(Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False) & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name__icontains='On Track'))),
                nb_Not_Started = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', filter=(Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False) & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name__icontains='Not Started'))),
                nb_Started_Behind_Issues = Count('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity', filter=(Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity__isnull=False) & Q(units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name__icontains='Stalled')))           
            ).values('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__completion_date__month','nb_Completed','nb_On_Track','nb_Not_Started','nb_Started_Behind_Issues').order_by('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__completion_date')
            
            month_completion__max=max(list(dt_fig_status_compl.values_list('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__completion_date__month')))
            
            subactivityByMonth= Statutworkplan.objects.filter(Q(statuts_operwork__sub_activity__isnull=False)).values('statuts_operwork__completion_date__month','statut_name').annotate(nb_sub_activity=Count('statuts_operwork__sub_activity'))
            # xfig=subactivityByDate.values_list('statuts_operwork__completion_date__month', flat=True)
            # yfig=subactivityByDate.values_list('nb_sub_activity', flat=True)
                    #  dt_fig_unit_date = pd.DataFrame(dt_fig_unit_date)       
            try:
                dt_fig_risk_category = RiskIdentification.objects.exclude(unit__unit_code__isnull=True).filter(Q(unit__unit_code__icontains=by_unit)).annotate(
                    nb_risk_Financial=Count('risk_name',filter=Q(risk_category__icontains='Financial')),
                    nb_risk_Political_governance=Count('risk_name',filter=Q(risk_category__icontains='Political_governance')),
                    nb_risk_Reputational=Count('risk_name',filter=Q(risk_category__icontains='Reputational')),
                    nb_risk_Staff_Systems_Structures=Count('risk_name',filter=Q(risk_category__icontains='Staff_Systems_Structures')),
                    nb_risk_Strategic=Count('risk_name',filter=Q(risk_category__icontains='Strategic')),
                    nb_risk_Technical=Count('risk_name',filter=Q(risk_category__icontains='Technical'))
                ).values('risk_criticality','nb_risk_Financial','nb_risk_Political_governance','nb_risk_Reputational','nb_risk_Staff_Systems_Structures','nb_risk_Strategic','nb_risk_Technical').order_by('-risk_criticality')
                

                dt_fig_risk_total = RiskIdentification.objects.filter(Q(unit__unit_code__icontains=by_unit)).annotate(
                    total_output_risk = Count('toptask__output__output_code',distinct=('toptask__output__output_code'), filter=Q(toptask__output__isnull=False)),
                    total_toptask_risk = Count('toptask',distinct=('toptask'), filter=Q(toptask__isnull=False)),
                    total_risk = Count('risk_name', filter=Q(risk_name__isnull=False)),
                    total_risk_Moderate = Count('risk_criticality', filter=(Q(risk_criticality__isnull=False) & Q(risk_criticality__icontains='Moderate'))),
                    total_risk_Critical = Count('risk_criticality', filter=(Q(risk_criticality__isnull=False) & Q(risk_criticality__icontains='Critical'))),
                    total_risk_Very_critical = Count('risk_criticality', filter=(Q(risk_criticality__isnull=False) & Q(risk_criticality__icontains='Very_critical'))) 
                ).values('unit__unit_code','total_output_risk','total_toptask_risk','total_risk','total_risk_Moderate','total_risk_Critical','total_risk_Very_critical').distinct()
            
                total_risks= dt_fig_risk_total.values('total_risk').aggregate(       
                    register_risk_total=Sum('total_risk')              
                )
                
                dta_PercentRisk= dt_fig_risk_total.values('total_output_risk','total_toptask_risk','total_risk','total_risk_Moderate','total_risk_Critical','total_risk_Very_critical').aggregate(       
                    before_Moderate=Sum('total_risk_Moderate')/Sum('total_risks')*100,
                    before_Critical=Sum('total_risk_Critical')/Sum('total_risks')*100,
                    Before_Very_critical=Sum('total_risk_Very_critical')/Sum('total_risks')*100   
                )
            except TypeError:
               response = JsonResponse({"messages":"data of risk management not conform"})
             # return response
                
            finally:
                
                # METHOD TO GENERATE WORD FILE
                if dta_NumberTask is not None:
                    if form.is_valid():
                        number_of_tables = [1]
                        document = Document()
                        style = document.styles['Normal']
                        font = style.font
                        font.name ='Arial'
                        # FORMAT OF DOCUMENT
                        sections = document.sections
                        for section in sections:
                                section.top_margin = Inches(0.95)
                                section.bottom_margin = Inches(0.95)
                                section.left_margin = Inches(0.79)
                                section.right_margin = Inches(0.79)
                            # section.orientation = WD_ORIENT.LANDSCAPE
                        # TITLE FOR DOCUMENT
                        Title_report = document.add_heading('Performance monitoring report', 0)
                        Title_report.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # document.add_heading('Performance monitoring report', 0)
                        EditReport= document.add_paragraph('Printed on: ')
                        EditReport.add_run("%s" % date.today().strftime('%B %d, %Y')).bold = True
                        EditReport.add_run(' From UHP website').italic = True
                        EditReport.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # PUT THE NAME OF UNIT IN THE DOCUMENT
                        AuthorReport = document.add_heading(str(units), level=1)
                        AuthorReport.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        
                        #    UnitName = document.add_paragraph('Unit: ' , style='ListBullet')
                        #   UnitName.add_run(str(units)).bold = True
                        
                        # SUB SECTION OR SUB TITLE IN THE DOCUMENT
                        document.add_paragraph('Operational workplan '+str(end_date), style='IntenseQuote')
                        
                        document.add_paragraph('Number of activities planned up to '+str(end_date), style='ListNumber')  
                        for s in range(len(dta_NumberTask)):
                            NumberOfOutput= document.add_paragraph('Number of output: ' , style='ListBullet')
                            NumberOfOutput.add_run(str(dta_NumberTask[s]['total_output'])).bold = True
                            
                            Kpis = document.add_paragraph('Number of KPIs in the workplan: ' , style='ListBullet')
                            Kpis .add_run(str(dta_NumberTask[s]['total_kpi'])).bold = True
                            
                            NumberOfTopTask= document.add_paragraph('Number of top tasks: ' , style='ListBullet')
                            NumberOfTopTask.add_run(str(dta_NumberTask[s]['total_TopTask'])).bold = True
                            
                            NumberOfLowestTask= document.add_paragraph('Number of lowest tasks: ' , style='ListBullet')
                            NumberOfLowestTask.add_run(str(dta_NumberTask[s]['total_lowest'])).bold = True
                            
                            NumbeOfSubActivity= document.add_paragraph('Number of sub-activities: ' , style='ListBullet')
                            NumbeOfSubActivity.add_run(str(dta_NumberTask[s]['total_subactivities'])).bold = True
                            
                            NumbeOfSubActivityCompleted= document.add_paragraph('Number of sub-activities completed: ' , style='ListBullet')
                            NumbeOfSubActivityCompleted.add_run(str(dta_NumberTask[s]['total_subactivity_completed'])).bold = True
                            
                            NumbeOfSubActivityOntrack= document.add_paragraph('Number of sub-activities on track: ' , style='ListBullet')
                            NumbeOfSubActivityOntrack.add_run(str(dta_NumberTask[s]['total_subactivity_onTrack'])).bold = True
                            
                            NumbeOfSubActivityNotStat= document.add_paragraph('Number of sub-activities not start: ' , style='ListBullet')
                            NumbeOfSubActivityNotStat.add_run(str(dta_NumberTask[s]['total_subactivity_notStart'])).bold = True
                            
                            NumbeOfSubActivityIssu= document.add_paragraph('Number of sub-activities Stalled ' , style='ListBullet')
                            NumbeOfSubActivityIssu.add_run(str(dta_NumberTask[s]['total_subactivity_issue'])).bold = True
                            
                            # str(dta_NumberTask[s]['percent_completed'])
                            #percent_subActivity_completed
                            #total_subactivity_completed
                #        document.add_page_break()
                    
                    #    document.add_paragraph(' ')   
                    
                        document.add_paragraph('Progress on implementation of activities planned to ' +str(end_date), style='ListNumber')  
                        plt.style.use('ggplot') 
                        # GRAPH NUMBER 1   
                        category_names = dt_fig_status_output.values_list('statut_name', flat=True).distinct()
                        results = {
                            'Output 1.1.1': dt_fig_status_output.values_list('nb_output_111', flat=True),
                            'Output 3.1.1': dt_fig_status_output.values_list('nb_output_311', flat=True),
                            'Output 3.1.2': dt_fig_status_output.values_list('nb_output_312', flat=True),
                            'Output 3.2.1': dt_fig_status_output.values_list('nb_output_321', flat=True),
                            'Output 3.2.2': dt_fig_status_output.values_list('nb_output_322', flat=True),
                            'Output 3.3.1': dt_fig_status_output.values_list('nb_output_331', flat=True),
                            'Output 3.3.2': dt_fig_status_output.values_list('nb_output_332', flat=True),
                            'Output 4.1.1': dt_fig_status_output.values_list('nb_output_411', flat=True),
                            'Output 4.2.1': dt_fig_status_output.values_list('nb_output_421', flat=True),
                            'Output 4.2.2': dt_fig_status_output.values_list('nb_output_422', flat=True)
                        }  
                        labels = list(results.keys())
                        data = np.array(list(results.values()))
                        data_cum = data.cumsum(axis=1)
                        category_colors = plt.colormaps['RdYlGn'](
                            np.linspace(0.15, 0.85, data.shape[1]))

                        fig, ax = plt.subplots(figsize=(9.2, 5))
                        ax.invert_yaxis()
                        ax.xaxis.set_visible(False)
                        ax.set_xlim(0, np.sum(data, axis=1).max())

                        for i, (colname, color) in enumerate(zip(category_names, category_colors)):
                            widths = data[:, i]
                            starts = data_cum[:, i] - widths
                            rects = ax.barh(labels, widths, left=starts, height=0.5,
                                            label=colname, color=color)
                            r, g, b, _ = color
                            text_color = 'white' if r * g * b < 0.5 else 'darkgrey'
                            ax.bar_label(rects, label_type='center', color=text_color)
                        ax.legend(ncols=len(category_names), bbox_to_anchor=(0.5, 0., 0.5, 0.5),
                                loc='lower left', fontsize='small')
                        ax.set_title("Number of sub-activities by status and output, n= " +str(list(total_subactivities_period.values())))
                        plt.savefig('test.png', bbox_inches = 'tight')
                        document.add_picture('test.png', width=Inches(6)) 
                        
                        
                        
                    #   document.add_paragraph(' ')   
                        
                    # GRAPH NUMBER 2               
                    
                        species_sub = list(dta_PercentTask.keys())
                        penguin_means_sub = {
                            'Percentage (%)': list(dta_PercentTask.values())
                        #  'Bill Length': (38.79, 48.83, 47.50),
                        #  'Flipper Length': (189.95, 195.82, 217.19),
                        }
                        x = np.arange(len(species_sub))  # the label locations
                        width = 0.25  # the width of the bars
                        multiplier = 0

                        fig, ax = plt.subplots(layout='constrained')
                        for attribute, measurement in penguin_means_sub.items():
                            offset = width * multiplier
                            rects = ax.bar(x + offset, measurement, width, label=attribute)
                            ax.bar_label(rects, padding=3)
                            multiplier += 1

                        # Add some text for labels, title and custom x-axis tick labels, etc.
                        ax.set_ylabel('Percent (%)')
                        ax.set_title("Sub-activities by implementation status (%), n= " +str(list(total_subactivities_period.values())))
                        ax.set_xticks(x + width, species_sub)
                        ax.legend(loc='upper left', ncols=3)
                        ax.set_ylim(0, 100)
                        plt.savefig('test.png', bbox_inches = 'tight')
                        document.add_picture('test.png', width=Inches(6)) 
                        
                        # GRAPH 3
                #     compl_date =dt_fig_status_Completed.values_list('statuts_operwork__completion_date__month', flat=True).distinct()
                #     subactivity_nb = dt_fig_status_Completed.values_list('nb_sub_activity', flat=True)

                #     fig, ax = plt.subplots()
                #     bar_container = ax.bar(compl_date, subactivity_nb)
                #     ax.set(ylabel='Number of sub-acitivity', title='Number of sub-activity by month', ylim=(0, max(subactivity_nb)+1))
                #     ax.bar_label(bar_container,
                                #   fmt=lambda x: f'{x * 1:.1f}',
                #                  padding=3
                #                  )
                        # creating the bar plot
                #     plt.savefig('test.png', bbox_inches = 'tight')
                #     document.add_picture('test.png', width=Inches(6))  
                        
                        
                        # GRAPH 4
                        species =dt_fig_status_compl.values_list('units_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__completion_date__month', flat=True).distinct()
                        sex_counts = {
                            'Completed': np.array(dt_fig_status_compl.values_list('nb_Completed', flat=True)),
                            'On Track': np.array(dt_fig_status_compl.values_list('nb_On_Track', flat=True)),
                            'Not Started': np.array(dt_fig_status_compl.values_list('nb_Not_Started', flat=True)),
                            'Stalled': np.array(dt_fig_status_compl.values_list('nb_Started_Behind_Issues', flat=True))
                        }
                        width = 0.6  # the width of the bars: can also be len(x) sequence


                        fig, ax = plt.subplots()
                        bottom = np.zeros(next(iter(month_completion__max),3))

                        for sex, sex_count in sex_counts.items():
                            p = ax.bar(species, sex_count, width, label=sex, bottom=bottom)
                            bottom += sex_count
                            ax.bar_label(p, label_type='center')

                    #  ax.set_title('Number of penguins by sex')
                        ax.set(ylabel='Number of sub-acitivity',xlabel='Month (Value)', title='Number of sub activities per month until the date of '+str(end_date))
                        ax.legend(loc='upper left', ncols=2)
                        #plt.xticks(rotation=90)
                #        plt.savefig('test.png', bbox_inches = 'tight')
                #        document.add_picture('test.png', width=Inches(6))
                    
                    
                        try:
                            if (not dt_fig_risk_total):
                                raise Exception("No risk register")
        
                            document.add_paragraph('Risk Identification during Operational Planning 2024 - 2025  ' +str(end_date), style='ListNumber')  
                            # GRAPHIQUE 5
                    
                            for q in range(len(dt_fig_risk_total)):
                                NumberOfOutputRisk= document.add_paragraph('Number of output with risk: ' , style='ListBullet')
                                NumberOfOutputRisk.add_run(str(dt_fig_risk_total[q]['total_output_risk'])).bold = True
                                
                                NumberOfToptaskputRisk = document.add_paragraph('Number of top task with risk in the workplan: ' , style='ListBullet')
                                NumberOfToptaskputRisk.add_run(str(dt_fig_risk_total[q]['total_toptask_risk'])).bold = True
                                
                                NumberOfRisk= document.add_paragraph('Number of risk: ' , style='ListBullet')
                                NumberOfRisk.add_run(str(dt_fig_risk_total[q]['total_risk'])).bold = True
                                
                                NumberOfRiskModerate= document.add_paragraph('Number of Moderate risk: ' , style='ListBullet')
                                NumberOfRiskModerate.add_run(str(dt_fig_risk_total[q]['total_risk_Moderate'])).bold = True
                                
                                NumberOfRiskCritical= document.add_paragraph('Number of Critical risk: ' , style='ListBullet')
                                NumberOfRiskCritical.add_run(str(dt_fig_risk_total[q]['total_risk_Critical'])).bold = True

                                NumberOfRiskVeryCritical= document.add_paragraph('Number of Very Critical risk: ' , style='ListBullet')
                                NumberOfRiskVeryCritical.add_run(str(dt_fig_risk_total[q]['total_risk_Very_critical'])).bold = True
                                
                    
                            category_names_criticaly = dt_fig_risk_category.values_list('risk_criticality', flat=True).distinct()
                            results_criticaly = {
                                'Financial': dt_fig_risk_category.values_list('nb_risk_Financial', flat=True),
                                'Political or governance': dt_fig_risk_category.values_list('nb_risk_Political_governance', flat=True),
                                'Reputational': dt_fig_risk_category.values_list('nb_risk_Reputational', flat=True),
                                'Staff, Systems & Structures': dt_fig_risk_category.values_list('nb_risk_Staff_Systems_Structures', flat=True),
                                'Strategic': dt_fig_risk_category.values_list('nb_risk_Strategic', flat=True),
                                'Technical or Public Health': dt_fig_risk_category.values_list('nb_risk_Technical', flat=True),
                            
                            }  
                            labels = list(results_criticaly.keys())
                            data = np.array(list(results_criticaly.values()))
                            data_cum = data.cumsum(axis=1)
                            category_colors = plt.colormaps['RdYlGn'](
                                np.linspace(0.15, 0.85, data.shape[1]))

                            fig, ax = plt.subplots(figsize=(9.2, 5))
                            ax.invert_yaxis()
                            ax.xaxis.set_visible(False)
                            ax.set_xlim(0, np.sum(data, axis=1).max())

                            for i, (colname, color) in enumerate(zip( category_names_criticaly, category_colors)):
                                widths = data[:, i]
                                starts = data_cum[:, i] - widths
                                rects = ax.barh(labels, widths, left=starts, height=0.5,
                                                label=colname, color=color)
                                r, g, b, _ = color
                                text_color = 'white' if r * g * b < 0.5 else 'darkgrey'
                                ax.bar_label(rects, label_type='center', color=text_color)
                            ax.legend(ncols=len(category_names_criticaly), bbox_to_anchor=(0.5, 0., 0.5, 0.5),
                                    loc='lower left', fontsize='small')
                            ax.set_title("Risk Criticality (Probability * Impact) by category, n= " +str(list(total_risks.values())))
                            plt.savefig('test.png', bbox_inches = 'tight')
                            document.add_picture('test.png', width=Inches(6)) 
                            
                                # Gragphique 6
                                
                            species_risk = list(dta_PercentRisk.keys())
                            penguin_means_risk = {
                                'Percentage (%)': list(dta_PercentRisk.values())
                            #  'Bill Length': (38.79, 48.83, 47.50),
                            #  'Flipper Length': (189.95, 195.82, 217.19),
                            }
                            x = np.arange(len(species_risk))  # the label locations
                            width = 0.25  # the width of the bars
                            multiplier = 0

                            fig, ax = plt.subplots(layout='constrained')
                            for attribute, measurement in penguin_means_risk.items():
                                offset = width * multiplier
                                rects = ax.bar(x + offset, measurement, width, label=attribute)
                                ax.bar_label(rects, padding=3)
                                multiplier += 1

                            # Add some text for labels, title and custom x-axis tick labels, etc.
                            ax.set_ylabel('Percent (%)')
                            ax.set_title("Risk management (%), n= " +str(list(total_risks.values())))
                            ax.set_xticks(x + width, species_risk)
                            ax.legend(loc='upper left', ncols=3)
                            ax.set_ylim(0, 100)
                            plt.savefig('test.png', bbox_inches = 'tight')
                            document.add_picture('test.png', width=Inches(6))     
                        except:
                               response = JsonResponse({"messages": "Not Risk"})
                            #   return response
                        finally:

                            document.add_paragraph('Summary of achievement', style='ListNumber')   
                            for k in number_of_tables:
                                    # ADD THE TABLE 1 IN THE DOCUMENT
                                    table = document.add_table(rows=1, cols=5)
                                    table.alignment = WD_TABLE_ALIGNMENT.CENTER #or LEFT or RIGHT
                                    table.style='Table Grid'
                                        
                                    hdr_cells = table.rows[0].cells
                                    hdr_cells[0].text = 'Status'
                                    hdr_cells[1].text = 'Date'
                                    for k in range(len(date_completion__max)):
                                        hdr_cells[2].text = 'Number of sub activity up to '+str(date_completion__max[k])
                                    hdr_cells[3].text = 'Number of sub activity on month ago below '+str(one_month_ago.strftime('%Y-%m-%d'))
                                    hdr_cells[4].text = 'Number of sub activity with completion date '+str(end_date)
                                        
                                    for i in range(len(dt_fig_unit)):
                                            row_cells = table.add_row().cells
                                            row_cells[0].text = str(dt_fig_unit[i]['statut_name'])
                                            row_cells[1].text = str(dt_fig_unit[i]['statuts_operwork__completion_date'])
                                            row_cells[2].text = str(dt_fig_unit[i]['total_planned'])
                                            row_cells[3].text = str(dt_fig_unit[i]['until_thisDay_planned'])
                                            row_cells[4].text = str(dt_fig_unit[i]['select_day_planned'])
                                    # Adding style to a table 
                                    table.style = 'ColorfulShading-Accent5'
                                    
                                # table.style = 'Colorful List'
                                    document.add_paragraph(' ') 
                                    document.add_paragraph('Details of sub-activity completed or on track with completion date '+str(end_date), style='ListNumber')    
                        # chart = dt_pie_status.plot(kind='bar', x='planed_notStart', y='this_period')
                        # plt.savefig(chart.png)
                                
                                    # ADD THE TABLE 2 IN THE DOCUMENT
                                    table_1 = document.add_table(rows=1, cols=9)
                                    table_1.style='Table Grid'
                                    hdr1_cells = table_1.rows[0].cells
                                    hdr1_cells[0].text = 'Output'
                                    hdr1_cells[1].text = 'Top Task'
                                    hdr1_cells[2].text = 'Lowest Task'
                                    hdr1_cells[3].text = 'Sub activity'
                                    hdr1_cells[4].text = 'Responsible'
                                    hdr1_cells[5].text = 'Completion date'
                                    hdr1_cells[6].text = 'Status'
                                    hdr1_cells[7].text = 'Country'
                                    hdr1_cells[8].text = 'Comments'
                                    for p in range(len(dataset_byUnit)):    
                                            row1_cells = table_1.add_row().cells
                                            row1_cells[0].text = str(dataset_byUnit[p]['statuts_operwork__gsmWorkplan__toptask__output__output_code'])
                                            row1_cells[1].text = str(dataset_byUnit[p]['statuts_operwork__gsmWorkplan__toptask__top_task_description'])
                                            row1_cells[2].text = str(dataset_byUnit[p]['statuts_operwork__gsmWorkplan__lowest_task_description'])
                                            row1_cells[3].text = str(dataset_byUnit[p]['statuts_operwork__sub_activity'])
                                            row1_cells[4].text = str(dataset_byUnit[p]['statuts_operwork__responsable'])
                                            row1_cells[5].text = str(dataset_byUnit[p]['statuts_operwork__completion_date'])
                                            row1_cells[6].text = str(dataset_byUnit[p]['statut_name'])
                                            row1_cells[7].text = str(dataset_byUnit[p]['statuts_operwork__country__country_code'])
                                            row1_cells[8].text = str(dataset_byUnit[p]['statuts_operwork__comments'])
                                            
                                    table_1.style = 'ColorfulShading-Accent5'
                                #   table_1.style = 'ColorfulGrid-Accent3'
                                #  table_1.style = 'ColorfulGrid'
                                #  section.start_type = WD_SECTION.NEW_PAGE
                                
                                    document.add_paragraph(' ') 
                                    document.add_paragraph('Sub-activities not started and to be implemented by the following month before '+str(one_month_next.strftime('%Y-%m-%d')), style='ListNumber') 
                                    
                                    # ADD THE TABLE 3 IN THE DOCUMENT
                                    table_2 = document.add_table(rows=1, cols=7)
                                    table_2.style='Table Grid'
                                    hdr2_cells = table_2.rows[0].cells
                                    hdr2_cells[0].text = 'Output'
                                    hdr2_cells[1].text = 'Top Task'
                                    hdr2_cells[2].text = 'Lowest Task'
                                    hdr2_cells[3].text = 'Sub activity'
                                    hdr2_cells[4].text = 'Responsible'
                                    hdr2_cells[5].text = 'Completion date'
                                # hdr2_cells[6].text = 'Status'
                                    hdr2_cells[6].text = 'Country'
                                #  hdr2_cells[8].text = 'Comments'
                                    for p in range(len(dataset_byUnit_Next)):    
                                            row2_cells = table_2.add_row().cells
                                            row2_cells[0].text = str(dataset_byUnit_Next[p]['statuts_operwork__gsmWorkplan__toptask__output__output_code'])
                                            row2_cells[1].text = str(dataset_byUnit_Next[p]['statuts_operwork__gsmWorkplan__toptask__top_task_description'])
                                            row2_cells[2].text = str(dataset_byUnit_Next[p]['statuts_operwork__gsmWorkplan__lowest_task_description'])
                                            row2_cells[3].text = str(dataset_byUnit_Next[p]['statuts_operwork__sub_activity'])
                                            row2_cells[4].text = str(dataset_byUnit_Next[p]['statuts_operwork__responsable'])
                                            row2_cells[5].text = str(dataset_byUnit_Next[p]['statuts_operwork__completion_date'])
                                        #    row2_cells[6].text = str(dataset_byUnit_Next[p]['statut_name'])
                                            row2_cells[6].text = str(dataset_byUnit_Next[p]['statuts_operwork__country__country_code'])
                                    #     row2_cells[8].text = str(dataset_byUnit_Next[p]['statuts_operwork__comments'])
                                    
                                    table_2.style = 'ColorfulShading-Accent2'
                                    
                                    document.add_paragraph(' ') 
                                    
                                    # NEW SUB SECTION OR SUB TITLE IN THE DOCUMENT
                                    document.add_paragraph('Result of the Key Performance Indicators (KPI). Updated on '+str(end_date), style='IntenseQuote')
                                                    # ADD THE TABLE 3 IN THE DOCUMENT
                                    table_3 = document.add_table(rows=1, cols=6)
                                    table_3.style='Table Grid'
                                    hdr3_cells = table_3.rows[0].cells
                                    hdr3_cells[0].text = 'KPI'
                                    hdr3_cells[1].text = 'Reporting date'
                                    hdr3_cells[2].text = 'Baseline'
                                    hdr3_cells[3].text = 'Target'
                                    hdr3_cells[4].text = 'Resut'
                                    hdr3_cells[5].text = 'Comment'
                                
                                    for i in range(len(kpi_results)):    
                                        row3_cells = table_3.add_row().cells
                                        row3_cells[0].text = str(kpi_results[i]['kpi_description'])
                                        row3_cells[1].text = str(kpi_results[i]['kpis_kpiAchive__report_date'])
                                        row3_cells[2].text = str(kpi_results[i]['kpis_kpiAchive__kpi_baseline'])
                                        row3_cells[3].text = str(kpi_results[i]['kpis_kpiAchive__kpi_target'])
                                        row3_cells[4].text = str(kpi_results[i]['kpis_kpiAchive__report_resut'])
                                        row3_cells[5].text = str(kpi_results[i]['kpis_kpiAchive__report_comment'])
                                    
                                    table_3.style = 'ColorfulShading-Accent1'   
                                    
                                    document.add_paragraph(' ')   

                        document.add_page_break()
                        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        response['Content-Disposition'] = 'attachment; filename=Monitoring_report_'+str(by_unit)+'.docx'
                        document.save(response)   

                        return response  
                
        except:
            messages.success(request, 'Please select UNIT')
                    
            data={'form': SelectUnitForm(),
              'units':units,
              'dt_fig_status_compl':dt_fig_status_compl,
              'month_completion__max':month_completion__max,
              'dt_fig_risk_total':dt_fig_risk_total,
              'total_risks':total_risks,
              'dta_PercentTask':dta_PercentTask
              } 
                    
        return render(request,'pages/forms/workplans/download_monitoring.html',data)
    
    
    

 # METHOD TO EXPORT KPI RESULT TO WORD FILE
def download_kpi_result(request, pk): 
    
    kpis = Kpi.objects.get(id=pk) 
    kpiAchieves = KpiAchieve.objects.select_related('kpi').filter(Q(kpi__id__icontains=kpis.id) & Q(kpi__kpi_code__icontains=kpis.kpi_code) & Q(report_resut__isnull=False)).values(
            
            'kpi__expected_result',
            'kpi__kpi_code',
            'kpi__kpi_description',
            'kpi__kpi_country_level',
            'kpi__data_description',
            'kpi__data_source',
            'kpi__collection_methods',
            'kpi__data_frequency',
            'kpi__kpi_link',
            'kpi_baseline',
            'kpi_target',
            'report_date',
            'report_resut',
            'report_comment').order_by('kpi')
    
    #kpiAchieves = Kpi.objects.filter(Q(id__icontains=pk) & Q(kpis_kpiAchive__report_resut__isnull=False)).values(
    #        'expected_result',
    #        'kpi_code',
    #        'kpi_description',
    #        'kpi_country_level',
    #        'data_description',
    #        'data_source',
    #        'collection_methods',
    #        'data_frequency',
    #        'kpi_link',
    #        'kpis_kpiAchive__kpi_baseline',
    #        'kpis_kpiAchive__kpi_target',
    #        'kpis_kpiAchive__report_date',
    #        'kpis_kpiAchive__report_resut',
    #        'kpis_kpiAchive__report_comment').order_by('kpi_code')

    if kpiAchieves is not None:
        document = Document()
        section = document.sections
        section.page_height = 8
        section.page_width = 11
        section = document.sections[-1]
        
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height=section.page_height, section.page_width
        
        section.page_width = new_width
        section.page_height = new_height
        
       # document.add_heading('Briefing note background submitted', 0)
        Title_report = document.add_heading('Key Performance Indicator (KPI) manage by '+str(kpis.unit.unit_code), 0)
        Title_report.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # document.add_heading('Performance monitoring report', 0)
        EditReport= document.add_paragraph('Printed on: ')
        EditReport.add_run("%s" % date.today().strftime('%B %d, %Y')).bold = True
        EditReport.add_run(' From UHP website').italic = True
        EditReport.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        KPi_name = document.add_heading('KPI: '+str(kpis.kpi_description), 1)
        KPi_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
                      
        
        DataDescritption = document.add_paragraph('Data description: ' , style='ListBullet')
        DataDescritption.add_run(str(kpis.data_description)).bold = True
        
        DataSource = document.add_paragraph('Data source: ' , style='ListBullet')
        DataSource.add_run(str(kpis.data_source)).bold = True
        
        CollectionMethod = document.add_paragraph('Collection methods: ' , style='ListBullet')
        CollectionMethod.add_run(str(kpis.collection_methods)).bold = True
        
        DataFrequency = document.add_paragraph('Data frequency: ' , style='ListBullet')
        DataFrequency.add_run(str(kpis.data_frequency)).bold = True
        
        LinkedWith = document.add_paragraph('Kpi linked with: ' , style='ListBullet')
        LinkedWith.add_run(str(kpis.kpi_link)).bold = True
                     
                     
        table = document.add_table(rows=1, cols=5)
        table.style='Table Grid'
        table.autofit = False
        table.allow_autofit = False
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Report date'
        hdr_cells[1].text = 'Baseline of the KPI'
        hdr_cells[2].text = 'Target of the KPI'
        hdr_cells[3].text = 'Result of the KPI'
        hdr_cells[4].text = 'Comment'
       
        for i in range(len(kpiAchieves)):
            row_cells = table.add_row().cells
            row_cells[0].text = str(kpiAchieves[i]['report_date'])
            row_cells[1].text = str(kpiAchieves[i]['kpi_baseline'])
            row_cells[2].text = str(kpiAchieves[i]['kpi_target'])
            row_cells[3].text = str(kpiAchieves[i]['report_resut'])
            row_cells[4].text = str(kpiAchieves[i]['report_comment'])
        
        document.add_page_break()

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=UHP_Result_'+str(kpis.kpi_code) +'.docx'
        document.save(response)

        return response
    else:
        return HttpResponse("Failed to Download Kpi result") 

# METHOD TO EXPORT WORKPLAN FOR EACH UNIT IN EXCEL FORMAT
@login_required
@group_required('Technical officer')
def export_to_excel_workplan(request,by_unit,end_date ):

    data_outputs = Outputworkplan.objects.exclude(outputs_toptask__unit__unit_code__isnull=True).filter(Q(outputs_toptask__unit__unit_code__icontains=by_unit) & Q(outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__completion_date__lte=end_date)).prefetch_related('outputs_toptask').distinct().order_by("outputs_toptask__output__output_code").values(
                'outputs_toptask__output__output_code',
                'outputs_toptask__top_task_description',
                'outputs_toptask__toptask_gsmWorkplan__lowest_task_description',
                'outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity',
                'outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__country__country_code',
                'outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__responsable',
                'outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__coworkers',
                'outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__completion_date',
                'outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name',
                'outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__comments'
            ).order_by("outputs_toptask__output__output_code")
    
    color_scale_rule = ColorScaleRule(start_type="formula",
                                        start_value="Not Started",
                                        start_color="00FF0000",  # Red
                                        mid_type="formula",
                                        mid_value="On Track",
                                        mid_color="00FFFF00",  # Yellow
                                        end_type="formula",
                                        end_value="Completed",
                                        end_color="0000FF00")  # Green
    
    red_background = PatternFill(fgColor="ff00ff")
    diff_style = DifferentialStyle(fill=red_background)
    rule = Rule(type="expression", dxf=diff_style)
    rule.formula = ['$I4="Not Started"']

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="'+'Workplan_'+ by_unit+"_"+end_date +'.xlsx"'  
    workbook=Workbook()
    worksheet = workbook.active
    
        # Edit page Setup
    worksheet.page_setup.orientation = worksheet.ORIENTATION_LANDSCAPE
    worksheet.page_setup.paperSize = worksheet.PAPERSIZE_TABLOID
    worksheet.page_setup.fitToHeight = 0
    worksheet.page_setup.fitToWidth = 1
    
    worksheet.merge_cells('A1:D1')
    worksheet.merge_cells('A2:D2')
    worksheet.freeze_panes = "D4"
    worksheet.auto_filter.ref = "A3:J300"
    worksheet["F1"] = "TOTAL"
    worksheet["G1"] = '=COUNTA(E4:E200)'
    worksheet["G1"].font = Font(bold=True,color="0000ff", size=14)
    worksheet["G1"].alignment = Alignment(horizontal ="center", vertical="center")
    worksheet["F2"] = "COMPLETED"
    worksheet["G2"] = '=COUNTIF(I4:I200, "Completed")'
    worksheet["G2"].font = Font(bold=True, color="00b300", size=14)
    worksheet["G2"].alignment = Alignment(horizontal ="center", vertical="center")
    worksheet["H1"] = "NOT STARTED"
    worksheet["I1"] = '=COUNTIF(I4:I200, "Not Started")'
    worksheet["I1"].font = Font(bold=True, color="00FF0000", size=14)
    worksheet["I1"].alignment = Alignment(horizontal ="center", vertical="center")
    worksheet["H2"] = "ON TRACK"
    worksheet["I2"] = '=COUNTIF(I4:I200, "On Track")'
    worksheet["I2"].font = Font(bold=True, color="ff00ff", size=14)
    worksheet["I2"].alignment = Alignment(horizontal ="center", vertical="center")
    worksheet.conditional_formatting.add("I4:I200", color_scale_rule)
    worksheet.conditional_formatting.add("A4:J200", rule)
    first_cell = worksheet['A1']
    first_cell.value ="Operational workplan with the sub-activity submitted by " +""+ by_unit +" and completion date: "+end_date
    first_cell.fill = PatternFill("solid",fgColor="246ba1")
    first_cell.font =  Font(bold=True, color="F7F6FA", size=16)
    first_cell.alignment = Alignment(horizontal ="center", vertical="center")
     
    second_cell = worksheet['A2']
    second_cell.value = by_unit
    second_cell.font = Font(bold=True, color="246ba1", size=16)
    second_cell.alignment = Alignment(horizontal ="center", vertical="center")
    
    # Name of Excel sheet
    worksheet.title='Workplan_' + "" + by_unit+"_"+end_date
  
     # Define the titles for columns
    columns = ['Output','Top task','Lowest task','Sub activity','Country where supported','Responsible','Co workers','Completion date','Status',
                'Comments']
    row_num = 3
    
    # Assign the titles for each cell of the header
    for col_num, column_title in enumerate(columns, 1):
        cell = worksheet.cell(row=row_num, column=col_num)
        cell.value= column_title
        cell.fill = PatternFill("solid", fgColor="50c878")
        cell.font = Font(bold=True, color="F7F6FA")
        third_cell=worksheet['D3']
        third_cell.alignment =Alignment(horizontal="center", vertical="center")
        
    for workplan in data_outputs:
        row_num +=1
        
        #Define the data for each cell in the row
        row=[workplan['outputs_toptask__output__output_code'],
             workplan['outputs_toptask__top_task_description'],
             workplan['outputs_toptask__toptask_gsmWorkplan__lowest_task_description'],
             workplan['outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__sub_activity'],
             workplan['outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__country__country_code'],
             workplan['outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__responsable'],
             workplan['outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__coworkers'],
             workplan['outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__completion_date'],
             workplan['outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__statut_name__statut_name'],
             workplan['outputs_toptask__toptask_gsmWorkplan__gsmWorkplan_operw__comments']
             ]
          #Assign the data for each cell in the row
        for col_num, cell_value in enumerate(row, 1):
            cell = worksheet.cell(row=row_num, column=col_num)
            cell.value = cell_value
            cell.alignment =Alignment(vertical="center")
    #          if isinstance(cell_value, decimal.Decimal):
    #              cell.number_format = numbers.FORMAT_NUMBER_COMMA_SEPARATED1
                  
    workbook.save(response)
    return response




# METHOD TO EXPORT KPI RESULTS FOR EACH UNIT IN EXCEL FORMAT
@login_required
@group_required('Technical officer')
def export_to_excel_kpiResults(request,by_unit,end_date):
    
    data_kpi_results = Kpi.objects.exclude(unit__unit_code__isnull=True).filter(Q(unit__unit_code__icontains=by_unit)).prefetch_related('kpis_kpiAchive').order_by("kpis_kpiAchive__report_date").values(
                'unit__unit_code',
                'expected_result',
                'kpi_description',
                'kpi_country_level',
                'data_description',
                'data_source',
                'collection_methods',
                'data_frequency',
                'kpi_link',
                'kpis_kpiAchive__kpi_baseline',
                'kpis_kpiAchive__kpi_target',
                'kpis_kpiAchive__report_date',
                'kpis_kpiAchive__report_resut',
                'kpis_kpiAchive__report_comment'
            ).order_by("kpi_description")
    
    
    
  
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="'+'KPI_results_'+ by_unit+"_"+end_date +'.xlsx"'  
    workbook=Workbook()
    worksheet = workbook.active
    
    red_background = PatternFill(fgColor="00FFFF00")
    diff_style = DifferentialStyle(fill=red_background)
    rule = Rule(type="expression", dxf=diff_style)
    rule.formula = ["$H4<3"]
    
    color_scale_rule = ColorScaleRule(start_type="num",
                                    start_value=1,
                                    start_color="00FFCC99",  # Red
                                    mid_type="num",
                                    mid_value=2,
                                    mid_color="00FFFF00",  # Yellow
                                    end_type="num",
                                    end_value=5,
                                    end_color="0000FF00")  # Green
    # Edit page Setup
    worksheet.page_setup.orientation = worksheet.ORIENTATION_LANDSCAPE
    worksheet.page_setup.paperSize = worksheet.PAPERSIZE_TABLOID
    worksheet.page_setup.fitToHeight = 0
    worksheet.page_setup.fitToWidth = 1
    
    worksheet.merge_cells('A1:D1')
    worksheet.merge_cells('A2:D2')
    worksheet.freeze_panes = "D4"
    worksheet.auto_filter.ref = "A3:J300"
    worksheet["F1"] = "TOTAL KPI"
    worksheet["G1"] = '=COUNTA(C4:C200)'
    worksheet["G1"].font = Font(bold=True,color="0000ff", size=14)
    worksheet["G1"].alignment = Alignment(horizontal ="center", vertical="center")
    worksheet["F2"] = "KPI WITH RESULTS"
    worksheet["G2"] = '=COUNTA(M4:M200)'
    worksheet["G2"].font = Font(bold=True, color="00b300", size=14, italic=True)
    worksheet["G2"].alignment = Alignment(horizontal ="center", vertical="center")
 #   worksheet["H1"] = "NOT STARTED"
 #   worksheet["I1"] = '=COUNTIF(I4:I200, "Not Started")'
 #   worksheet["I1"].font = Font(bold=True, color="00FF0000", size=14)
 #   worksheet["I1"].alignment = Alignment(horizontal ="center", vertical="center")
 #   worksheet["H2"] = "ON TRACK"
 #   worksheet["I2"] = '=COUNTIF(I4:I200, "On Track")'
 #   worksheet["I2"].font = Font(bold=True, color="ff00ff", size=14)
 #   worksheet["I2"].alignment = Alignment(horizontal ="center", vertical="center")
    worksheet.conditional_formatting.add("C4:N10", rule)
    worksheet.conditional_formatting.add("G2", color_scale_rule)
   
    first_cell = worksheet['A1']
    first_cell.value ="Key Performance Indicator (KPI) manage by " +""+ by_unit +" and report date: "+end_date
    first_cell.fill = PatternFill("solid",fgColor="246ba1")
    first_cell.font =  Font(bold=True, color="F7F6FA", size=16)
    first_cell.alignment = Alignment(horizontal ="center", vertical="center")
     
    second_cell = worksheet['A2']
    second_cell.value = by_unit
    second_cell.font = Font(bold=True, color="246ba1", size=16)
    second_cell.alignment = Alignment(horizontal ="center", vertical="center")
        
    worksheet.title='Result_Kpi_' + "" + by_unit+"_"+end_date
     # Define the titles for columns
    columns = ['Unit code',
                'Output description',
                'Kpi AFRO level',
                'Kpi country level',
                'Data description',
                'Data source',
                'Collection methods',
                'Data frequency',
                'Kpi link',
                'Baseline of the AFRO KPI',
                'Target of the AFRO KPI',
                'Report date AFRO KPI',
                'Result of the AFRO KPI',
                'Comment']
    row_num = 3
    
    # Assign the titles for each cell of the header
    for col_num, column_title in enumerate(columns, 1):
        cell = worksheet.cell(row=row_num, column=col_num)
        cell.value= column_title
        cell.fill = PatternFill("solid", fgColor="50c878")
        cell.font = Font(bold=True, color="F7F6FA")
        cell.alignment = Alignment(horizontal ="center", vertical="center",wrap_text=False)
        third_cell=worksheet['D3']
        third_cell.alignment =Alignment(horizontal="center", vertical="center",wrap_text=False)
        
    for kpiresult in data_kpi_results:
        row_num +=1    
        #Define the data for each cell in the row
        row=[kpiresult['unit__unit_code'],
             kpiresult['expected_result'],
             kpiresult['kpi_description'],
             kpiresult['kpi_country_level'],
             kpiresult['data_description'],
             kpiresult['data_source'],
             kpiresult['collection_methods'],
             kpiresult['data_frequency'],
             kpiresult['kpi_link'],
             kpiresult['kpis_kpiAchive__kpi_baseline'],
             kpiresult['kpis_kpiAchive__kpi_target'],
             kpiresult['kpis_kpiAchive__report_date'],
             kpiresult['kpis_kpiAchive__report_resut'],
             kpiresult['kpis_kpiAchive__report_comment']
             ]
          #Assign the data for each cell in the row
        for col_num, cell_value in enumerate(row, 1):
            cell = worksheet.cell(row=row_num, column=col_num)
            cell.value = cell_value
            cell.alignment =Alignment(vertical="center",wrap_text=False)
    #          if isinstance(cell_value, decimal.Decimal):
    #              cell.number_format = numbers.FORMAT_NUMBER_COMMA_SEPARATED1
    workbook.save(response)
    return response


# METHOD TO EXPORT RISK REGISTER FOR EACH UNIT IN EXCEL FORMAT
@login_required
@group_required('Team Leader')
def export_to_excel_riskRegister(request,by_unit,end_date):
         
    data_risk_register = RiskIdentification.objects.all().filter(Q(unit__unit_code__icontains=by_unit)& Q(response_date__lte=end_date)).values(
               "unit__unit_code", 
               "toptask__top_task_short", 
               "risk_name",
               "risk_cause", 
               "risk_impact", 
               "risk_category", 
               "risk_status",
               "risk_occuring",
               "risk_rating",
               "risk_criticality",
               "response_decision",
               "risk_action",
               "budget",
               "response_date",
               "review_plan",
               "comments",
               "responsible",
               "monitoring",
               "risk_criticality_after"
            ).order_by("toptask")
    
    
    # Create fill

    color_scale_rule = ColorScaleRule(start_type="num",
                                      start_value=5,
                                      start_color="00FF0000",  # Red
                                      mid_type="num",
                                      mid_value=10,
                                      mid_color="00FFFF00",  # Yellow
                                      end_type="num",
                                      end_value=80,
                                      end_color="0000FF00")  # Green
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="'+'Risk_register_'+ by_unit+"_"+end_date +'.xlsx"'  
    workbook=Workbook()
    worksheet = workbook.active
    
    # Edit page Setup
    worksheet.page_setup.orientation = worksheet.ORIENTATION_LANDSCAPE
    worksheet.page_setup.paperSize = worksheet.PAPERSIZE_TABLOID
    worksheet.page_setup.fitToHeight = 0
    worksheet.page_setup.fitToWidth = 1
    
    worksheet.merge_cells('A1:D1')
    worksheet.merge_cells('A2:D2')
    worksheet.freeze_panes = "D4"
    worksheet.auto_filter.ref = "A3:J300"
    worksheet["F1"] = "TOTAL RISK"
    worksheet["G1"] = '=COUNTA(C4:C200)'
    worksheet["G1"].font = Font(bold=True,color="0000ff", size=14)
    worksheet["G1"].alignment = Alignment(horizontal ="center", vertical="center")
    worksheet["F2"] = "NEW RISK"
    worksheet["G2"] = '=COUNTIF(G4:G200, "New")'
    worksheet["G2"].font = Font(bold=True, color="00FF0000", size=14, italic=True)
    worksheet["G2"].alignment = Alignment(horizontal ="center", vertical="center")
    worksheet["H2"] = "EXISTING RISK"
    worksheet["I2"] = '=COUNTIF(G4:G200, "Existing")'
    worksheet["I2"].font = Font(bold=True, color="00993300", size=14, italic=True)
    worksheet["I2"].alignment = Alignment(horizontal ="center", vertical="center")
    worksheet["J2"] = "CLOSED"
    worksheet["K2"] = '=COUNTIF(G4:G200, "closed")'
    worksheet["K2"].font = Font(bold=True, color="ff00ff", size=14, italic=True)
    worksheet["K2"].alignment = Alignment(horizontal ="center", vertical="center")
    
    worksheet["L2"] = "MODERATE"
    worksheet["M2"] = '=COUNTIF(J4:J200, "Moderate")'
    worksheet["M2"].font = Font(bold=True, color="0000CCFF", size=14, italic=True)
    worksheet["M2"].alignment = Alignment(horizontal ="center", vertical="center")
    
    worksheet["N2"] = "CRITICAL"
    worksheet["O2"] = '=COUNTIF(J4:J200, "Critical")'
    worksheet["O2"].font = Font(bold=True, color="00993366", size=14, italic=True)
    worksheet["O2"].alignment = Alignment(horizontal ="center", vertical="center")
    
    worksheet["P2"] = "VERY CRITICAL"
    worksheet["Q2"] = '=COUNTIF(J4:J200, "Very_critical")'
    worksheet["Q2"].font = Font(bold=True, color="00FF6600", size=14, italic=True)
    worksheet["Q2"].alignment = Alignment(horizontal ="center", vertical="center")
    
    worksheet["R1"] = "MONITORING (%)"
    worksheet["S1"] ='=AVERAGE(R4:R20)'
    worksheet["S1"].font = Font(bold=True, color="0000FF", size=14, italic=True)
    worksheet["S1"].alignment = Alignment(horizontal ="center", vertical="center")
    #worksheet.cell['S1'].number_format = '0.00%'
    worksheet.conditional_formatting.add("S1", color_scale_rule)
    
   
  
  # Again, let's add this gradient to the star ratings, column "H"
    worksheet.conditional_formatting.add("R4:R20", color_scale_rule)
   
    first_cell = worksheet['A1']
    first_cell.value ="Risk Register submitted by " +""+ by_unit +" and mitigate report date: "+end_date
    first_cell.fill = PatternFill("solid",fgColor="246ba1")
    first_cell.font =  Font(bold=True, color="F7F6FA", size=16)
    first_cell.alignment = Alignment(horizontal ="center", vertical="center")
     
    second_cell = worksheet['A2']
    second_cell.value = by_unit
    second_cell.font = Font(bold=True, color="246ba1", size=16)
    second_cell.alignment = Alignment(horizontal ="center", vertical="center")
    
   
        
    worksheet.title='Risk_register_' + "" + by_unit+"_"+end_date
     # Define the titles for columns
    columns = [
               'Unit', 
               'Top task',
               'Risk name',
               'Risk Cause',
               'Risk Impact', 
               'Category', 
               'Status',
               'Occuring',
               'Rating',
               'Criticality',
               'Decision',
               'Action',
               'Budget',
               'Date',
               'Review plan',
               'Comments',
               'Responsible',
               'Monitoring (%)',
               'After action'
                ]
 
    row_num = 3
    
    # Assign the titles for each cell of the header
    for col_num, column_title in enumerate(columns, 1):
        cell = worksheet.cell(row=row_num, column=col_num)
        cell.value= column_title
        cell.fill = PatternFill("solid", fgColor="50c878")
        cell.font = Font(bold=True, color="F7F6FA")
        cell.alignment = Alignment(horizontal ="center", vertical="center",wrap_text=False)
        third_cell=worksheet['D3']
        third_cell.alignment =Alignment(horizontal="center", vertical="center",wrap_text=False)
        
    for riskRegister in data_risk_register:
        row_num +=1    
        #Define the data for each cell in the row
        row=[riskRegister['unit__unit_code'],
             riskRegister['toptask__top_task_short'],
             riskRegister['risk_name'],
             riskRegister['risk_cause'],
             riskRegister['risk_impact'],
             riskRegister['risk_category'],
             riskRegister['risk_status'],
             riskRegister['risk_occuring'],
             riskRegister['risk_rating'],
             riskRegister['risk_criticality'],
             riskRegister['response_decision'],
             riskRegister['risk_action'],
             riskRegister['budget'],
             riskRegister['response_date'],
             riskRegister['review_plan'],
             riskRegister['comments'],
             riskRegister['responsible'],
             riskRegister['monitoring'],
             riskRegister['risk_criticality_after']
             ]
        
          #Assign the data for each cell in the row
        for col_num, cell_value in enumerate(row, 1):
            cell = worksheet.cell(row=row_num, column=col_num)
            cell.value = cell_value
            cell.alignment =Alignment(vertical="center",wrap_text=False)
    #          if isinstance(cell_value, decimal.Decimal):
    #              cell.number_format = numbers.FORMAT_NUMBER_COMMA_SEPARATED1
    workbook.save(response)
    return response


# METHOD TO EXPORT SURVEY DATASET FOR EACH UNIT IN EXCEL FORMAT
def export_to_excel_survey_dataset(request,by_survey,end_day):
         
    survey_project=SurveyProject.objects.get(id=by_survey)
     
    data_survey = SurveyProject.objects.all().prefetch_related('project_surveyData').filter(Q(id__icontains=by_survey)& Q(end_date__lte=end_day)).values(
               "responsible", 
               "title_surv", 
               "start_date",
               "end_date", 
               "location_survey", 
               "project_surveyData__quest_code", 
               "project_surveyData__question",
               "project_surveyData__response_text",
               "project_surveyData__response_num",
               "project_surveyData__level_1",
               "project_surveyData__level_2"
            ).order_by("title_surv")
    
   # number_register =data_survey.annotate(total_sample=Count('project_surveyData__level_1')).values('total_sample')
    
   
    # Create fill

    color_scale_rule = ColorScaleRule(start_type="num",
                                      start_value=5,
                                      start_color="00FF0000",  # Red
                                      mid_type="num",
                                      mid_value=10,
                                      mid_color="00FFFF00",  # Yellow
                                      end_type="num",
                                      end_value=80,
                                      end_color="0000FF00")  # Green
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="'+'Survey_dataset_'+ survey_project.title_surv+"_"+end_day +'.xlsx"'  
    workbook=Workbook()
    worksheet = workbook.active
    
    # Edit page Setup
    worksheet.page_setup.orientation = worksheet.ORIENTATION_LANDSCAPE
    worksheet.page_setup.paperSize = worksheet.PAPERSIZE_TABLOID
    worksheet.page_setup.fitToHeight = 0
    worksheet.page_setup.fitToWidth = 1
    
    worksheet.merge_cells('A1:F1')
    worksheet.merge_cells('A2:F2')
    worksheet.freeze_panes = "F4"
    worksheet.auto_filter.ref = "A3:J5000"
    worksheet["H1"] = "TOTAL REGISTER"
    worksheet["I1"] = '=COUNTA(C4:C5000)'
    worksheet["I1"].font = Font(bold=True,color="0000ff", size=14)
    worksheet["I1"].alignment = Alignment(horizontal ="center", vertical="center")
 #   worksheet["F2"] = "NEW RISK"
 #   worksheet["G2"] = '=COUNTIF(G4:G200, "New")'
 #   worksheet["G2"].font = Font(bold=True, color="00FF0000", size=14, italic=True)
 #   worksheet["G2"].alignment = Alignment(horizontal ="center", vertical="center")
 #   worksheet["H2"] = "EXISTING RISK"
 #   worksheet["I2"] = '=COUNTIF(G4:G200, "Existing")'
 #   worksheet["I2"].font = Font(bold=True, color="00993300", size=14, italic=True)
 #   worksheet["I2"].alignment = Alignment(horizontal ="center", vertical="center")
 #   worksheet["J2"] = "CLOSED"
 #   worksheet["K2"] = '=COUNTIF(G4:G200, "closed")'
 #   worksheet["K2"].font = Font(bold=True, color="ff00ff", size=14, italic=True)
 #   worksheet["K2"].alignment = Alignment(horizontal ="center", vertical="center")
    
 #   worksheet["L2"] = "MODERATE"
 #   worksheet["M2"] = '=COUNTIF(J4:J200, "Moderate")'
 #   worksheet["M2"].font = Font(bold=True, color="0000CCFF", size=14, italic=True)
 #   worksheet["M2"].alignment = Alignment(horizontal ="center", vertical="center")
    
 #   worksheet["N2"] = "CRITICAL"
 #   worksheet["O2"] = '=COUNTIF(J4:J200, "Critical")'
 #   worksheet["O2"].font = Font(bold=True, color="00993366", size=14, italic=True)
 #   worksheet["O2"].alignment = Alignment(horizontal ="center", vertical="center")
    
 #   worksheet["P2"] = "VERY CRITICAL"
 ##   worksheet["Q2"] = '=COUNTIF(J4:J200, "Very_critical")'
 #   worksheet["Q2"].font = Font(bold=True, color="00FF6600", size=14, italic=True)
 #   worksheet["Q2"].alignment = Alignment(horizontal ="center", vertical="center")
    
 #  worksheet["R1"] = "MONITORING (%)"
 #   worksheet["S1"] ='=AVERAGE(R4:R20)'
 #   worksheet["S1"].font = Font(bold=True, color="0000FF", size=14, italic=True)
 #   worksheet["S1"].alignment = Alignment(horizontal ="center", vertical="center")
    #worksheet.cell['S1'].number_format = '0.00%'
 #   worksheet.conditional_formatting.add("S1", color_scale_rule)
    
   
  
  # Again, let's add this gradient to the star ratings, column "H"
  #  worksheet.conditional_formatting.add("R4:R20", color_scale_rule)
   
    first_cell = worksheet['A1']
    first_cell.value ="Dataset for survey: " +""+ survey_project.title_surv +" and End date: "+end_day
    first_cell.fill = PatternFill("solid",fgColor="246ba1")
    first_cell.font =  Font(bold=True, color="F7F6FA", size=16)
    first_cell.alignment = Alignment(horizontal ="center", vertical="center")
     
    second_cell = worksheet['A2']
    second_cell.value = survey_project.responsible
    second_cell.font = Font(bold=True, color="246ba1", size=16)
    second_cell.alignment = Alignment(horizontal ="center", vertical="center")
    
   
        
    worksheet.title='Survey_dataset_' + "" +"_"+end_day
     # Define the titles for columns
    columns = [
               "Organization", 
               "Survey title", 
               "Start date",
               "End date", 
               "Survey location", 
               "Question code", 
               "Questions",
               "Response in text format",
               "Response in integer format",
               "First area of data",
               "Seconde area of data'"
                ]
 
    row_num = 3
    
    # Assign the titles for each cell of the header
    for col_num, column_title in enumerate(columns, 1):
        cell = worksheet.cell(row=row_num, column=col_num)
        cell.value= column_title
        cell.fill = PatternFill("solid", fgColor="50c878")
        cell.font = Font(bold=True, color="F7F6FA")
        cell.alignment = Alignment(horizontal ="center", vertical="center",wrap_text=False)
        third_cell=worksheet['F3']
        third_cell.alignment =Alignment(horizontal="center", vertical="center",wrap_text=False)
        
    for surveyDataset in data_survey:
        row_num +=1    
        #Define the data for each cell in the row
        row=[surveyDataset['responsible'],
             surveyDataset['title_surv'],
             surveyDataset['start_date'],
             surveyDataset['end_date'],
             surveyDataset['location_survey'],
             surveyDataset['project_surveyData__quest_code'],
             surveyDataset['project_surveyData__question'],
             surveyDataset['project_surveyData__response_text'],
             surveyDataset['project_surveyData__response_num'],
             surveyDataset['project_surveyData__level_1'],
             surveyDataset['project_surveyData__level_2']
             ]
               
          #Assign the data for each cell in the row
        for col_num, cell_value in enumerate(row, 1):
            cell = worksheet.cell(row=row_num, column=col_num)
            cell.value = cell_value
            cell.alignment =Alignment(vertical="center",wrap_text=False)
    #          if isinstance(cell_value, decimal.Decimal):
    #              cell.number_format = numbers.FORMAT_NUMBER_COMMA_SEPARATED1
    workbook.save(response)
    return response
